#!/usr/bin/env python3
"""Build a DOCX English correction handbook from exam correction PDFs.

The generator intentionally scopes extraction to the essay section headed
`作文（占 20 分）` so translation and choice-question feedback cannot leak into
the manual corpus.
"""

from __future__ import annotations

import csv
import difflib
import hashlib
import json
import math
import os
import re
import urllib.error
import urllib.parse
import urllib.request
from collections import Counter, defaultdict
from dataclasses import asdict, dataclass
from datetime import date
from pathlib import Path
from statistics import mean, median
from typing import Callable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader


ROOT = Path("/Users/philo/Documents/英語正則POC")
PDF_DIR = ROOT / "站前模D-20260505-results"
OUTPUT_DIR = ROOT / "output"
MANUAL_PATH = OUTPUT_DIR / "114下學測作文英語正則手冊_站前模D.docx"
SUMMARY_CSV = OUTPUT_DIR / "essay_summary_站前模D.csv"
ITEMS_CSV = OUTPUT_DIR / "essay_correction_items_站前模D.csv"
ITEMS_JSON = OUTPUT_DIR / "essay_correction_items_站前模D.json"
BATCH_NAME = "站前模D 2026-05-05"
ProgressCallback = Callable[[dict], None]
AI_ANALYSIS_STATUS: dict[str, str] = {
    "mode": "rules",
    "provider": "",
    "model": "",
    "error": "",
}
GEMINI_ENDPOINT = "https://generativelanguage.googleapis.com/v1beta"

PRESET = {
    "name": "compact_reference_guide",
    "page": {
        "width": Inches(8.5),
        "height": Inches(11),
        "margin": Inches(1),
        "header_footer": Inches(0.492),
    },
    "font": "Calibri",
    "east_asia_font": "Microsoft JhengHei",
    "body_size": 11,
    "body_after": 6,
    "body_line": 1.25,
    "h1": {"size": 16, "color": RGBColor(0x11, 0x11, 0x11), "before": 18, "after": 10},
    "h2": {"size": 13, "color": RGBColor(0x11, 0x11, 0x11), "before": 14, "after": 7},
    "h3": {"size": 12, "color": RGBColor(0x11, 0x11, 0x11), "before": 10, "after": 5},
    "muted": RGBColor(0x55, 0x55, 0x55),
    "ink": RGBColor(0x11, 0x11, 0x11),
    "table_header_fill": "F1F1F1",
    "callout_fill": "F7F7F7",
    "table_width_dxa": 9360,
    "table_indent_dxa": 120,
    "cell_margins_dxa": {"top": 80, "bottom": 80, "start": 120, "end": 120},
}


@dataclass
class EssayRecord:
    file_id: str
    file_name: str
    student_id: str
    essay_score: float
    content_score: float
    structure_score: float
    grammar_score: float
    vocabulary_score: float
    error_count: int
    error_stats: dict[str, int]
    marked_item_count: int
    unanswered: bool
    teacher_feedback: str


@dataclass
class CorrectionItem:
    file_id: str
    item_no: int
    category: str
    wrong: str
    correction: str
    explanation: str
    source_sentence: str
    score_band: str


@dataclass
class KnowledgeCluster:
    cluster_key: str
    category: str
    teaching_family: str
    title: str
    member_entries: list[dict[str, str]]
    frequency: int
    source_count: int
    patterns: list[str]
    representative_source_sentence: str


MANUAL_UNITS = [
    {
        "unit": 1,
        "title": "遣詞與搭配",
        "description": "處理 Word Choice、搭配詞與中文直譯。本次資料中 Word Choice 是最高頻錯誤類型。",
        "entries": [
            {
                "title": "feel 後面接形容詞，不接情緒名詞",
                "wrong": "feel depression",
                "correct": "feel depressed",
                "explanation": "feel 當連綴動詞時，後面通常接形容詞描述主詞狀態。depression 是名詞，若要說「感到沮喪」，應改成 depressed。",
                "example": "I felt depressed after hearing those harsh comments.",
                "tip": "feel + adj.：feel happy / nervous / embarrassed / relieved。",
            },
            {
                "title": "「隱私被侵犯」要讓 privacy 當被動主詞",
                "wrong": "I felt like a violation of privacy.",
                "correct": "I felt that my privacy was being violated.",
                "explanation": "人不會「像一個侵犯」。英文應說「我的隱私被侵犯」，讓 privacy 成為被動語態的主詞。",
                "example": "Her questions made me feel that my privacy was being violated.",
                "tip": "表達感受時先問：主詞到底是人、感受，還是被影響的事物？",
            },
            {
                "title": "concerned about 是擔心，不是被追問",
                "wrong": "I found myself concerned by my aunt.",
                "correct": "I was questioned by my aunt.",
                "explanation": "concerned about 表示「擔心」。若語境是親戚一直問學業或生活，questioned / interrogated / asked repeated questions 更準確。",
                "example": "At the family gathering, I was questioned about my grades.",
                "tip": "中文「被關切」常不是 concern，考場上要依情境換字。",
            },
            {
                "title": "poke 要放進慣用語，不可直接接 business",
                "wrong": "poke others' business",
                "correct": "poke your nose into others' business / pry into others' business",
                "explanation": "poke 單獨使用不能表示「管閒事」。常用片語是 poke one's nose into something 或 pry into something。",
                "example": "It is impolite to pry into others' private lives.",
                "tip": "寫作文時，片語比單字更重要；不要只把中文逐字翻成英文。",
            },
            {
                "title": "well-meaning 表「出於好意」，well-genuine 不自然",
                "wrong": "well-genuine concern",
                "correct": "genuine concern / well-meaning advice",
                "explanation": "genuine 已經表示真誠，不需要加 well。若要說「出於好意的建議」，可用 well-meaning advice。",
                "example": "Although it was well-meaning advice, it still made me uncomfortable.",
                "tip": "well- + 形容詞不是萬用公式；先確認是否為英文既有搭配。",
            },
            {
                "title": "收到評語用 receive，不一定用 accept",
                "wrong": "I accepted negative comments from my teacher.",
                "correct": "I received negative comments from my teacher.",
                "explanation": "accept 偏向主動接受、同意或願意採納；receive 只是客觀地「收到」。批評、評論、建議多用 receive。",
                "example": "I received some unexpected comments about my drawing.",
                "tip": "receive comments / feedback / advice；accept an offer / invitation / apology。",
            },
            {
                "title": "全班不是 the whole classmate",
                "wrong": "the whole classmate",
                "correct": "the whole class / all my classmates",
                "explanation": "classmate 是「一位同學」。要說「全班」可用 the whole class；要說「所有同學」則用 all my classmates。",
                "example": "I felt embarrassed in front of the whole class.",
                "tip": "class 是班級；classmates 是多位同學。",
            },
            {
                "title": "annoyed 形容人，annoying 形容事物",
                "wrong": "I felt annoying.",
                "correct": "I felt annoyed.",
                "explanation": "-ed 形容人的感受，-ing 形容造成感受的人事物。學生作文常把兩者混用。",
                "example": "The repeated questions were annoying, so I felt annoyed.",
                "tip": "bored/boring、embarrassed/embarrassing、shocked/shocking 也同理。",
            },
            {
                "title": "private and respect 要改成同詞性的形容詞",
                "wrong": "a private and respect conversation",
                "correct": "a private and respectful conversation",
                "explanation": "and 連接的兩個修飾語應詞性平行。private 是形容詞，respect 是名詞或動詞，應改為 respectful。",
                "example": "Good communication should be private and respectful.",
                "tip": "A and B 並列時，檢查 A/B 是否同詞性。",
            },
            {
                "title": "「內心」不用 interior heart",
                "wrong": "in my interior heart",
                "correct": "deep down / in my heart",
                "explanation": "interior 常用於建築、空間或物體內部；描述人的內心，用 deep down 或 in my heart 較自然。",
                "example": "Deep down, I still felt hurt by his words.",
                "tip": "遇到抽象中文詞，不要硬套字典第一個意思。",
            },
        ],
    },
    {
        "unit": 2,
        "title": "文法與句構",
        "description": "處理 run-on sentence、comma splice、fragment、缺主詞與子句接法。",
        "entries": [
            {
                "title": "兩個完整句不能只用空白或逗號接起來",
                "wrong": "We receive advice from others In spite of their positive intention, we feel depressed.",
                "correct": "We receive advice from others. In spite of their positive intentions, we may feel depressed.",
                "explanation": "兩個獨立子句之間需要句點、分號或對等連接詞。只用逗號或直接接下一句，就是 run-on sentence 或 comma splice。",
                "example": "The advice was useful, but it still made me uncomfortable.",
                "tip": "每寫完一個長句，先圈出主詞和動詞；若有兩組，就檢查連接方式。",
            },
            {
                "title": "when 子句不能單獨成句",
                "wrong": "When I received the comment. I felt shocked.",
                "correct": "When I received the comment, I felt shocked.",
                "explanation": "When 引導的是從屬子句，必須和主要子句連在一起，不能獨立成完整句。",
                "example": "When my aunt asked about my grades, I did not know how to respond.",
                "tip": "because / although / when / if 開頭的句子，後面通常還需要主句。",
            },
            {
                "title": "Before I could... 可修補破碎句",
                "wrong": "I even still not finish my answering.",
                "correct": "Before I could even finish my answer, she asked another question.",
                "explanation": "原句缺少助動詞與清楚結構。若要表達「還沒來得及回答」，Before I could... 是自然且完整的句型。",
                "example": "Before I could explain myself, everyone started laughing.",
                "tip": "中文「我都還沒...」常可用 Before I could even...。",
            },
            {
                "title": "Personally 是副詞，不能當主詞",
                "wrong": "Personally, offered their opinions on my lifestyle.",
                "correct": "Some people offered their opinions on my lifestyle.",
                "explanation": "句子必須有主詞。Personally 只能修飾整句，不能代替「誰做了這件事」。",
                "example": "Personally, I think advice should be given only when it is needed.",
                "tip": "用 Personally 開頭時，後面常接 I think / I believe / I feel。",
            },
            {
                "title": "When 不是 must distinguish 的主詞",
                "wrong": "When must distinguish between advice and intrusion.",
                "correct": "We must distinguish between advice and intrusion.",
                "explanation": "When 是連接詞或疑問副詞，不能執行 distinguish 這個動作。這類句子通常是缺主詞。",
                "example": "We must distinguish between helpful advice and unnecessary criticism.",
                "tip": "每個英文句子的主要動詞前面都要找到真正的主詞。",
            },
            {
                "title": "be based on 不要直譯成 is stand on",
                "wrong": "Good communication is stand on respect.",
                "correct": "Good communication is based on respect.",
                "explanation": "中文「建立在...之上」不是 stand on。英文常用 be based on / be built on。",
                "example": "A healthy relationship is based on mutual respect.",
                "tip": "抽象概念的「建立在」優先用 based on。",
            },
            {
                "title": "感官動詞後可接 V-ing 表示正在發生",
                "wrong": "I saw a friend criticized my work.",
                "correct": "I saw a friend criticizing my work.",
                "explanation": "see/hear/watch 等感官動詞後接 V-ing，可表示看見某動作正在進行；接原形動詞則偏向看完整個動作。",
                "example": "I noticed my classmates staring at my painting.",
                "tip": "看到「正在...」的畫面時，用 see + O + V-ing。",
            },
            {
                "title": "I didn't behave anything 不符合 behave 用法",
                "wrong": "I didn't behave anything.",
                "correct": "I didn't show any reaction. / I didn't say anything.",
                "explanation": "behave 是不及物動詞，通常說 behave well/badly，不直接接 anything。若要說「沒有反應」，用 show any reaction。",
                "example": "I did not show any reaction, but I felt hurt.",
                "tip": "不要把中文「表現」一律翻成 behave。",
            },
            {
                "title": "that + known 缺 be 動詞",
                "wrong": "an art teacher that known for strict standards",
                "correct": "an art teacher who was known for strict standards / an art teacher known for strict standards",
                "explanation": "關係子句中 known 前需要 be 動詞；若省略 who was，known 可作過去分詞修飾名詞。",
                "example": "She was taught by a teacher known for honest feedback.",
                "tip": "that/who 後面通常要有完整動詞；省略時要改成分詞片語。",
            },
        ],
    },
    {
        "unit": 3,
        "title": "一致性與時態",
        "description": "處理 Agreement、Tense、代名詞一致、單複數與敘事時態。",
        "entries": [
            {
                "title": "過去事件要統一用過去式",
                "wrong": "It is a family gathering and my aunt asks me many questions.",
                "correct": "It was a family gathering and my aunt asked me many questions.",
                "explanation": "描述已發生的個人經驗時，主要敘事動詞應統一用過去式。",
                "example": "Last year, my teacher gave me advice that changed my attitude.",
                "tip": "看到 once / last year / that day / when I was...，先檢查過去式。",
            },
            {
                "title": "leaves 要配合過去敘事改成 left",
                "wrong": "Her words leaves a bitter taste in my mouth.",
                "correct": "Her words left a bitter taste in my mouth.",
                "explanation": "若前文正在描述過去某次事件，結果或感受也應用過去式。",
                "example": "The remark left an invisible burden on me.",
                "tip": "作文常見錯誤不是不會過去式，而是段落中途忘了維持時態。",
            },
            {
                "title": "複數主詞 colors 要配 are",
                "wrong": "These colors you used is terrible.",
                "correct": "These colors you used are terrible.",
                "explanation": "主詞是複數 colors，be 動詞應用 are。插入的 you used 不會改變主詞單複數。",
                "example": "The suggestions she gave were useful.",
                "tip": "先找真正主詞，不要被中間修飾語干擾。",
            },
            {
                "title": "some 後面接可數名詞複數",
                "wrong": "some question",
                "correct": "some questions",
                "explanation": "question 是可數名詞；some 後若接可數名詞，通常要用複數形。",
                "example": "My relatives asked me some questions about my future.",
                "tip": "some advice 例外，因 advice 不可數；some questions 才是可數名詞。",
            },
            {
                "title": "代名詞前後要一致",
                "wrong": "my friend poked their nose into my picture, and he said...",
                "correct": "my friend poked his nose into my picture, and he said...",
                "explanation": "同一個人若前後用 their 和 he，指代會不一致。可以全句使用 singular they，也可以依語境使用 he/she。",
                "example": "My classmate gave his opinion before I asked for it.",
                "tip": "修稿時專門掃一遍 he/she/they/it，檢查指的是誰。",
            },
            {
                "title": "人不能用 it 指稱",
                "wrong": "If the listener wants to speak, it should say something.",
                "correct": "If listeners want to speak, they should say something.",
                "explanation": "it 通常不用來指人。若不限定性別，可改成複數 listeners + they。",
                "example": "When people feel uncomfortable, they may stop sharing their thoughts.",
                "tip": "泛指人時，把單數改成複數常可避開代名詞問題。",
            },
            {
                "title": "expectation 通常依語意用複數",
                "wrong": "Parents often have high expectation.",
                "correct": "Parents often have high expectations.",
                "explanation": "談到父母對課業、未來、表現等多方面期待時，expectations 常用複數。",
                "example": "High expectations can become pressure when they are expressed carelessly.",
                "tip": "抽象名詞也可能可數，要看是否是一個或多個具體期待。",
            },
            {
                "title": "relationship 泛指人際關係常用複數",
                "wrong": "Advice can hurt relationship.",
                "correct": "Advice can hurt relationships.",
                "explanation": "若不是某一段特定關係，而是泛指人際關係，relationship 常用複數。",
                "example": "Careless comments may damage relationships.",
                "tip": "泛指一類事物時，可數名詞常用複數。",
            },
        ],
    },
    {
        "unit": 4,
        "title": "冠詞、介系詞與可數性",
        "description": "處理 Article、Preposition、不可數名詞與固定介系詞搭配。",
        "entries": [
            {
                "title": "advice 是不可數名詞",
                "wrong": "unsolicited advices",
                "correct": "unsolicited advice",
                "explanation": "advice 在英文中不可數，不能加 s。若要表達一則建議，用 a piece of advice。",
                "example": "Unsolicited advice can make people feel pressured.",
                "tip": "advice 不可數；suggestion 可數。",
            },
            {
                "title": "communication 表概念時不可數",
                "wrong": "a good communication",
                "correct": "good communication",
                "explanation": "communication 表「溝通」這個概念時通常不可數，不加 a。",
                "example": "Good communication requires respect and patience.",
                "tip": "a communication 多指一則正式訊息或通訊，不是一般作文中的「溝通」。",
            },
            {
                "title": "food 作餐桌食物時通常不可數",
                "wrong": "delicious-looking foods",
                "correct": "delicious-looking food",
                "explanation": "food 泛指食物時不可數；若強調多種食品種類，才可能用 foods。",
                "example": "We were enjoying the delicious-looking food at the feast.",
                "tip": "學測作文中多數情境用 food 即可。",
            },
            {
                "title": "單數可數名詞 gathering 前要有冠詞",
                "wrong": "during family gathering",
                "correct": "during a family gathering",
                "explanation": "gathering 是可數名詞；單數使用時前面需要 a/an/the 或所有格。",
                "example": "During a family gathering, my relatives asked about my future.",
                "tip": "看到單數可數名詞裸奔，是冠詞錯誤高風險。",
            },
            {
                "title": "in the future 是固定說法",
                "wrong": "in future",
                "correct": "in the future",
                "explanation": "美式與一般高中英文寫作中，表示「未來」通常寫 in the future。",
                "example": "In the future, I hope adults will respect teenagers' choices.",
                "tip": "future 前常有 the；不要受中文「在未來」影響而省略。",
            },
            {
                "title": "go 後接目的地要加 to",
                "wrong": "go university",
                "correct": "go to university",
                "explanation": "go 是不及物動詞；接目的地時通常要加 to。",
                "example": "My aunt kept asking which university I wanted to go to.",
                "tip": "go home 是少數例外，home 可當副詞。",
            },
            {
                "title": "enjoy 是及物動詞，不加 in",
                "wrong": "enjoying in the feast",
                "correct": "enjoying the feast",
                "explanation": "enjoy 後面直接接受詞，不需要介系詞 in。",
                "example": "Everyone was enjoying the feast when the conversation became awkward.",
                "tip": "enjoy + N/V-ing：enjoy dinner, enjoy talking with friends。",
            },
            {
                "title": "出於真誠用 out of sincerity",
                "wrong": "from his sincerity",
                "correct": "out of sincerity",
                "explanation": "表示動機或出發點時，常用 out of + 抽象名詞，例如 out of kindness / curiosity / concern。",
                "example": "He gave the advice out of sincerity, but it still hurt me.",
                "tip": "中文「來自真誠」不要直譯成 from sincerity。",
            },
            {
                "title": "talk 後接主題要保留 about",
                "wrong": "talk our future",
                "correct": "talk about our future",
                "explanation": "talk 是不及物動詞；若要接談論主題，需用 talk about。",
                "example": "Parents should talk about teenagers' future with respect.",
                "tip": "discuss 可直接接受詞；talk about 才能直接接主題。",
            },
        ],
    },
    {
        "unit": 5,
        "title": "拼字、大小寫與標點",
        "description": "處理 Spelling、Capitalization、Punctuation 與考場修稿檢查。",
        "entries": [
            {
                "title": "receive 的拼字",
                "wrong": "recieve",
                "correct": "receive",
                "explanation": "receive 是本批資料中多次出現的拼字錯誤。請記得 c 後面是 ei。",
                "example": "We often receive advice from people around us.",
                "tip": "receive / receipt / deceive 都要特別檢查 ei/ie。",
            },
            {
                "title": "occasion 的拼字",
                "wrong": "ocassion",
                "correct": "occasion",
                "explanation": "occasion 有兩個 c、一個 s。",
                "example": "On one occasion, my aunt asked me about my grades.",
                "tip": "寫個人經驗常用 On one occasion，拼字要背熟。",
            },
            {
                "title": "calm down 不要拼成 clam down",
                "wrong": "clam down",
                "correct": "calm down",
                "explanation": "calm 表示冷靜；clam 是另一個完全不同的字。",
                "example": "I tried to calm down before answering the question.",
                "tip": "考場上可直接用 stay calm，較短也較安全。",
            },
            {
                "title": "unnecessary 的拼字",
                "wrong": "unnessary",
                "correct": "unnecessary",
                "explanation": "unnecessary 中間是 c + essary，常見錯誤是漏 c 或 s 數量錯。",
                "example": "Unnecessary comments can create pressure.",
                "tip": "necessary / unnecessary 是作文高頻字，建議整組背。",
            },
            {
                "title": "business 中間只有一個 s",
                "wrong": "bussiness",
                "correct": "business",
                "explanation": "business 是高頻拼字陷阱，中間不是 ssiness。",
                "example": "It is rude to pry into others' business.",
                "tip": "mind your own business 也是常用片語。",
            },
            {
                "title": "appearance 的拼字",
                "wrong": "apperance",
                "correct": "appearance",
                "explanation": "appearance 由 appear + ance 組成，中間保留 appear 的 ea。",
                "example": "Comments about appearance can be hurtful.",
                "tip": "appear / appearance 一起記。",
            },
            {
                "title": "happening 只有一個 n",
                "wrong": "happenning",
                "correct": "happening",
                "explanation": "happen 加 -ing 時不雙寫 n。",
                "example": "The conversation was happening during Chinese New Year.",
                "tip": "happened 有 -ed；happening 不加第二個 n。",
            },
            {
                "title": "Chinese New Year 要大寫",
                "wrong": "chinese new year",
                "correct": "Chinese New Year",
                "explanation": "節慶與專有名詞的主要單字需大寫。",
                "example": "During Chinese New Year, relatives often gather for dinner.",
                "tip": "國籍、語言、節慶、地名首字母大寫。",
            },
            {
                "title": "a lot 永遠分開寫",
                "wrong": "alot",
                "correct": "a lot",
                "explanation": "a lot 是兩個字，不可連寫。",
                "example": "I learned a lot from that uncomfortable experience.",
                "tip": "若想更正式，可改用 many / much / numerous / considerable。",
            },
        ],
    },
    {
        "unit": 6,
        "title": "作文架構與表達策略",
        "description": "從老師總評萃取出的共通寫作問題：分段、內容深度、轉承與收尾。",
        "entries": [
            {
                "title": "不要全篇只寫一段",
                "wrong": "把背景、三個例子、感受與結論全部塞在同一段。",
                "correct": "引言一段、主體一至兩段、結論一段。",
                "explanation": "多份總評指出，單段作文會讓讀者難以追蹤主題、例子與反思。學測作文即使篇幅短，也應有清楚段落。",
                "example": "Paragraph 1: introduce the issue. Paragraph 2: explain one experience. Paragraph 3: reflect and conclude.",
                "tip": "交卷前先看版面：若整篇只有一個大段落，結構分通常會受影響。",
                "source_sentence": "In modern social interaction, we usually recieve some \"unsolicited advices\" about our life from other people In spite of their positive intention, we would sometimes feel depression. On one ocassion, I found myself concerned by my aunt during chinese new year gathering. She offered her opinions on my study...",
            },
            {
                "title": "引言要交代主題，不要立刻跳例子",
                "wrong": "第一句直接開始講親戚問問題。",
                "correct": "先用一至兩句介紹 unsolicited advice 或 communication 的共同情境，再進入個人經驗。",
                "explanation": "引言的功能是讓讀者知道文章要討論什麼，而不是只提供第一個事件。",
                "example": "In daily life, advice is often given with good intentions, but it can still make people uncomfortable.",
                "tip": "引言 = 主題背景 + 文章立場。",
                "source_sentence": "It was happenning in a Chinese New Year, my family members were enjoying in the feast and one of my relatives started asking me some question about which university I want to go in future.",
            },
            {
                "title": "每個主體段要有主題句",
                "wrong": "連續列出事件，沒有說明此段重點。",
                "correct": "用主題句先指出這個例子要證明什麼。",
                "explanation": "主題句能讓例子服務論點，而不是變成流水帳。",
                "example": "One reason unsolicited advice can be harmful is that it may cross personal boundaries.",
                "tip": "主體段第一句可用 One reason... / This experience taught me that...。",
                "source_sentence": "She offered her opinions on my study, which leaves a bitter taste in my mouth. Another event is my classmate thought my hairstyle is ugly, which makes me feel like a violation of privacy.",
            },
            {
                "title": "例子後要分析原因、影響與反思",
                "wrong": "只寫事件發生了什麼和我很難過。",
                "correct": "補上為何難過、造成什麼影響、後來學到什麼。",
                "explanation": "老師總評多次指出內容停留在表面描述。高分作文需要把事件推進到解釋與反思。",
                "example": "The comment hurt me not because it was direct, but because it was made in front of others.",
                "tip": "例子後至少加一句 because / as a result / therefore。",
                "source_sentence": "The remark placed an invisible burden on me. To prevent the conflict, I was forced to clam down to avoid unnessary conflict.",
            },
            {
                "title": "例子之間要有轉承",
                "wrong": "One event... Another event... The other event...",
                "correct": "First,... More importantly,... This experience also shows that...",
                "explanation": "轉承詞能標示段落邏輯，讓讀者知道下一個例子和主題的關係。",
                "example": "More importantly, even well-meaning advice can become pressure when it is given publicly.",
                "tip": "不要只換例子；要說明例子如何推進論點。",
            },
            {
                "title": "結論要收束觀點，不只重複一句話",
                "wrong": "In short, it is impolite to pry into others' business.",
                "correct": "In short, advice should be offered with respect for timing, privacy, and the listener's feelings.",
                "explanation": "結論應回扣文章核心，並提出較完整的判準或啟示。",
                "example": "In short, good communication begins when people learn to listen before they judge.",
                "tip": "結論可回答：那麼我們應該怎麼做？",
            },
            {
                "title": "不要只追求難字，先確保搭配正確",
                "wrong": "用 perceive depression、research 代替所有 study。",
                "correct": "先選語意自然、搭配正確的字，再考慮升級詞彙。",
                "explanation": "詞彙升級若忽略搭配，反而會讓句子不自然。作文評分看的是精準度，不是單字難度。",
                "example": "I felt depressed is better than I perceived depression.",
                "tip": "高分詞彙 = 合語境 + 合搭配 + 不破壞清楚度。",
            },
            {
                "title": "修稿優先順序：句構、時態、單複數、拼字",
                "wrong": "只檢查有沒有寫滿字數。",
                "correct": "先檢查每句是否完整，再檢查過去式、主謂一致、可數名詞和高頻拼字。",
                "explanation": "本批資料的高頻錯誤集中在 Word Choice、Grammar、Agreement、Spelling、Tense。考場修稿應優先處理這些會明顯扣分的問題。",
                "example": "After finishing an essay, underline verbs and check whether the story is consistently in the past tense.",
                "tip": "最後三分鐘不要重寫內容；用清單掃掉高頻錯誤。",
                "source_sentence": "In modern social interaction, we usually recieve some \"unsolicited advices\" about our life from other people In spite of their positive intention, we would sometimes feel depression.",
            },
        ],
    },
]


def pdf_file_id(path: Path) -> str:
    return hashlib.sha256(path.stem.encode("utf-8")).hexdigest()[:12]


def infer_student_id(path: Path) -> str:
    match = re.search(r"批改結果[-_](\d{4,})", path.stem)
    if match:
        return match.group(1)
    match = re.search(r"(?<!\d)(\d{5,})(?!\d)", path.stem)
    return match.group(1) if match else ""


def configure_paths(pdf_dir: Path | str, output_dir: Path | str, batch_name: str) -> None:
    global PDF_DIR, OUTPUT_DIR, MANUAL_PATH, SUMMARY_CSV, ITEMS_CSV, ITEMS_JSON, BATCH_NAME
    PDF_DIR = Path(pdf_dir)
    OUTPUT_DIR = Path(output_dir)
    MANUAL_PATH = OUTPUT_DIR / "manual.docx"
    SUMMARY_CSV = OUTPUT_DIR / "essay_summary.csv"
    ITEMS_CSV = OUTPUT_DIR / "essay_correction_items.csv"
    ITEMS_JSON = OUTPUT_DIR / "essay_correction_items.json"
    BATCH_NAME = batch_name.strip() or "未命名批次"


def score_band(score: float) -> str:
    if score == 0:
        return "0"
    if score <= 5:
        return "1-5"
    if score <= 10:
        return "5.5-10"
    if score <= 15:
        return "10.5-15"
    return "15.5-20"


def read_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def compact_cjk(text: str) -> str:
    text = text.replace("\u3000", " ")
    text = re.sub(r"(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])", "", text)
    text = re.sub(r"\s+([，。：；！？）])", r"\1", text)
    text = re.sub(r"([（])\s+", r"\1", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalize_lookup(text: str) -> str:
    text = text.lower()
    text = re.sub(r"[\u2018\u2019]", "'", text)
    text = re.sub(r"[\u201c\u201d]", '"', text)
    text = re.sub(r"[^a-z0-9'\"]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


MARKED_ITEMS_RE = re.compile(
    r"(?:標\s*[注註]\s*修\s*改\s*與\s*建\s*議|批\s*改\s*建\s*議|修\s*改\s*建\s*議)\s*[:：]",
    flags=re.S,
)

ESSAY_TEXT_RE = re.compile(
    r"(?:作文\s*內容\s*與\s*標\s*[注註]|作文\s*內容|學生\s*作文|原\s*文)\s*[:：]",
    flags=re.S,
)
STRICT_ESSAY_TEXT_RE = re.compile(r"作文\s*內容\s*與\s*標\s*[注註]\s*[:：]", flags=re.S)

TEACHER_FEEDBACK_RE = re.compile(
    r"(?:老師\s*總\s*評|教師\s*總\s*評|批\s*改\s*總\s*評|整體\s*回饋|總\s*評)\s*[:：]",
    flags=re.S,
)


def find_essay_section(text: str) -> str:
    match = re.search(r"作文\s*（\s*占\s*20\s*分\s*）", text)
    if match:
        return text[match.start() :]

    candidates: list[int] = []
    for marker in [
        r"作文\s*評\s*分\s*[:：]",
        r"作文\s*得\s*分\s*[:：]",
        r"寫作\s*得\s*分\s*[:：]",
    ]:
        score_match = re.search(marker, text)
        if score_match:
            candidates.append(score_match.start())
    for marker in [ESSAY_TEXT_RE, TEACHER_FEEDBACK_RE, MARKED_ITEMS_RE]:
        marker_match = marker.search(text)
        if marker_match:
            candidates.append(marker_match.start())

    if not candidates:
        return ""
    section = text[min(candidates) :]
    signal_count = sum(
        1
        for pattern in [
            r"作文\s*(?:評\s*分|得\s*分)\s*[:：]",
            r"寫作\s*得\s*分\s*[:：]",
            ESSAY_TEXT_RE,
            TEACHER_FEEDBACK_RE,
            MARKED_ITEMS_RE,
        ]
        if (pattern.search(section) if hasattr(pattern, "search") else re.search(pattern, section))
    )
    return section if signal_count >= 2 else ""


def extract_student_essay_text(section: str) -> str:
    start = STRICT_ESSAY_TEXT_RE.search(section) or ESSAY_TEXT_RE.search(section)
    end = MARKED_ITEMS_RE.search(section)
    if not start or not end or end.start() <= start.end():
        return ""
    essay = section[start.end() : end.start()]
    essay = re.sub(r"共\s*\d+\s*字\s*[，,]\s*\d+\s*處\s*修\s*正", "", essay)
    return compact_cjk(essay)


def split_source_sentences(essay_text: str) -> list[str]:
    if not essay_text:
        return []
    protected = re.sub(r"\s+", " ", essay_text).strip()
    chunks = re.split(r"(?<=[.!?])\s+(?=[A-Z\"'])", protected)
    sentences: list[str] = []
    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue
        if len(chunk) > 380:
            # Long run-on originals are still useful, but keep the handbook readable.
            chunk = chunk[:377].rstrip() + "..."
        sentences.append(chunk)
    return sentences


def find_source_sentence(wrong: str, essay_text: str) -> str:
    sentences = split_source_sentences(essay_text)
    if not sentences:
        return ""
    wrong_norm = normalize_lookup(wrong)
    if wrong_norm:
        for sentence in sentences:
            if wrong_norm in normalize_lookup(sentence):
                return sentence

    wrong_tokens = [tok for tok in wrong_norm.split() if len(tok) >= 3]
    if wrong_tokens:
        best_sentence = ""
        best_score = 0
        for sentence in sentences:
            sentence_norm = normalize_lookup(sentence)
            score = sum(1 for tok in wrong_tokens if tok in sentence_norm)
            if score > best_score:
                best_sentence = sentence
                best_score = score
        if best_score >= max(1, min(3, len(wrong_tokens))):
            return best_sentence
    return ""


def parse_error_stats(section: str) -> tuple[int, dict[str, int]]:
    match = re.search(
        r"錯\s*誤\s*統\s*計\s*[:：]\s*(\d+)\s*項\s*(.*?)(?=老師\s*總\s*評|老師總評)",
        section,
        flags=re.S,
    )
    if not match:
        return 0, {}
    count = int(match.group(1))
    stats_text = match.group(2)
    stats = {
        name.strip(): int(num)
        for name, num in re.findall(r"([A-Za-z][A-Za-z ]+?)\s*[:：]\s*(\d+)\s*\(", stats_text)
    }
    return count, stats


def extract_teacher_feedback(section: str) -> str:
    start = TEACHER_FEEDBACK_RE.search(section)
    if not start:
        return ""
    end = ESSAY_TEXT_RE.search(section[start.end() :])
    if end:
        feedback = section[start.end() : start.end() + end.start()]
    else:
        alt_end = MARKED_ITEMS_RE.search(section[start.end() :])
        feedback = section[start.end() : start.end() + alt_end.start()] if alt_end else section[start.end() :]
    return compact_cjk(feedback)


EXPLANATION_TAIL_MARKERS = {
    "Advice",
    "Suggestion",
    "Suggestions",
    "Recommendation",
    "Recommend",
    "Compromise",
    "Keep",
    "Seem",
    "Harm",
    "That",
    "Talk",
    "Cared",
    "Obtain",
    "Perceive",
    "Indifferent",
    "Revolve around",
    "Good begin",
    "Sale",
    "Chat",
    "Cut",
    "Choose",
    "Feel",
    "Improve",
}


def strip_explanation_tail(correction: str) -> str:
    value = correction.strip()
    value = re.sub(r"(?<=[a-z])(?=[A-Z])", " ", value)
    tokens = value.split()
    if len(tokens) >= 2 and tokens[-1].strip(".,;:!?").lower() == tokens[-2].strip(".,;:!?").lower():
        value = " ".join(tokens[:-1])
    for marker in sorted(EXPLANATION_TAIL_MARKERS, key=len, reverse=True):
        if re.search(rf"\s+{re.escape(marker)}$", value):
            value = re.sub(rf"\s+{re.escape(marker)}$", "", value).strip()
            break
    return value


def strip_glued_explanation_tail(correction: str, wrong: str = "") -> str:
    value = correction.strip()
    if not value or value == DELETE_CORRECTION:
        return value

    compact_letters = re.sub(r"[^A-Za-z]", "", value.lower())
    if compact_letters and len(compact_letters) % 2 == 0:
        midpoint = len(compact_letters) // 2
        if compact_letters[:midpoint] == compact_letters[midpoint:]:
            return compact_letters[:midpoint]

    wrong_tokens = [token for token in normalize_lookup(wrong).split() if len(token) >= 3]
    for token in sorted(set(wrong_tokens), key=len, reverse=True):
        if not value.lower().endswith(token):
            continue
        start = len(value) - len(token)
        if start <= 0 or not value[start - 1].isalpha():
            continue
        candidate = value[:start].strip()
        if candidate and normalize_lookup(candidate) != normalize_lookup(wrong):
            value = candidate
            break

    for marker in sorted(GLUED_EXPLANATION_TAIL_MARKERS, key=len, reverse=True):
        if not value.lower().endswith(marker):
            continue
        start = len(value) - len(marker)
        if start <= 0 or not value[start - 1].isalpha():
            continue
        candidate = value[:start].strip()
        if len(normalize_lookup(candidate).replace(" ", "")) >= 3:
            value = candidate
            break

    lower = value.lower()
    if lower.endswith("to") and " " not in value and lower not in {"into", "onto"}:
        candidate = value[:-2].strip()
        if len(candidate) >= 4:
            value = candidate

    return value.strip()


def split_correction_and_explanation(rest: str, wrong: str = "") -> tuple[str, str]:
    raw = rest.replace("\r\n", "\n").replace("\r", "\n").strip()
    lines = [compact_cjk(line) for line in raw.split("\n") if compact_cjk(line)]
    if len(lines) >= 2 and not re.search(r"[\u4e00-\u9fff]", lines[0]) and re.search(r"[\u4e00-\u9fff]", "\n".join(lines[1:])):
        correction_lines: list[str] = []
        explanation_lines: list[str] = []
        for index, line in enumerate(lines):
            if re.search(r"[\u4e00-\u9fff]", line):
                explanation_lines = lines[index:]
                break
            correction_lines.append(line)
        correction = " ".join(correction_lines).strip()
        explanation = compact_cjk("\n".join(explanation_lines))
    else:
        rest = compact_cjk(raw)
        first_cjk = re.search(r"[\u4e00-\u9fff]", rest)
        if first_cjk:
            correction = rest[: first_cjk.start()].strip()
            explanation = rest[first_cjk.start() :].strip()
        else:
            parts = rest.split("\n", 1)
            correction = parts[0].strip()
            explanation = parts[1].strip() if len(parts) > 1 else ""

    correction = correction.strip(" '\"「」")
    correction = re.sub(r"(?<=[a-z])(?=[A-Z])", " ", correction)
    correction = re.sub(r"\s+", " ", correction)
    quoted_explanation = re.match(r"(.+?)(\s*[\"“”「][A-Z][A-Za-z].*)$", correction)
    if quoted_explanation:
        correction = quoted_explanation.group(1).strip(" '\"「」")
        explanation = (quoted_explanation.group(2).strip() + " " + explanation).strip()
    for clean_prefix in ["studying", "chatting", "customers", "writing", "exams"]:
        if correction.lower().startswith(clean_prefix) and len(correction) > len(clean_prefix):
            explanation = (correction[len(clean_prefix) :].strip() + " " + explanation).strip()
            correction = clean_prefix
            break
    wrong_head = normalize_lookup(wrong).split(" ")[:2]
    if wrong_head and '"' in correction:
        correction = correction.split('"', 1)[0].strip()
    wrong_key = normalize_lookup(wrong)
    correction_key = normalize_lookup(correction)
    wrong_tokens = wrong_key.split()
    if wrong_key and len(wrong_tokens) >= 2 and wrong_key in correction_key:
        wrong_match = re.search(rf"(?<![A-Za-z]){re.escape(wrong)}(?![A-Za-z])", correction, flags=re.I)
        if wrong_match and wrong_match.start() > 0:
            prefix = correction[: wrong_match.start()].strip().lower()
            if prefix in {"a", "an", "the"}:
                wrong_match = None
        if wrong_match and wrong_match.start() > 0:
            explanation = (correction[wrong_match.start() :].strip() + " " + explanation).strip()
            correction = correction[: wrong_match.start()].strip()
    correction = strip_explanation_tail(correction.strip(" '\"「」"))
    correction = strip_glued_explanation_tail(correction, wrong)
    correction = strip_explanation_tail(correction.strip(" '\"「」"))
    if not correction and (explanation or raw):
        correction = DELETE_CORRECTION
    if len(correction) > 120:
        correction = correction[:117].rstrip() + "..."
    return correction, explanation


def recover_tail_arrow_correction(wrong_chunk: str) -> tuple[str, str, str] | None:
    text = compact_cjk(wrong_chunk)
    match = re.match(r"^([A-Za-z][A-Za-z'\-]*)\s+(.+?)([\u4e00-\u9fff].*)$", text)
    if not match:
        return None
    wrong = match.group(1).strip()
    correction = strip_glued_explanation_tail(match.group(2).strip(), wrong)
    explanation = match.group(3).strip()
    if not wrong or not correction or is_dirty_correction(correction):
        return None
    return wrong, correction, explanation


def infer_category(wrong: str, correction: str, explanation: str) -> str:
    haystack = f"{wrong} {correction} {explanation}"
    lower = haystack.lower()
    if any(key in haystack for key in ["分段", "段落", "主題句", "引言", "結論", "例子", "轉承", "內容深度", "架構", "修稿"]):
        return "Writing Strategy"
    if correction == DELETE_CORRECTION and is_punctuation_pair(wrong, correction):
        return "Punctuation"
    if wrong and correction and wrong.lower() == correction.lower() and wrong != correction:
        return "Capitalization"
    if re.search(r"[A-Za-z]'$", wrong.strip()) and "cafe" in normalize_lookup(correction):
        return "Spelling"
    wrong_norm = normalize_lookup(wrong)
    correction_norm = normalize_lookup(correction.split("/")[0])
    if is_article_pair(wrong_norm, correction_norm) or any(key in haystack for key in ["冠詞", "不可數", "可數名詞", "不定冠詞"]):
        return "Article"
    if is_preposition_pair(wrong_norm, correction_norm) or "介系詞" in haystack:
        return "Preposition"
    if is_punctuation_pair(wrong, correction) or any(key in lower for key in ["comma splice", "run-on", "fragment"]) or any(key in haystack for key in ["標點", "逗號", "句點"]):
        return "Punctuation"
    if any(key in haystack for key in ["時態", "過去式", "現在式"]):
        return "Tense"
    if any(key in haystack for key in ["主謂", "主詞", "複數", "單數", "一致", "代名詞"]):
        return "Agreement"
    if any(key in haystack for key in ["詞性", "形容詞", "副詞", "動名詞"]):
        return "Grammar"
    if wrong_norm.replace(" ", "") == correction_norm.replace(" ", "") and wrong_norm != correction_norm:
        return "Spelling"
    spelling_words = {
        "recieve",
        "ocassion",
        "clam",
        "unnessary",
        "bussiness",
        "apperance",
        "happenning",
        "alot",
    }
    if any(word in lower for word in spelling_words) or "拼字" in haystack:
        return "Spelling"
    if is_likely_spelling_pair(wrong_norm, correction_norm):
        return "Spelling"
    if any(key in haystack for key in ["不自然", "慣用", "用詞", "中式", "搭配", "語意"]):
        return "Word Choice"
    return "Grammar"


PREPOSITIONS = {"in", "on", "at", "to", "for", "with", "about", "of", "from", "into", "by", "as"}
ARTICLES = {"a", "an", "the"}


def is_preposition_pair(wrong: str, correction: str) -> bool:
    wrong_tokens = wrong.split()
    correction_tokens = correction.split()
    if not wrong_tokens or not correction_tokens:
        return False
    if wrong_tokens[0] == "to" and any(tok in {"mr", "ms", "teacher"} for tok in wrong_tokens):
        return False
    wrong_preps = [tok for tok in wrong_tokens if tok in PREPOSITIONS]
    correction_preps = [tok for tok in correction_tokens if tok in PREPOSITIONS]
    if wrong_tokens[0] in PREPOSITIONS and correction_tokens[0] in PREPOSITIONS and wrong_tokens[0] != correction_tokens[0]:
        return True
    if wrong_preps == correction_preps or not (wrong_preps or correction_preps):
        return False
    wrong_content = {tok for tok in wrong_tokens if tok not in PREPOSITIONS}
    correction_content = {tok for tok in correction_tokens if tok not in PREPOSITIONS}
    overlap = len(wrong_content & correction_content)
    needed = max(1, min(len(wrong_content), len(correction_content)) - 1)
    return overlap >= needed


def is_article_pair(wrong: str, correction: str) -> bool:
    wrong_tokens = wrong.split()
    correction_tokens = correction.split()
    if not wrong_tokens or not correction_tokens:
        return False
    wrong_article = wrong_tokens[0] in ARTICLES
    correction_article = correction_tokens[0] in ARTICLES
    if wrong_article != correction_article:
        return True
    if wrong_article and correction_article and wrong_tokens[0] != correction_tokens[0]:
        return True
    return False


def is_punctuation_pair(wrong: str, correction: str) -> bool:
    wrong_clean = wrong.strip()
    correction_clean = correction.strip()
    if not wrong_clean or not correction_clean:
        return False
    punctuation = set(".,;:!?")
    if all(ch in punctuation for ch in wrong_clean):
        return True
    if wrong_clean.lower() in {"and", "but", "so"} and any(mark in correction_clean for mark in [",", "."]):
        return True
    if re.sub(r"\s+([.,;:!?])", r"\1", wrong_clean) == correction_clean:
        return True
    if normalize_lookup(wrong_clean) == normalize_lookup(correction_clean) and any(
        mark in wrong_clean + correction_clean for mark in punctuation
    ):
        return True
    return False


def is_likely_spelling_pair(wrong: str, correction: str) -> bool:
    wrong_words = wrong.split()
    correction_words = correction.split()
    if len(wrong_words) != 1 or len(correction_words) != 1:
        return False
    wrong_word = wrong_words[0]
    correction_word = correction_words[0]
    if len(wrong_word) < 4 or len(correction_word) < 4:
        return False
    if wrong_word == correction_word:
        return False
    if inflection_stems(wrong_word) & inflection_stems(correction_word):
        return False
    if abs(len(wrong_word) - len(correction_word)) > 2 and (
        wrong_word.startswith(correction_word) or correction_word.startswith(wrong_word)
    ):
        return False
    simple_inflections = {
        wrong_word + "s",
        wrong_word + "es",
        wrong_word + "ed",
        wrong_word + "d",
        wrong_word + "ing",
    }
    if correction_word in simple_inflections:
        return False
    ratio = difflib.SequenceMatcher(None, wrong_word, correction_word).ratio()
    return ratio >= 0.74


def inflection_stem(word: str) -> str:
    stems = inflection_stems(word)
    return sorted(stems, key=len)[0]


def inflection_stems(word: str) -> set[str]:
    stems = {word}
    if word.endswith("ied") and len(word) > 5:
        stems.add(word[:-3] + "y")
    if word.endswith("ing") and len(word) > 6:
        base = word[:-3]
        stems.add(base)
        stems.add(base + "e" if not base.endswith("e") else base)
    if word.endswith("ed") and len(word) > 5:
        stems.add(word[:-2])
        stems.add(word[:-1])
    if word.endswith("es") and len(word) > 5:
        stems.add(word[:-2])
    if word.endswith("s") and len(word) > 4:
        stems.add(word[:-1])
    if word.endswith("d") and len(word) > 4:
        stems.add(word[:-1])
    return stems


UNIT_GROUPS = [
    {
        "unit": 1,
        "title": "遣詞與搭配",
        "categories": ["Word Choice", "Style", "Redundancy"],
        "description": "整理用詞、搭配與中文直譯問題。",
    },
    {
        "unit": 2,
        "title": "文法與句構",
        "categories": ["Grammar"],
        "description": "整理句構、子句、片語與完整句問題。",
    },
    {
        "unit": 3,
        "title": "一致性與時態",
        "categories": ["Agreement", "Tense"],
        "description": "整理主謂一致、單複數、代名詞與時態問題。",
    },
    {
        "unit": 4,
        "title": "冠詞、介系詞與可數性",
        "categories": ["Article", "Preposition"],
        "description": "整理冠詞、介系詞、可數與不可數名詞問題。",
    },
    {
        "unit": 5,
        "title": "拼字、大小寫與標點",
        "categories": ["Spelling", "Capitalization", "Punctuation"],
        "description": "整理拼字、大小寫與標點問題。",
    },
]

CATEGORY_LABELS = {
    "Word Choice": "遣詞搭配",
    "Style": "表達自然度",
    "Redundancy": "冗詞",
    "Grammar": "文法句構",
    "Agreement": "一致性",
    "Tense": "時態",
    "Article": "冠詞可數性",
    "Preposition": "介系詞",
    "Spelling": "拼字",
    "Capitalization": "大小寫",
    "Punctuation": "標點",
    "Writing Strategy": "作文策略",
}

CATEGORY_TIPS = {
    "Word Choice": "先確認英文搭配是否自然，再追求高級詞彙。",
    "Style": "修稿時讀一次英文句子，刪掉中文直譯或不自然表達。",
    "Redundancy": "刪掉重複資訊，讓句子保留一個清楚重點。",
    "Grammar": "每句先檢查主詞、主要動詞與連接方式。",
    "Agreement": "圈出主詞和動詞，確認人稱、單複數與代名詞一致。",
    "Tense": "敘述過去經驗時，動詞時態要一路維持一致。",
    "Article": "單數可數名詞前通常需要 a/an/the 或所有格。",
    "Preposition": "介系詞多為固定搭配，建議直接背完整片語。",
    "Spelling": "拼字錯誤適合整理成個人清單，考前反覆檢查。",
    "Capitalization": "專有名詞、節日與句首字母要大寫。",
    "Punctuation": "逗號不能任意連接兩個完整句。",
    "Writing Strategy": "每個例子後至少補一句原因、影響或反思。",
}

DELETE_CORRECTION = "[刪除]"

GLUED_EXPLANATION_TAIL_MARKERS = {
    "host",
    "important",
    "lifestyle",
    "lots",
    "marvel",
    "mini",
    "some",
    "usually",
}


def preview_text(text: str, limit: int = 46) -> str:
    text = compact_cjk(text)
    if len(text) <= limit:
        return text
    return text[: limit - 3].rstrip() + "..."


def normalize_group_key(text: str) -> str:
    normalized = normalize_lookup(text)
    if normalized:
        return normalized
    return compact_cjk(text).lower()


NOISY_WRONG_KEYS = {
    "",
    ".",
    ",",
    ";",
    ":",
    "!",
    "?",
    "and",
    "but",
    "so",
    "a",
    "an",
    "the",
    "in",
    "on",
    "at",
    "to",
    "of",
    "for",
    "with",
    "about",
    "am",
    "is",
    "are",
    "was",
    "were",
    "be",
    "been",
    "being",
    "do",
    "does",
    "did",
    "have",
    "has",
    "had",
    "i",
    "you",
    "he",
    "she",
    "it",
    "we",
    "they",
    "this",
    "that",
    "these",
    "those",
    "my",
    "your",
    "his",
    "her",
    "our",
    "their",
    "me",
    "him",
    "us",
    "them",
    "i'm",
    "im",
}
SHORT_CORRECTION_ONLY_KEYS = {
    "i",
    "you",
    "he",
    "she",
    "it",
    "we",
    "they",
    "my",
    "your",
    "his",
    "her",
    "our",
    "their",
    "me",
    "him",
    "us",
    "them",
}

LIST_ENTRY_CATEGORIES = {"Spelling", "Capitalization", "Punctuation", "Preposition"}

MERGED_LIST_LIMITS = {
    "Spelling": 14,
    "Capitalization": 12,
    "Punctuation": 12,
    "Preposition": 10,
}

MERGED_LIST_TITLES = {
    "Spelling": "本批常見拼字清單",
    "Capitalization": "本批大小寫清單",
    "Punctuation": "本批標點修正清單",
    "Preposition": "本批常見介系詞搭配",
    "Article": "本批冠詞與可數性清單",
}

CATEGORY_MAIN_LIMITS = {
    "Word Choice": 10,
    "Style": 3,
    "Redundancy": 3,
    "Grammar": 12,
    "Agreement": 10,
    "Tense": 7,
    "Article": 8,
    "Preposition": 7,
}
DEDUPE_CATEGORY_PRIORITY = {
    "Grammar": 10,
    "Agreement": 9,
    "Tense": 8,
    "Word Choice": 7,
    "Article": 6,
    "Preposition": 5,
    "Style": 4,
    "Redundancy": 3,
}
MAIN_CATEGORY_ORDER = [
    "Word Choice",
    "Style",
    "Redundancy",
    "Grammar",
    "Agreement",
    "Tense",
    "Article",
    "Preposition",
]
MAX_MAIN_ENTRIES = 42
TEACHER_TARGET_ENTRIES_PER_UNIT = 11
TEACHER_MAX_ENTRIES_PER_UNIT = 12
TEACHER_TARGET_MANUAL_ENTRIES = 44
TEACHER_MAX_MANUAL_ENTRIES = 50
CORE_SOURCE_COUNT_THRESHOLD = 20
SUPPLEMENTAL_SOURCE_COUNT_THRESHOLD = 5
CLUSTER_PATTERN_LIMIT = 8
CLUSTER_MEMBER_PATTERN_MIN_SOURCE = 2

TEACHING_FAMILY_META = {
    "countable_plural": {
        "title": "可數名詞要依語意使用複數形",
        "explanation": "這組錯誤都和可數名詞的單複數判斷有關。當句子泛指多個人、物或座位時，名詞通常要改成複數，而不是保留單數形式。",
        "example": "Many customers use their phones while waiting for seats.",
        "tip": "看到 many、some、two、people、customers 或泛指一群人時，先檢查後面的可數名詞是否要加 s/es。",
    },
    "plural_after_number_quantifier": {
        "title": "數量詞後的可數名詞要用複數",
        "explanation": "two、many、some、several、no enough 等數量概念後，若接的是可數名詞，通常要使用複數形。",
        "example": "Two students are using several seats.",
        "tip": "看到數字或 many/some/several，立刻檢查後面的可數名詞是否為複數。",
    },
    "subject_verb_agreement": {
        "title": "主詞與動詞要依單複數一致",
        "explanation": "這組錯誤都和主詞、動詞一致有關。動詞形式要跟真正主詞的單複數一致，單數主詞常要加 s/es，複數主詞則不要多加。",
        "example": "This behavior causes problems for the cafe owner.",
        "tip": "先圈出真正主詞，再看主要動詞是否跟它的單複數一致。",
    },
    "third_person_singular_verb": {
        "title": "第三人稱單數現在式動詞要加 s/es",
        "explanation": "主詞若是 he、she、it、this action、the picture、a student 這類單數概念，現在式主要動詞通常要加 s/es。",
        "example": "This action causes trouble for other customers.",
        "tip": "主詞可代換成 it/he/she 時，現在式動詞多半要加 s/es。",
    },
    "be_verb_agreement": {
        "title": "be 動詞要和主詞單複數一致",
        "explanation": "is、are、was、were 的選擇要看真正主詞。複數主詞用 are/were；單數主詞用 is/was。",
        "example": "The seats are not enough for new customers.",
        "tip": "看到 is/are，往前找真正主詞，不要被介系詞片語或修飾語干擾。",
    },
    "pronoun_agreement": {
        "title": "代名詞要和前面的名詞一致",
        "explanation": "代名詞必須和它指代的名詞在人稱與單複數上保持一致。若前面是複數 seats、customers，就不能用 it 或 this 代替。",
        "example": "The seats are occupied, so other customers cannot use them.",
        "tip": "看到 it/they/this/these，回頭確認它指的是單數還是複數。",
    },
    "indefinite_pronoun_agreement": {
        "title": "everyone / another / other 的單複數要分清楚",
        "explanation": "everyone 視為單數；another 後接單數名詞；other 可修飾複數名詞；others 則可直接當代名詞使用。",
        "example": "Everyone needs a seat, and others may be waiting outside.",
        "tip": "another + 單數名詞；other + 複數名詞；others = other people。",
    },
    "subject_needed": {
        "title": "英文完整句通常不能省略主詞",
        "explanation": "正式作文中，主要子句通常需要明確主詞。中文可以省略「我希望」，英文常要補出 I hope 等主詞動詞。",
        "example": "I hope the owner can set a clearer rule.",
        "tip": "每句先檢查有沒有主詞與主要動詞。",
    },
    "there_be_agreement": {
        "title": "there be 句型要和後面的名詞一致",
        "explanation": "there is / there are 的 be 動詞要看後面真正主詞的單複數。後面若是複數名詞，應使用 there are。",
        "example": "There are not enough seats for new customers.",
        "tip": "看到 there is / there are，往後找第一個名詞來決定 be 動詞。",
    },
    "some_others": {
        "title": "some... others... 表達「有些人...其他人...」",
        "explanation": "這組錯誤都和代名詞 others 的用法有關。表達群體中另一部分人時，常用 some... others...，不要只寫 other。",
        "example": "Some customers are studying, while others are chatting with friends.",
        "tip": "other 後面通常接名詞；others 可以直接指「其他人」。",
    },
    "article_missing": {
        "title": "單數可數名詞前要有冠詞或限定詞",
        "explanation": "這組錯誤都和單數可數名詞前缺少 a/an/the 或限定詞有關。英文不能讓單數可數名詞孤立出現，必須依語境補冠詞。",
        "example": "A coffee shop should set a clear seating rule.",
        "tip": "看到單數可數名詞，先問：是泛指一個、特指那個，還是應改成複數？",
    },
    "article_place_single": {
        "title": "地點類單數可數名詞前要補冠詞",
        "explanation": "coffee shop、cafe、cafeteria、place 等地點名詞若以單數出現，前面通常需要 a/an/the 或改成複數。",
        "example": "The picture shows a coffee shop with many customers.",
        "tip": "描述一個地點時，先檢查是否要寫 a cafe / a coffee shop / the cafe。",
    },
    "article_object_single": {
        "title": "常見物品單數名詞前要補 a/an/the",
        "explanation": "seat、book、laptop、drink、nap、look 等可數名詞以單數出現時，前面通常需要冠詞或限定詞。",
        "example": "A customer is reading a book and using a laptop.",
        "tip": "看到單數物品名詞，問自己：是不是少了 a/an/the？",
    },
    "article_time_expression": {
        "title": "時間片語常有固定冠詞",
        "explanation": "for a long time、the whole day 等時間片語有固定冠詞搭配，不能直接逐字省略。",
        "example": "Some customers stay for a long time without ordering more food.",
        "tip": "時間片語不要逐字翻譯，直接背 for a long time、the whole day。",
    },
    "article_a_an_choice": {
        "title": "a / an 要看後面單字的發音",
        "explanation": "不定冠詞不是只看拼字，而是看後面單字的開頭發音。母音音素前用 an，子音音素前用 a。",
        "example": "An old customer is waiting for a seat.",
        "tip": "先念出後面的字，再決定 a 或 an。",
    },
    "article_specific_context": {
        "title": "特定圖片或角色可用 the",
        "explanation": "若句子指的是圖片中的特定人、物或地點，通常用 the；若是第一次泛指其中一個，則用 a/an。",
        "example": "The picture shows the owner of a cafe.",
        "tip": "the 表示讀者知道是哪一個；a/an 表示其中一個。",
    },
    "article_uncountable_noun": {
        "title": "不可數名詞不要誤加複數或冠詞",
        "explanation": "work、advice、information 等不可數名詞不能直接加 s，也不一定能用 a/an。需要數量時可改用 a piece of 或其他量詞。",
        "example": "Some students bring work to the cafe.",
        "tip": "看到抽象或工作類名詞，先確認它是不是不可數。",
    },
    "article_choice": {
        "title": "冠詞 a/the 要看是否特定",
        "explanation": "這組錯誤都和 a/an/the 的特定性判斷有關。第一次提到或泛指任一個時多用 a/an；讀者已知或特定對象才用 the。",
        "example": "A customer may need a seat, but the seat near the window is already taken.",
        "tip": "a/an 是「其中一個」；the 是「你我都知道的那一個」。",
    },
    "generic_plural_countable": {
        "title": "泛指一類事物時可用複數名詞",
        "explanation": "這組錯誤都和泛指類別的名詞形式有關。若不是指單一特定物，而是在談一整類事物，常用複數名詞表達泛指。",
        "example": "Coffee shops need rules during busy hours.",
        "tip": "談一整類人事物時，優先檢查是否要用複數名詞。",
    },
    "participle_ving": {
        "title": "名詞後的動作描述可用 V-ing",
        "explanation": "這組錯誤都和分詞結構有關。當句子已經有主要動詞，後面描述名詞正在做的動作時，常改成 V-ing 來修飾前面的名詞。",
        "example": "There is a student occupying a table for several hours.",
        "tip": "一句話只能有一個主要動詞；第二個動作若在修飾名詞，先考慮 V-ing。",
    },
    "there_is_n_ving": {
        "title": "There is/are + 名詞 + V-ing 描述畫面動作",
        "explanation": "描述圖片中某人正在做某事時，可用 there is/are + 名詞 + V-ing。不要把第二個動詞直接寫成原形或現在式。",
        "example": "There are two students sleeping at the table.",
        "tip": "There is/are 後面若還要補動作，先檢查是否要改成 V-ing。",
    },
    "second_verb_ving": {
        "title": "一句已有主要動詞時，第二個動作用 V-ing 修飾",
        "explanation": "英文一個子句通常只能有一個主要動詞。若後面的動作是補充描述前面的名詞，常要改成 V-ing。",
        "example": "A man is standing near the door, looking for a seat.",
        "tip": "看到連續兩個動詞，先判斷第二個是不是要改成分詞。",
    },
    "after_preposition_ving": {
        "title": "介系詞後接動詞時要用 V-ing",
        "explanation": "of、by、after、before、instead of、without 等介系詞後如果接動詞，動詞要改成 V-ing。",
        "example": "Instead of studying for hours, customers should order a drink.",
        "tip": "看到介系詞後面接動作，先改成 V-ing 檢查。",
    },
    "parallel_ving": {
        "title": "並列動作要維持相同形式",
        "explanation": "列舉多個動作時，動詞形式要平行一致。若前面用 V-ing，後面的動作也應維持 V-ing。",
        "example": "Some customers are chatting, studying, and scrolling on their phones.",
        "tip": "and / or 連接的項目，左右形式要對稱。",
    },
    "sit_seat_word_form": {
        "title": "sit 是動詞，seat 是名詞",
        "explanation": "sit 表示「坐」這個動作；seat 表示「座位」。寫作時要先判斷句中需要動詞還是名詞。",
        "example": "A customer wants to sit in a seat near the window.",
        "tip": "要說動作寫 sit；要說座位寫 seat。",
    },
    "adjective_participle_form": {
        "title": "描述狀態或感受時要用正確形容詞形式",
        "explanation": "full、crowded、depressed、relaxing 等詞在句中常作形容詞。不要把它們誤寫成不存在或不合語境的動詞形式。",
        "example": "The cafe is crowded, and some customers feel relaxed.",
        "tip": "be 動詞後面常接形容詞；先檢查該字是不是正確形容詞。",
    },
    "there_have_structure": {
        "title": "表示「有」某物時用 there is/are",
        "explanation": "中文說「那裡有」，英文存在句要用 there is/are，不用 there have。",
        "example": "There are many customers in the cafe.",
        "tip": "there have 幾乎不是標準存在句；先改成 there is/are。",
    },
    "because_so_although_but": {
        "title": "because / so 與 although / but 不要成對重複",
        "explanation": "英文中 because 已經表示原因，so 表示結果；although 已經表示讓步，but 表示轉折。兩組通常不能同時放在同一句主要結構中。",
        "example": "Because the cafe is crowded, new customers cannot find seats.",
        "tip": "中文的「因為...所以...」「雖然...但是...」翻成英文時通常二選一。",
    },
    "relative_clause_place": {
        "title": "地點先行詞後可用 where 引導子句",
        "explanation": "place、cafe、shop 等地點名詞後，若子句要表達「在那裡做某事」，常用 where 或 in which。",
        "example": "A cafe is a place where people can study or relax.",
        "tip": "先行詞是地點，後面缺的是地點副詞時，優先檢查 where。",
    },
    "quantifier_no_any": {
        "title": "no 與 not any 不要混用",
        "explanation": "no 本身已經有否定與數量含義，通常不再和 any 疊在一起。可寫 no seats 或 not any seats。",
        "example": "There are no seats for new customers.",
        "tip": "no any 看到就改成 no 或 not any。",
    },
    "subjunctive_if": {
        "title": "假設語氣要使用 were / would / could",
        "explanation": "這組錯誤都和與現在事實相反的假設語氣有關。If 子句常用 were，主要子句常搭配 would 或 could。",
        "example": "If I were the owner, I would set a clear time limit.",
        "tip": "看到 If I were...，後面主句通常要接 would / could，而不是 will / can。",
    },
    "modal_base_verb": {
        "title": "情態助動詞後接原形動詞",
        "explanation": "can、could、should、would、must 等情態助動詞後面要接原形動詞，不再加過去式、第三人稱單數或分詞變化。",
        "example": "We can see many customers in the cafe.",
        "tip": "看到 can / should / must，下一個主要動詞先還原成原形。",
    },
    "picture_present_tense": {
        "title": "描述圖片時通常使用現在式",
        "explanation": "這組錯誤都和圖片描述的時態一致有關。描述圖片中正在呈現的狀態時，通常使用現在式。",
        "example": "The picture shows several customers using a cafe as a study space.",
        "tip": "圖片描述不是回憶過去，動詞先用現在式檢查。",
    },
    "picture_present_progressive": {
        "title": "描述圖片中的正在進行動作用現在進行式",
        "explanation": "若要描述圖片中人物當下正在做的動作，常使用 be + V-ing，而不是單純現在式或過去式。",
        "example": "Two customers are sleeping at the table.",
        "tip": "圖片中正在發生的動作，可先試 be + V-ing。",
    },
    "present_perfect_current_result": {
        "title": "到目前為止的狀態可用現在完成式",
        "explanation": "若句子強調某動作到現在仍有影響，例如尚未點餐、剛走進來，現在完成式會比過去式更精確。",
        "example": "Some customers have not ordered anything yet.",
        "tip": "看到 yet、just、到目前為止的結果，先檢查現在完成式。",
    },
    "subjunctive_tense": {
        "title": "假設語氣中的助動詞時態要一致",
        "explanation": "這組錯誤都和假設語氣中的時態搭配有關。If 子句與主要子句要維持同一套假設語氣，不要混用 can / will。",
        "example": "If I were the owner, I could ask customers to follow the rule politely.",
        "tip": "假設語氣中看到 were，後面多半要搭配 would / could。",
    },
    "spelling_list": {
        "title": "本批常見拼字清單",
        "explanation": "這些是本批反覆出現的拼字問題，適合整理成考前清單反覆核對。",
        "example": "Customers are chatting in a crowded coffee shop.",
        "tip": "把常錯字做成個人清單，寫完作文後用 30 秒逐字掃描。",
    },
    "spelling_customer_list": {
        "title": "customer / customers 相關拼字",
        "explanation": "這組錯誤都和 customer / customers 的拼字或易混字有關。costumer 是服裝師，customs 是海關或習俗；若要說咖啡店顧客，應寫 customer / customers。",
        "example": "Many customers were waiting for a seat in the coffee shop.",
        "tip": "顧客是 customer；服裝師是 costumer；海關或習俗是 customs。",
    },
    "spelling_ing_list": {
        "title": "V-ing 拼字變化",
        "explanation": "這組錯誤都和 V-ing 拼字變化有關。短母音 CVC 字尾常需要重複尾音再加 -ing；study 加 -ing 時保留 y。",
        "example": "Some customers are chatting, while others are studying.",
        "tip": "加 -ing 前先檢查：是否要雙寫字尾？y 是否要保留？",
    },
    "spelling_cafe_coffee_list": {
        "title": "cafe / coffee 相關拼字",
        "explanation": "這組錯誤都和本批主題中的 cafe / coffee 拼字有關。cafe 不需要撇號；coffee 要有兩個 f 與兩個 e。",
        "example": "The coffee shop was crowded in the afternoon.",
        "tip": "寫 cafe / coffee 這類主題核心字時，務必在完稿後逐字檢查。",
    },
    "spelling_confusable_words_list": {
        "title": "易混淆拼字與同音字",
        "explanation": "這組錯誤多半不是單純少一個字母，而是把另一個英文單字誤用進句子，例如 their/there、sign/sigh、scene/scent。",
        "example": "The sign says that customers should not occupy seats for too long.",
        "tip": "如果拼字檢查沒有紅線，仍要確認該字的意思是否符合語境。",
    },
    "spelling_topic_words_list": {
        "title": "本批主題核心字拼字",
        "explanation": "這組錯誤集中在本批作文主題常用字。主題核心字一旦拼錯，會明顯影響老師對內容掌握度的判斷。",
        "example": "A modern cafe should provide a comfortable space for customers.",
        "tip": "寫完後先掃主題關鍵字，例如 customer、coffee、modern、minimum、business。",
    },
    "spelling_double_letter_list": {
        "title": "常見雙寫與漏字母拼字",
        "explanation": "這組錯誤多半是少寫、錯寫或重複字母位置不對，例如 business、address、approach、inconvenience。",
        "example": "The owner should address the inconvenience with a clear approach.",
        "tip": "遇到長字先切音節，再確認雙寫字母與常漏字母。",
    },
    "spelling_people_words_list": {
        "title": "人物與日常名詞常錯拼字",
        "explanation": "這組收錄 people、other、girls 等日常高頻字的拼字錯誤。這些字不難，但出現頻率高，考場最容易因粗心失分。",
        "example": "Other people may be waiting for seats.",
        "tip": "越常用的字越要檢查，因為拼錯會讓句子看起來很不穩。",
    },
    "spelling_common_words_list": {
        "title": "其他高頻常錯拼字",
        "explanation": "這組收錄本批反覆出現、但不屬於單一拼字規則的常錯字，適合整理成考前個人清單。",
        "example": "A modern coffee shop should provide separate spaces for different needs.",
        "tip": "無法用規則推導的拼字，最適合用清單記憶。",
    },
    "preposition_list": {
        "title": "本批常見介系詞搭配",
        "explanation": "這組錯誤都和介系詞固定搭配或空間位置表達有關。介系詞通常不能逐字翻譯，應直接背完整片語。",
        "example": "The customer is sitting at the table and working on a laptop.",
        "tip": "介系詞要連同前後詞一起記：at the table、on the left、in front of、work on。",
    },
    "preposition_quantity_list": {
        "title": "數量與飲品表達不要直譯 cup of drink",
        "explanation": "這組錯誤都和中文「一杯飲料」的直譯有關。英文通常直接說 a drink，或在需要強調容器時說 a cup of coffee / tea。",
        "example": "Customers should order a drink before using a seat for a long time.",
        "tip": "drink 本身就是飲料；不要把「一杯飲料」硬翻成 a cup of drink。",
    },
    "preposition_table_position_list": {
        "title": "人在桌邊用 at the table",
        "explanation": "描述人在桌邊坐著、工作或休息時，通常用 at the table；on the table 會變成在桌面上。",
        "example": "Two students are sleeping at the table.",
        "tip": "人坐在桌邊用 at；東西放在桌面上才常用 on。",
    },
    "preposition_picture_context_list": {
        "title": "描述圖片內容用 in the picture",
        "explanation": "描述圖片裡出現的人事物時，常用 in the picture。不要受中文「在畫面上」影響誤用 on。",
        "example": "In the picture, several customers are using laptops.",
        "tip": "圖片裡有什麼，英文通常是 in the picture。",
    },
    "preposition_perspective_list": {
        "title": "觀點片語使用 from my perspective",
        "explanation": "表達「從我的觀點」時，常用 from my perspective 或 in my opinion，不要直接寫 in my perspective。",
        "example": "From my perspective, the owner should set a seating limit.",
        "tip": "perspective 常搭配 from；opinion 常搭配 in。",
    },
    "preposition_because_of_list": {
        "title": "because 接子句，because of 接名詞",
        "explanation": "because 後面要接完整子句；because of 後面接名詞或名詞片語。兩者不能只照中文「因為」互換。",
        "example": "Customers cannot find seats because the cafe is crowded.",
        "tip": "because + S + V；because of + N。",
    },
    "preposition_passive_by_list": {
        "title": "被動語態中的 by / with 要看語意",
        "explanation": "描述座位被某人或物品佔用時，常用 occupied by；with 多表示伴隨或工具，不能直接取代 by。",
        "example": "The seats are occupied by customers' bags.",
        "tip": "被誰佔用用 by；用什麼工具或伴隨狀態才考慮 with。",
    },
    "preposition_location_list": {
        "title": "空間位置介系詞要用固定說法",
        "explanation": "這組錯誤都和位置表達有關。英文描述桌邊、左邊、前方、角落時，介系詞常是固定搭配，不能照中文逐字選。",
        "example": "The student is sitting at the table near the corner.",
        "tip": "位置片語直接背整組：at the table、on the left、in front of、in the corner。",
    },
    "preposition_verb_phrase_list": {
        "title": "動詞後的介系詞搭配要一起記",
        "explanation": "這組錯誤都和動詞後面的介系詞搭配有關。sit、work 等動詞接不同介系詞時，意思或自然度會改變。",
        "example": "Some customers are working on their laptops and sitting in their seats.",
        "tip": "背單字時不要只背動詞，連後面的介系詞一起背。",
    },
    "capitalization_list": {
        "title": "本批大小寫清單",
        "explanation": "這些是本批反覆出現的大小寫問題，通常和句首、專有名詞或固定名稱有關。",
        "example": "Taiwan has many famous landmarks.",
        "tip": "句首、I、專有名詞與地名要特別檢查大小寫。",
    },
    "punctuation_list": {
        "title": "本批標點修正清單",
        "explanation": "這些是本批反覆出現的標點問題，常見原因是逗號連接兩個完整句或標點位置不自然。",
        "example": "Some customers are studying, while others are chatting.",
        "tip": "逗號不能單獨連接兩個完整句；必要時改用句點、分號或連接詞。",
    },
    "punctuation_comma_splice_list": {
        "title": "逗號不能連接兩個完整句",
        "explanation": "兩個可以獨立成立的句子不能只用逗號連接。應改用句點、分號，或加入 and、but、so 等連接詞。",
        "example": "The cafe is crowded. New customers cannot find seats.",
        "tip": "逗號前後若各自都有主詞動詞，先檢查是否是 comma splice。",
    },
    "punctuation_transition_list": {
        "title": "轉折連接詞前後標點要完整",
        "explanation": "however、but、therefore、furthermore 等轉折或連接詞要配合正確標點。however 通常不能直接用逗號夾在兩個完整句中間。",
        "example": "The cafe is crowded; however, some customers still stay for hours.",
        "tip": "however 前常用句點或分號；but 連接兩個完整句時前面常加逗號。",
    },
    "punctuation_listing_list": {
        "title": "列舉多個動作時要用連接詞或分號",
        "explanation": "列舉多個人物動作或子句時，不能一直用逗號串接。最後一項前常需要 and，較長的並列子句可用分號。",
        "example": "Some customers are studying, others are chatting, and still others are sleeping.",
        "tip": "列舉到最後一項時，檢查是否需要 and。",
    },
    "punctuation_spacing_list": {
        "title": "英文標點前不空格，標點後空一格",
        "explanation": "英文逗號、句點、分號應緊接前一個單字，後面再空一格。中文輸入習慣常會造成標點前多空格。",
        "example": "In the picture, many customers are sitting in a cafe.",
        "tip": "完稿後掃描逗號與句點前面是否多打一格。",
    },
    "punctuation_fragment_list": {
        "title": "Because / such as 片語不能單獨成句",
        "explanation": "because 引導從屬子句，such as 引導舉例片語，兩者通常不能單獨當完整句。要補主要子句或併入前後句。",
        "example": "This is because many customers stay for a long time.",
        "tip": "句首看到 Because 或 Such as，先檢查整句是否有主要子句。",
    },
    "word_choice_chinese_translation": {
        "title": "避免中文直譯，改用英文自然結構",
        "explanation": "這組錯誤多半是把中文語序或中文邏輯直接搬進英文，導致句子雖有單字卻不自然。修正時要改用英文慣用結構。",
        "example": "Their behavior leaves other customers with no place to sit.",
        "tip": "翻譯後讀一次英文，如果像中文句子，就改用英文固定句型。",
    },
    "word_choice_causative_result": {
        "title": "表達「導致」時選對 cause / lead to / leave",
        "explanation": "中文的「讓、造成、導致」不能一律翻成 make 或 let。英文常依語意使用 cause + O + to-V、lead to + N/V-ing、leave + O + 狀態。",
        "example": "This behavior leaves new customers without seats.",
        "tip": "壞結果常用 cause/lead to/leave，不要每次都用 make 或 let。",
    },
    "word_choice_collocation": {
        "title": "常用名詞與動詞搭配要整組記",
        "explanation": "英文自然度常來自固定搭配，例如 set a time limit、create an atmosphere、place an order。不要只背單字本身。",
        "example": "The owner can set a time limit during busy hours.",
        "tip": "背動詞時一起背後面的名詞搭配。",
    },
    "word_choice_redundancy": {
        "title": "刪除語意重複或贅字",
        "explanation": "英文句子若同時放入兩個功能相同的片語，會顯得累贅。例如 from my perspective 和 I think 通常擇一即可。",
        "example": "From my perspective, cafes should set clearer rules.",
        "tip": "同一句中若兩個片語都在表達「我認為」，保留一個即可。",
    },
    "word_choice_precision": {
        "title": "選字要符合語意角色",
        "explanation": "用字不只看中文意思，也要看英文中的語意角色。人、行為、結果與物品若角色放錯，句子會變得不清楚。",
        "example": "Their actions prevent the owner from serving more customers.",
        "tip": "先確認主詞是誰、動作是什麼、結果落在誰身上，再選字。",
    },
    "word_choice_cafe_terms": {
        "title": "咖啡店情境常用語要自然",
        "explanation": "本批作文常談咖啡店、座位、低消與久坐，相關詞彙可直接用 seating limit、minimum purchase、occupy seats、turnover 等自然說法。",
        "example": "A minimum purchase and a seating limit can improve table turnover.",
        "tip": "主題相關詞彙直接背英文慣用說法，不逐字翻「低消、佔位、翻桌率」。",
    },
}


def effective_essay_count(records: list[EssayRecord] | None = None, items: list[CorrectionItem] | None = None) -> int:
    if records is not None:
        return sum(1 for record in records if not record.unanswered and record.essay_score > 0)
    if items is not None:
        return len({item.file_id for item in items if item.file_id})
    return 0


def teacher_correction_source_threshold(effective_count: int) -> int:
    if effective_count < 50:
        return 2
    if effective_count < 200:
        return 3
    return 5


def teacher_list_source_threshold(effective_count: int) -> int:
    return max(3, math.ceil(max(0, effective_count) * 0.015))


def teacher_strategy_source_threshold(effective_count: int) -> int:
    if effective_count <= 0:
        return 0
    return max(5, math.ceil(effective_count * 0.02))

STRATEGY_RULES = [
    {
        "title": "提供具體證據支撐成績申訴",
        "keywords": ["具體", "證據", "細節", "數據", "佐證", "例子", "考試成績", "評分標準"],
        "wrong": "I think my grade should be higher because I worked hard.",
        "correct": "I would like to explain my concern by referring to my test scores, assignments, and class participation.",
        "explanation": "正式申訴不能只說自己很努力，必須用考試、作業、課堂表現或評分標準作為證據，讓請求更有說服力。",
        "example": "For example, my quiz scores improved from 68 to 86, and I submitted every assignment on time.",
        "tip": "態度是背景，證據才是申訴的核心。",
    },
    {
        "title": "正式信件要維持禮貌、清楚、具體",
        "keywords": ["正式", "禮貌", "語氣", "尊重", "清晰", "目的明確"],
        "wrong": "I don't agree with my grade, and I hope you can change it.",
        "correct": "I respectfully ask whether you could review my final grade and clarify the grading criteria.",
        "explanation": "寫給老師的信件應避免命令語氣，改用 respectfully ask、would like to、could you 等較正式的請求方式。",
        "example": "I would appreciate the opportunity to discuss my performance with you.",
        "tip": "正式信件的語氣要禮貌，但請求仍要明確。",
    },
    {
        "title": "請求要明確：重新評估、澄清或約時間討論",
        "keywords": ["請求", "重新評估", "澄清", "會面", "討論", "re-evaluate", "review"],
        "wrong": "I hope you can help me.",
        "correct": "Could you please review my grade or let me know when I could discuss it with you?",
        "explanation": "結尾若只寫希望得到幫助，讀者不清楚下一步。應明確提出要重新評估、說明標準，或安排討論時間。",
        "example": "If possible, I would like to meet with you this week to understand how my grade was calculated.",
        "tip": "結尾要讓老師知道你希望他做哪一件事。",
    },
    {
        "title": "段落架構要符合正式信件流程",
        "keywords": ["結構", "開頭", "主體", "結尾", "段落", "格式"],
        "wrong": "I studied hard this semester. My grade is low. Please check it.",
        "correct": "First state the purpose, then provide evidence, and finally make a polite request.",
        "explanation": "正式信件建議分成三層：開頭說明目的，主體提出證據與理由，結尾提出具體請求與感謝。",
        "example": "I am writing to ask about my final grade. I believe my assignments and exam scores may not be fully reflected. Could you please review them?",
        "tip": "目的、證據、請求三段式最穩定。",
    },
    {
        "title": "論點不能只停留在學習態度",
        "keywords": ["論點", "說服力", "深度", "具體", "學習態度", "連結性"],
        "wrong": "I always listened carefully and tried my best.",
        "correct": "In addition to my effort, my improvement in exams and assignments shows steady progress.",
        "explanation": "努力學習是好的材料，但若沒有連到成績、表現或評分標準，說服力會不足。要把態度轉化成可判斷的成果。",
        "example": "My participation also improved because I volunteered answers more often in the second half of the semester.",
        "tip": "每個態度句後面補一個可觀察的成果。",
    },
]

CAFE_STRATEGY_RULES = [
    {
        "title": "先交代圖片場景與核心問題",
        "keywords": ["圖片", "場景", "咖啡", "咖啡廳", "咖啡店", "座位", "顧客", "佔用"],
        "source_keywords": ["coffee", "cafe", "shop", "seat", "seats", "customers", "occupy", "table"],
        "wrong": "There are many people in a coffee shop.",
        "correct": "The picture shows a crowded coffee shop where some customers occupy seats without ordering enough food or drinks.",
        "explanation": "描述圖片時，不只要列出人物，也要點出核心衝突。這批作文的核心多半是咖啡店座位被長時間佔用、店家營運與其他顧客權益受到影響。",
        "example": "In the picture, several customers are using the cafe as a study or work space, which leaves few seats for people who want to buy drinks.",
        "tip": "第一段要讓讀者立刻知道：地點、人物、問題。",
    },
    {
        "title": "解決方案要具體到可執行",
        "keywords": ["解決方案", "具體", "規定", "措施", "時間", "低消", "提醒", "執行"],
        "source_keywords": ["rule", "rules", "limit", "time limit", "minimum", "order", "owner", "shopkeeper"],
        "wrong": "The boss should make a rule.",
        "correct": "The owner could set a two-hour seating limit and ask staff to remind customers politely during busy hours.",
        "explanation": "只寫 make a rule 太籠統。高分作文需要說明規則內容、執行方式，以及它如何改善座位不足的問題。",
        "example": "For example, the cafe could require each customer to order at least one drink during peak hours.",
        "tip": "每個 solution 後面補：誰執行、怎麼執行、為什麼有效。",
    },
    {
        "title": "論述要平衡店家與顧客需求",
        "keywords": ["顧客", "店家", "老闆", "公平", "權益", "收入", "消費", "其他人"],
        "source_keywords": ["customers", "owner", "business", "consume", "order", "right", "fair", "other customers"],
        "wrong": "People who study there should leave immediately.",
        "correct": "Although customers may need a place to study, they should also respect the business and other customers.",
        "explanation": "這類題目不是只責備某一方。較成熟的論述會同時考量店家收入、顧客使用空間的需求，以及其他顧客找不到座位的困擾。",
        "example": "A reasonable time limit can protect the cafe's business while still allowing customers to enjoy the space.",
        "tip": "把雙方立場都寫進去，說服力會更強。",
    },
    {
        "title": "避免只羅列人物動作，要補原因與影響",
        "keywords": ["描述", "羅列", "原因", "影響", "深入", "分析", "說服力"],
        "source_keywords": ["sleeping", "chatting", "laptop", "study", "studying", "occupy", "seats", "tables"],
        "wrong": "A woman is reading, and two people are talking.",
        "correct": "Some customers are staying for a long time, so the cafe may lose potential customers who cannot find a seat.",
        "explanation": "圖片描述不能停在人物動作清單。應進一步說明這些行為造成什麼問題，例如翻桌率下降、收入受影響、真正想消費的人無位可坐。",
        "example": "If most tables are occupied by people who are not ordering, the cafe may become less welcoming to new customers.",
        "tip": "看到一個畫面細節，就往後補一句 consequence。",
    },
    {
        "title": "語氣要專業，不要過度強硬",
        "keywords": ["語氣", "專業", "禮貌", "強硬", "報警", "提醒", "溝通"],
        "source_keywords": ["politely", "remind", "warn", "leave", "policy", "rule", "customers", "staff"],
        "wrong": "If they do not leave, I will call the police.",
        "correct": "If customers stay too long without ordering, the staff could politely remind them of the seating policy.",
        "explanation": "作為店主提出方案時，語氣要合理且專業。過度強硬的做法會讓文章顯得不成熟，也可能偏離題目期待的解決問題能力。",
        "example": "A polite reminder is more practical than forcing customers to leave immediately.",
        "tip": "解法要可行，也要符合現實中的服務情境。",
    },
]

LANDMARK_STRATEGY_RULES = [
    {
        "title": "選擇景點後，要說清楚代表台灣的理由",
        "keywords": ["代表", "選擇", "理由", "特色", "獨特", "台灣", "景點", "自然", "文化", "landscape", "scenery"],
        "source_keywords": ["represent", "taiwan", "choose", "unique", "landscape", "scenery", "natural", "culture"],
        "wrong": "I would choose this place because it is beautiful.",
        "correct": "I would choose Taroko Gorge because its marble cliffs, ecology, and cultural value show Taiwan's natural beauty.",
        "explanation": "景點型作文不能只寫 beautiful 或 famous。要說明這個景點如何代表台灣，例如自然地形、文化背景、生態特色或國際辨識度。",
        "example": "This landmark represents Taiwan because it combines striking natural scenery with local culture.",
        "tip": "選定景點後，至少補兩個具體理由支撐代表性。",
    },
    {
        "title": "推廣策略要寫出媒介、對象與做法",
        "keywords": ["社群媒體", "推廣", "策略", "宣傳", "影片", "活動", "國際", "投票", "campaign", "promote"],
        "source_keywords": ["social media", "campaign", "promote", "promotion", "video", "international", "vote", "voting"],
        "wrong": "I will promote it on the Internet.",
        "correct": "I will launch a short-video campaign on social media to introduce the landmark to international travelers.",
        "explanation": "只說 use social media 太籠統。高分寫法要交代使用哪種媒介、鎖定誰，以及用什麼內容吸引對方。",
        "example": "A campaign featuring short videos and local stories could encourage more international voters to support the landmark.",
        "tip": "策略句可用：媒介 + 內容 + 目標讀者 + 預期效果。",
    },
    {
        "title": "自然與文化特色要用具體細節支撐",
        "keywords": ["具體", "細節", "自然", "文化", "生態", "地形", "原住民", "歷史", "描述"],
        "source_keywords": ["marble", "erosion", "mountain", "sea", "cloud", "indigenous", "ecology", "geological"],
        "wrong": "The place has nature and culture.",
        "correct": "The place is known for its marble cliffs, ecological diversity, and indigenous cultural traditions.",
        "explanation": "抽象名詞如 nature、culture、beauty 需要後面接具體內容，讀者才知道特色在哪裡。",
        "example": "Instead of saying the scenery is special, describe the cliffs, coastline, wildlife, or cultural festival.",
        "tip": "每個抽象特色後面補一個看得見的細節。",
    },
    {
        "title": "結論要回扣全球投票與台灣形象",
        "keywords": ["結論", "回扣", "比賽", "投票", "全球", "國際", "形象", "recognition", "competition"],
        "source_keywords": ["global", "world", "international", "competition", "recognition", "vote", "voting"],
        "wrong": "In conclusion, I like this place very much.",
        "correct": "In conclusion, this landmark deserves global recognition because it shows Taiwan's natural beauty and cultural depth.",
        "explanation": "結尾不要只重複個人喜好。應回到題目任務，說明此景點為何值得在國際活動中代表台灣。",
        "example": "With a clear campaign, the landmark could help global audiences see a richer image of Taiwan.",
        "tip": "最後一句回答：它如何代表台灣、為何值得被世界看見。",
    },
]

ADVICE_STRATEGY_RULES = [
    {
        "title": "先界定 advice 與 intrusion 的差別",
        "keywords": ["advice", "unsolicited", "intrusion", "boundary", "boundaries", "隱私", "界線", "建議", "干涉"],
        "wrong": "Advice is good, but sometimes it is bad.",
        "correct": "The key issue is whether advice respects the listener's privacy, timing, and personal boundaries.",
        "explanation": "這類題目不能只說建議有好有壞；高分文章要清楚界定善意建議何時會變成干涉，並提出判準。",
        "example": "Advice becomes intrusive when it is given publicly, repeatedly, or without understanding the listener's feelings.",
        "tip": "先定義界線，再寫例子，文章立場會更清楚。",
    },
    {
        "title": "個人經驗後要補原因與影響",
        "keywords": ["example", "experience", "原因", "影響", "反思", "感受", "privacy", "pressure", "embarrassed"],
        "wrong": "My aunt asked me many questions, and I felt bad.",
        "correct": "My aunt's repeated questions embarrassed me because they exposed my private choices in front of others.",
        "explanation": "只寫事件和心情會停在表面描述。例子後應補上為何受傷、造成什麼壓力，以及這件事如何支撐主題。",
        "example": "The comment hurt me not because it was direct, but because it was made in public without considering my feelings.",
        "tip": "每個例子後至少加一句 because / as a result / this made me realize。",
    },
    {
        "title": "把善意與表達方式分開討論",
        "keywords": ["good intention", "well-meaning", "positive intention", "善意", "語氣", "方式", "respect"],
        "wrong": "They had good intentions, so I should accept their advice.",
        "correct": "Even well-meaning advice can hurt others if it is delivered at the wrong time or in the wrong tone.",
        "explanation": "本題常見盲點是把善意等同於合理。文章應指出動機可能是好的，但表達方式、場合與尊重同樣重要。",
        "example": "A suggestion can be helpful when it is private and respectful, but hurtful when it sounds like public criticism.",
        "tip": "用 although / even if 寫出讓步，論述會更成熟。",
    },
    {
        "title": "提出可執行的溝通原則",
        "keywords": ["solution", "communicate", "respect", "listen", "ask", "溝通", "尊重", "傾聽", "建議"],
        "wrong": "People should not give advice.",
        "correct": "People should ask whether advice is welcome, listen first, and speak in a respectful way.",
        "explanation": "結論不要只禁止別人給建議；更好的寫法是提出可實行的溝通原則，讓文章有解決問題的方向。",
        "example": "Before giving advice, we can ask, 'Would you like to hear my suggestion?'",
        "tip": "結論回答「那應該怎麼做」，不要只重複問題。",
    },
    {
        "title": "避免連續列事件，讓例子服務同一論點",
        "keywords": ["structure", "paragraph", "topic sentence", "主題句", "段落", "轉承", "例子"],
        "wrong": "One event is about my aunt. Another event is about my classmate.",
        "correct": "Use one topic sentence to explain how both events show the same problem: advice can cross personal boundaries.",
        "explanation": "若連續列出親戚、同學或朋友的事件，文章容易變成流水帳。主體段應用主題句統整例子，說明它們共同證明什麼。",
        "example": "These experiences show that comments about grades, appearance, or relationships can become pressure when they are unsolicited.",
        "tip": "例子之間要有共同論點，不只是換人、換場景。",
    },
]

GENERIC_STRATEGY_RULES = [
    {
        "title": "每段只處理一個清楚重點",
        "keywords": ["段落", "結構", "主題句", "銜接", "架構"],
        "wrong": "The paragraph describes the picture, gives opinions, and adds a new example at the same time.",
        "correct": "Use one paragraph for the situation and another paragraph for your solution or opinion.",
        "explanation": "段落功能清楚，讀者才容易跟上論述。每段應有主題句，後面用例子或理由支撐。",
        "example": "The first paragraph can describe the problem, while the second paragraph explains two possible solutions.",
        "tip": "一段一重點，重點後面接例子或原因。",
    },
    {
        "title": "例子後要補原因、影響或反思",
        "keywords": ["例子", "具體", "原因", "影響", "反思", "說服力"],
        "wrong": "I give a solution, but I do not explain why it works.",
        "correct": "After giving a solution, explain how it solves the problem and who benefits from it.",
        "explanation": "只有例子仍不夠，還需要補上原因或影響，才能把內容推深。",
        "example": "This strategy would work because it addresses the main problem and gives readers a clear reason to agree.",
        "tip": "每個例子後補一句 because / so / as a result。",
    },
    {
        "title": "結論要回扣題目，不要只重複前文",
        "keywords": ["結論", "總結", "回扣", "題目", "完整"],
        "wrong": "That is my opinion.",
        "correct": "Restate your main point and connect it back to the task, instead of ending with a vague sentence.",
        "explanation": "結論應收束整篇文章，回到題目核心，並強化你的立場或解決方案。",
        "example": "In this way, the solution becomes clearer and the whole essay feels complete.",
        "tip": "最後一句要回答：所以這件事可以怎麼變好？",
    },
]


def correction_pattern_key(item: CorrectionItem) -> tuple[str, str, str]:
    return (
        item.category,
        normalize_group_key(item.wrong),
        normalize_group_key(clean_correction_variant(item.correction)),
    )


def is_dirty_correction(text: str) -> bool:
    value = compact_cjk(text)
    lower = value.lower()
    if value == DELETE_CORRECTION:
        return False
    if not value:
        return True
    if len(value) > 120:
        return True
    if value.count("(") != value.count(")") or value.count("[") != value.count("]"):
        return True
    if "..." in value or "…" in value:
        return True
    normalized_letters = re.sub(r"[^A-Za-z]", "", value.lower())
    if len(normalized_letters) >= 8:
        for size in range(4, min(24, len(normalized_letters) // 2) + 1):
            for index in range(0, len(normalized_letters) - (size * 2) + 1):
                chunk = normalized_letters[index : index + size]
                if chunk and chunk == normalized_letters[index + size : index + (size * 2)]:
                    return True
    if re.search(r"\b([A-Za-z]{4,})\s+\1\b", value, flags=re.I):
        return True
    if re.match(r"^[,.;:!?]\s*[A-Z]", value):
        return True
    if re.search(r"[\"“”「」].{8,}", value):
        return True
    if re.search(r"[A-Za-z][^。]*[\"“”「」]", value):
        return True
    if any(marker in lower for marker in ["comma splice", "run-on", "fragment"]):
        return True
    if len(re.findall(r"\b[A-Z][a-z]{3,}\b", value)) >= 3 and len(value.split()) >= 8:
        return True
    return False


def is_noisy_wrong(wrong: str) -> bool:
    key = normalize_group_key(wrong)
    if key in NOISY_WRONG_KEYS:
        return True
    if len(key) == 1 and not key.isalnum():
        return True
    return False


def is_unhelpful_pair(wrong: str, correct: str) -> bool:
    if compact_cjk(correct) == DELETE_CORRECTION:
        return False
    if is_capitalization_case_pair(wrong, correct):
        return False
    wrong_key = normalize_group_key(wrong)
    correct_key = normalize_group_key(correct)
    wrong_tokens = wrong_key.split()
    correct_tokens = correct_key.split()
    if len(wrong_tokens) > 1 and len(correct_tokens) == 1 and correct_key in SHORT_CORRECTION_ONLY_KEYS:
        return True
    if wrong_key == correct_key:
        return True
    return False


def meaningful_token_count(text: str) -> int:
    return len([tok for tok in normalize_lookup(text).split() if tok not in NOISY_WRONG_KEYS])


def clean_correction_variant(correction: str) -> str:
    value = compact_cjk(correction)
    value = re.split(r"\s*/\s*|；|;", value, maxsplit=1)[0].strip()
    return value


def clean_entry_value(value: str, limit: int = 140) -> str:
    value = compact_cjk(value)
    value = re.split(r"；|;", value, maxsplit=1)[0].strip()
    value = re.sub(r"\s+", " ", value)
    if len(value) > limit:
        value = value[: limit - 3].rstrip() + "..."
    return value


def encode_file_ids(file_ids: set[str] | list[str]) -> str:
    return ",".join(sorted({file_id for file_id in file_ids if file_id}))


def decode_file_ids(value: object) -> set[str]:
    if isinstance(value, list):
        return {str(item) for item in value if str(item)}
    return {part for part in str(value or "").split(",") if part}


def choose_representative(items: list[CorrectionItem]) -> CorrectionItem:
    clean_items = [item for item in items if not is_dirty_correction(item.correction)] or items
    correction_counts = Counter(normalize_group_key(clean_correction_variant(item.correction)) for item in clean_items)

    def rank(item: CorrectionItem) -> tuple[int, int, int, int]:
        correction_key = normalize_group_key(clean_correction_variant(item.correction))
        return (
            correction_counts[correction_key],
            1 if item.source_sentence.strip() else 0,
            1 if item.explanation.strip() else 0,
            -len(item.correction),
        )

    return max(clean_items, key=rank)


def make_rule_title(item: CorrectionItem) -> str:
    label = CATEGORY_LABELS.get(item.category, item.category)
    return f"{label}：{preview_text(item.wrong, 30)} → {preview_text(clean_correction_variant(item.correction), 30)}"


def make_rule_explanation(item: CorrectionItem, frequency: int) -> str:
    explanation = compact_cjk(item.explanation)
    if explanation:
        if not explanation.endswith(("。", ".", "！", "!", "?", "？")):
            explanation += "。"
        return explanation
    label = CATEGORY_LABELS.get(item.category, item.category)
    return f"這類{label}問題需要先判斷錯誤片語在句中的詞性、搭配與語意，再套用建議改法。"


def make_rule_example(item: CorrectionItem) -> str:
    source = item.source_sentence.strip()
    wrong = item.wrong.strip()
    correction = clean_correction_variant(item.correction)
    if correction == DELETE_CORRECTION:
        if source and wrong:
            pattern = re.escape(wrong)
            if re.match(r"^[A-Za-z0-9' -]+$", wrong):
                pattern = rf"(?<![A-Za-z0-9']){re.escape(wrong)}(?![A-Za-z0-9'])"
            replaced = re.sub(pattern, "", source, count=1, flags=re.I)
            replaced = re.sub(r"\s+([,.;:!?])", r"\1", replaced)
            replaced = re.sub(r"\s{2,}", " ", replaced).strip()
            if replaced != source:
                return preview_text(replaced, 180)
        return f"刪除：{wrong}"
    if len(wrong) > 70 or len(correction) > 90 or "." in correction:
        return f"修正方向：{wrong} → {correction}"
    if source and wrong and correction:
        pattern = re.escape(wrong)
        if re.match(r"^[A-Za-z0-9' -]+$", wrong):
            pattern = rf"(?<![A-Za-z0-9']){re.escape(wrong)}(?![A-Za-z0-9'])"
        replaced = re.sub(pattern, correction, source, count=1, flags=re.I)
        replaced = re.sub(r"\b([A-Za-z]{3,})\s+\1\b", r"\1", replaced, flags=re.I)
        if replaced != source:
            return preview_text(replaced, 180)
    return f"修正方向：{wrong} → {correction}"


def make_rule_tip(item: CorrectionItem, frequency: int) -> str:
    return CATEGORY_TIPS.get(item.category, "把這類錯誤整理成固定檢查點，寫完作文後逐句檢查。")


def build_entry_from_group(grouped_items: list[CorrectionItem]) -> dict[str, str]:
    representative = choose_representative(grouped_items)
    frequency = len(grouped_items)
    file_ids = {item.file_id for item in grouped_items if item.file_id}
    source_count = len(file_ids)
    return {
        "title": make_rule_title(representative),
        "wrong": representative.wrong,
        "correct": clean_correction_variant(representative.correction),
        "explanation": make_rule_explanation(representative, frequency),
        "example": make_rule_example(representative),
        "tip": make_rule_tip(representative, frequency),
        "source_sentence": representative.source_sentence,
        "category": representative.category,
        "frequency": str(frequency),
        "source_count": str(source_count),
        "file_ids": encode_file_ids(file_ids),
    }


def build_grouped_entries(items: list[CorrectionItem]) -> list[dict[str, str]]:
    groups: dict[tuple[str, str, str], list[CorrectionItem]] = defaultdict(list)
    for item in items:
        if not item.wrong.strip() or not item.correction.strip():
            continue
        groups[correction_pattern_key(item)].append(item)

    entries = [build_entry_from_group(grouped_items) for grouped_items in groups.values()]
    entries.sort(
        key=lambda entry: (
            int(entry["frequency"]),
            int(entry["source_count"]),
            meaningful_token_count(entry["wrong"]),
        ),
        reverse=True,
    )
    return entries


def dedupe_entry_pairs(entries: list[dict[str, str]]) -> list[dict[str, str]]:
    best_by_pair: dict[tuple[str, str], dict[str, str]] = {}
    for entry in entries:
        key = (normalize_group_key(entry["wrong"]), normalize_group_key(entry["correct"]))
        current = best_by_pair.get(key)
        if not current:
            best_by_pair[key] = entry
            continue
        current_score = (
            int(current["frequency"]),
            int(current.get("source_count", "0")),
            DEDUPE_CATEGORY_PRIORITY.get(current["category"], 0),
        )
        entry_score = (
            int(entry["frequency"]),
            int(entry.get("source_count", "0")),
            DEDUPE_CATEGORY_PRIORITY.get(entry["category"], 0),
        )
        if entry_score > current_score:
            best_by_pair[key] = entry
    deduped = list(best_by_pair.values())
    deduped.sort(
        key=lambda entry: (
            int(entry["frequency"]),
            int(entry["source_count"]),
            meaningful_token_count(entry["wrong"]),
        ),
        reverse=True,
    )
    return deduped


def is_main_rule_entry(entry: dict[str, str]) -> bool:
    frequency = int(entry.get("frequency", "0"))
    category = entry.get("category", "")
    wrong = entry.get("wrong", "")
    correct = entry.get("correct", "")
    token_count = meaningful_token_count(wrong)
    is_cluster_entry = entry.get("is_cluster") == "true" and bool(entry.get("patterns"))
    if (not is_cluster_entry and is_noisy_wrong(wrong)) or is_dirty_correction(correct):
        return False
    if is_unhelpful_pair(wrong, correct):
        return False
    if is_cluster_entry and not is_exact_teaching_family(entry.get("teaching_family", "")):
        return int(entry.get("source_count", "0") or 0) >= SUPPLEMENTAL_SOURCE_COUNT_THRESHOLD
    if category in {"Spelling", "Capitalization", "Punctuation"}:
        return False
    if category in {"Article", "Preposition"}:
        return frequency >= 2 and token_count >= 2
    if category in {"Agreement", "Tense"}:
        return frequency >= 2 or token_count >= 3
    if category in {"Word Choice", "Style", "Grammar", "Redundancy"}:
        return frequency >= 2 or token_count >= 4
    return frequency >= 2 and token_count >= 2


def pair_summary(entry: dict[str, str], *, include_frequency: bool = False) -> str:
    text = f"{preview_text(entry['wrong'], 34)} → {preview_text(entry['correct'], 34)}"
    if include_frequency:
        frequency = int(entry.get("frequency", "0") or 0)
        if frequency > 0:
            text += f"（出現 {frequency} 次）"
    return text


def limit_pattern_text(patterns: str, limit: int = CLUSTER_PATTERN_LIMIT, part_limit: int = 90) -> str:
    parts: list[str] = []
    seen: set[str] = set()
    for raw_part in re.split(r"；|;", patterns or ""):
        part = compact_cjk(raw_part)
        if not part or "→" not in part:
            continue
        key = normalize_group_key(strip_frequency_suffix(part))
        if key in seen:
            continue
        seen.add(key)
        parts.append(preview_text(part, part_limit))
        if len(parts) >= limit:
            break
    return "；".join(parts)


def is_category_list_candidate(entry: dict[str, str], category: str) -> bool:
    wrong = entry["wrong"]
    correct = entry["correct"]
    wrong_tokens = normalize_lookup(wrong).split()
    correct_tokens = normalize_lookup(correct).split()
    if category == "Spelling":
        if "'" in wrong + correct and "cafe" not in normalize_lookup(wrong + " " + correct):
            return False
        if len(wrong_tokens) != 1 or len(correct_tokens) != 1:
            return False
        return is_likely_spelling_pair(wrong_tokens[0], correct_tokens[0])
    if category == "Preposition":
        if len(wrong_tokens) > 5 or len(correct_tokens) > 6:
            return False
        rewrite_markers = {"thank", "hope", "would", "need", "fixed", "dear"}
        if rewrite_markers & set(correct_tokens):
            return False
        return bool(set(wrong_tokens) & set(correct_tokens))
    if category == "Punctuation":
        if "..." in wrong + correct or "…" in wrong + correct:
            return False
        if len(wrong) > 34 or len(correct) > 34:
            return False
        if "\"" in wrong + correct or "“" in wrong + correct or "”" in wrong + correct:
            return False
        if len(wrong_tokens) > 5 or len(correct_tokens) > 6:
            return False
        if wrong.strip() in {",", ".", ";", ":"} or correct.strip() in {",", ".", ";", ":"}:
            return True
        explanation = entry.get("explanation", "")
        if any(marker in explanation.lower() for marker in ["comma splice", "run-on", "fragment"]):
            return True
        if any(marker in explanation for marker in ["逗號連", "逗號點", "標點", "句點", "分號"]):
            return True
        wrong_without_punctuation = re.sub(r"[,.;:!?]", "", normalize_lookup(wrong))
        correct_without_punctuation = re.sub(r"[,.;:!?]", "", normalize_lookup(correct))
        return bool(wrong_without_punctuation) and wrong_without_punctuation == correct_without_punctuation
    return True


def build_pair_list_entry(title: str, category: str, entries: list[dict[str, str]], limit: int) -> dict[str, str] | None:
    usable = [
        entry
        for entry in entries
        if not is_dirty_correction(entry["correct"])
        and not is_noisy_wrong(entry["wrong"])
        and not is_unhelpful_pair(entry["wrong"], entry["correct"])
        and is_category_list_candidate(entry, category)
    ]
    if not usable:
        return None
    usable.sort(
        key=lambda entry: (
            int(entry["frequency"]),
            int(entry["source_count"]),
            meaningful_token_count(entry["wrong"]),
        ),
        reverse=True,
    )
    selected = usable[:limit]
    wrongs = "；".join(preview_text(entry["wrong"], 34) for entry in selected)
    corrections = "；".join(preview_text(entry["correct"], 34) for entry in selected)
    pairs = "；".join(pair_summary(entry, include_frequency=True) for entry in selected)
    total_frequency = sum(int(entry["frequency"]) for entry in selected)
    file_ids: set[str] = set()
    for entry in selected:
        file_ids.update(decode_file_ids(entry.get("file_ids", "")))
    total_sources = len(file_ids) if file_ids else sum(int(entry.get("source_count", "0")) for entry in selected)
    source_sentence = next((entry.get("source_sentence", "") for entry in selected if entry.get("source_sentence")), "")
    label = CATEGORY_LABELS.get(category, category)
    return {
        "title": title,
        "wrong": wrongs,
        "correct": corrections,
        "explanation": f"這些是可直接合併複習的{label}修正。它們多半不需要各自獨立成一條規則，適合整理成考前檢查清單。",
        "example": pairs,
        "tip": CATEGORY_TIPS.get(category, "把這類錯誤整理成固定檢查點，寫完作文後逐句檢查。"),
        "source_sentence": source_sentence,
        "category": category,
        "frequency": str(total_frequency),
        "source_count": str(total_sources),
        "file_ids": encode_file_ids(file_ids),
        "patterns": pairs,
        "is_list": "true",
        "source_lookup": "disabled",
    }


def normalized_word_tokens(text: str) -> list[str]:
    return [token for token in normalize_lookup(text).split() if token]


def single_word(text: str) -> str:
    tokens = normalized_word_tokens(text)
    if len(tokens) != 1:
        return ""
    token = tokens[0]
    return token if re.match(r"^[a-z][a-z']*$", token) else ""


def regular_plural_bases(word: str) -> set[str]:
    bases = {word}
    if word.endswith("ies") and len(word) > 4:
        bases.add(word[:-3] + "y")
    if word.endswith("es") and len(word) > 4:
        bases.add(word[:-2])
    if word.endswith("s") and len(word) > 3 and not word.endswith("ss"):
        bases.add(word[:-1])
    return bases


def regular_plural_base(word: str) -> str:
    return sorted(regular_plural_bases(word), key=len)[0]


def is_regular_plural_word_pair(wrong: str, correct: str) -> bool:
    wrong_word = single_word(wrong)
    correct_word = single_word(correct)
    if not wrong_word or not correct_word or wrong_word == correct_word:
        return False
    if wrong_word.endswith("s") and not correct_word.endswith("s"):
        return False
    return wrong_word in regular_plural_bases(correct_word)


def phrase_has_regular_plural_change(wrong: str, correct: str) -> bool:
    wrong_tokens = normalized_word_tokens(wrong)
    correct_tokens = normalized_word_tokens(correct)
    if not wrong_tokens or not correct_tokens:
        return False
    for wrong_token in wrong_tokens:
        for correct_token in correct_tokens:
            if is_regular_plural_word_pair(wrong_token, correct_token):
                return True
    return False


def article_count(tokens: list[str]) -> int:
    return sum(1 for token in tokens if token in ARTICLES)


def has_inserted_article(wrong: str, correct: str) -> bool:
    wrong_tokens = normalized_word_tokens(wrong)
    correct_tokens = normalized_word_tokens(correct)
    if not wrong_tokens or not correct_tokens:
        return False
    return article_count(correct_tokens) > article_count(wrong_tokens)


def has_article_choice_change(wrong: str, correct: str) -> bool:
    wrong_tokens = normalized_word_tokens(wrong)
    correct_tokens = normalized_word_tokens(correct)
    if not wrong_tokens or not correct_tokens:
        return False
    shared_article_count = min(len(wrong_tokens), len(correct_tokens))
    for index in range(shared_article_count):
        if wrong_tokens[index] in ARTICLES and correct_tokens[index] in ARTICLES and wrong_tokens[index] != correct_tokens[index]:
            return True
    return bool(wrong_tokens and correct_tokens and wrong_tokens[0] in ARTICLES and correct_tokens[0] in ARTICLES and wrong_tokens[0] != correct_tokens[0])


def has_subject_verb_marker(text: str) -> bool:
    return any(marker in text for marker in ["主謂", "主詞", "動詞", "第三人稱", "單數主詞"])


def has_noun_plural_marker(text: str) -> bool:
    return any(marker in text for marker in ["可數", "複數名詞", "複數形", "單複數", "名詞必須使用複數", "使用複數"])


def has_ving_pair(wrong: str, correct: str) -> bool:
    wrong_word = single_word(wrong)
    correct_word = single_word(correct)
    if not wrong_word or not correct_word or not correct_word.endswith("ing"):
        return False
    base = correct_word[:-3]
    if len(base) >= 3 and base[-1] == base[-2] and base[-1] not in "aeiou":
        if base[:-1] == wrong_word:
            return True
    return wrong_word in inflection_stems(correct_word)


def is_capitalization_case_pair(wrong: str, correct: str) -> bool:
    wrong_clean = compact_cjk(wrong)
    correct_clean = compact_cjk(correct)
    return bool(wrong_clean and correct_clean and wrong_clean.lower() == correct_clean.lower() and wrong_clean != correct_clean)


def normalized_entry_category(entry: dict[str, str]) -> str:
    category = entry.get("category", "")
    wrong = entry.get("wrong", "")
    correct = entry.get("correct", "")
    wrong_key = normalize_group_key(wrong)
    correct_key = normalize_group_key(correct)
    text = entry_family_text(entry).lower()
    if is_capitalization_case_pair(wrong, correct):
        return "Capitalization"
    if category == "Tense" and is_likely_spelling_pair(wrong_key, correct_key):
        return "Spelling"
    if category == "Tense" and wrong_key == "crowed" and correct_key == "crowded":
        return "Spelling"
    if category == "Punctuation" and any(marker in text for marker in ["with chat", "介系詞", "介係詞"]):
        return "Preposition"
    return category


def spelling_teaching_family(wrong: str, correct: str) -> str:
    wrong_key = normalize_group_key(wrong)
    correct_key = normalize_group_key(correct)
    joined = f"{wrong_key} {correct_key}"
    if any(token in joined for token in ["costumer", "costomers", "costumers", "customs", "customer", "customers"]):
        return "spelling_customer_list"
    if any(token in joined for token in ["chatting", "chating", "studying", "studing"]):
        return "spelling_ing_list"
    if any(token in joined for token in ["cafe", "café", "coffee", "coffe"]):
        return "spelling_cafe_coffee_list"
    if any(token in joined for token in ["there", "their", "sigh", "sign", "scent", "scene", "desert", "dessert"]):
        return "spelling_confusable_words_list"
    if any(
        token in joined
        for token in [
            "business",
            "bussiness",
            "approach",
            "appoarch",
            "address",
            "adress",
            "minimum",
            "minimun",
            "modern",
            "morden",
            "comfortable",
            "confortable",
            "inconvenience",
            "incovenience",
            "separate",
            "seperate",
        ]
    ):
        return "spelling_topic_words_list"
    if any(token in joined for token in ["people", "peope", "other", "othe", "girl", "girls", "grils"]):
        return "spelling_people_words_list"
    return "spelling_common_words_list"


def preposition_teaching_family(wrong: str, correct: str) -> str:
    joined = f"{normalize_group_key(wrong)} {normalize_group_key(correct)}"
    if "cup of drink" in joined or "a cup of drink" in joined:
        return "preposition_quantity_list"
    if "on the table" in joined or "at the table" in joined:
        return "preposition_table_position_list"
    if "in the picture" in joined or "on the picture" in joined or joined in {"in on", "on in"}:
        return "preposition_picture_context_list"
    if "perspective" in joined:
        return "preposition_perspective_list"
    if "because" in joined or "because of" in joined:
        return "preposition_because_of_list"
    if "occupied with" in joined or "occupied by" in joined:
        return "preposition_passive_by_list"
    if any(token in joined for token in ["left", "right", "corner", "table", "front", "behind", "beside"]):
        return "preposition_location_list"
    if any(token in joined for token in ["working", "work", "sit", "sitting"]):
        return "preposition_verb_phrase_list"
    return "preposition_list"


def punctuation_teaching_family(entry: dict[str, str]) -> str:
    wrong = entry.get("wrong", "")
    correct = entry.get("correct", "")
    text = entry_family_text(entry)
    text_lower = text.lower()
    wrong_key = normalize_group_key(wrong)
    correct_key = normalize_group_key(correct)
    if any(marker in text_lower for marker in ["comma splice", "逗號連", "逗號點", "獨立句子"]) or wrong_key in {",", ", this", "there are", "there is"} or correct_key.startswith("."):
        return "punctuation_comma_splice_list"
    if any(token in wrong_key for token in ["however", "but", "and", "because", "furthermore"]) or any(token in correct_key for token in ["however", "but", "and", "because", "furthermore"]):
        return "punctuation_transition_list"
    if any(marker in text for marker in ["列舉", "最後一項", "對等連接詞"]) or any(token in wrong_key for token in ["still others", "others", "some"]):
        return "punctuation_listing_list"
    if any(marker in text for marker in ["空格", "標點符號", "逗號前", "句點"]) and not any(marker in text_lower for marker in ["comma splice", "逗號連"]):
        return "punctuation_spacing_list"
    if any(token in wrong_key for token in ["because", "such as"]) or "fragment" in text_lower or "從屬子句" in text:
        return "punctuation_fragment_list"
    return "punctuation_list"


def word_choice_teaching_family(entry: dict[str, str]) -> str:
    wrong = entry.get("wrong", "")
    correct = entry.get("correct", "")
    text = entry_family_text(entry)
    text_lower = text.lower()
    wrong_key = normalize_group_key(wrong)
    correct_key = normalize_group_key(correct)
    joined = f"{wrong_key} {correct_key} {text_lower}"
    if any(marker in text for marker in ["贅字", "重複", "語意重複", "Redundant"]) or correct == DELETE_CORRECTION:
        return "word_choice_redundancy"
    if any(token in joined for token in ["cause", "lead", "let", "make", "prevent", "leave", "導致", "造成", "使役", "讓"]):
        return "word_choice_causative_result"
    if any(
        token in joined
        for token in [
            "time limit",
            "minimum purchase",
            "turnover",
            "occupy",
            "seating",
            "seat hogging",
            "coffee shop",
            "cafe",
            "cup of coffee",
        ]
    ):
        return "word_choice_cafe_terms"
    if any(marker in text for marker in ["中式", "直譯", "中文"]) or any(token in joined for token in ["by only", "exchange", "gain no one"]):
        return "word_choice_chinese_translation"
    if any(marker in text for marker in ["搭配", "慣用", "自然", "collocation"]) or any(token in joined for token in ["create", "set", "from my perspective", "but also"]):
        return "word_choice_collocation"
    if any(marker in text for marker in ["語意不清", "語意混淆", "語意錯置", "不自然", "模糊", "精準"]):
        return "word_choice_precision"
    return exact_teaching_family(entry)


def entry_family_text(entry: dict[str, str]) -> str:
    return compact_cjk(
        " ".join(
            [
                entry.get("wrong", ""),
                entry.get("correct", ""),
                entry.get("explanation", ""),
            ]
        )
    )


def exact_teaching_family(entry: dict[str, str]) -> str:
    return f"exact:{normalize_group_key(entry.get('wrong', ''))}->{normalize_group_key(entry.get('correct', ''))}"


def teaching_family_for_entry(entry: dict[str, str]) -> str:
    category = normalized_entry_category(entry)
    wrong = entry.get("wrong", "")
    correct = entry.get("correct", "")
    text = entry_family_text(entry)
    text_lower = text.lower()

    if category != entry.get("category", ""):
        entry["category"] = category

    if category in {"Word Choice", "Style", "Redundancy"}:
        family = word_choice_teaching_family(entry)
        if not is_exact_teaching_family(family):
            return family

    if category == "Spelling" and is_category_list_candidate(entry, category):
        return spelling_teaching_family(wrong, correct)
    if category == "Capitalization" and is_category_list_candidate(entry, category):
        return "capitalization_list"
    if category == "Punctuation":
        return punctuation_teaching_family(entry)
    if category == "Preposition" and is_category_list_candidate(entry, category):
        return preposition_teaching_family(wrong, correct)

    if category == "Agreement":
        wrong_key = normalize_group_key(wrong)
        correct_key = normalize_group_key(correct)
        if wrong_key.startswith("there ") or correct_key.startswith("there ") or "there is" in text_lower or "there are" in text_lower:
            return "there_be_agreement"
        if wrong_key in {"other", "the others"} or correct_key in {"others", "other people"} or "some... others" in text_lower:
            return "some_others"
        if any(token in {wrong_key, correct_key} for token in ["it", "them", "this", "these", "you", "they"]) or "代名詞" in text:
            return "pronoun_agreement"
        if any(marker in text_lower for marker in ["everyone", "another", "other people", "every people"]) or any(key in {wrong_key, correct_key} for key in ["another", "others", "other"]):
            return "indefinite_pronoun_agreement"
        if "缺少主詞" in text or wrong_key in {"hope"} or correct_key.startswith("i hope"):
            return "subject_needed"
        if wrong_key in {"is", "are", "was", "were", "there's", "there is"} or correct_key in {"is", "are", "was", "were", "there are"}:
            return "be_verb_agreement"
        if phrase_has_regular_plural_change(wrong, correct) and any(marker in text_lower for marker in ["two", "many", "some", "several", "no", "enough"]):
            return "plural_after_number_quantifier"
        if has_subject_verb_marker(text):
            if single_word(wrong) and single_word(correct) and correct_key.endswith("s") and not wrong_key.endswith("s"):
                return "third_person_singular_verb"
            return "subject_verb_agreement"
        if phrase_has_regular_plural_change(wrong, correct) and has_noun_plural_marker(text):
            return "countable_plural"
        if phrase_has_regular_plural_change(wrong, correct):
            return "countable_plural"
        if single_word(wrong) and single_word(correct) and correct_key.endswith("s") and not wrong_key.endswith("s"):
            return "third_person_singular_verb"

    if category == "Article":
        wrong_key = normalize_group_key(wrong)
        correct_key = normalize_group_key(correct)
        if has_article_choice_change(wrong, correct) and any(key in f"{wrong_key} {correct_key}" for key in [" an ", " an", "an "]):
            return "article_a_an_choice"
        if any(token in f"{wrong_key} {correct_key}" for token in ["coffee shop", "cafe", "café", "cafeteria", "place"]):
            if has_inserted_article(wrong, correct):
                return "article_place_single"
            if phrase_has_regular_plural_change(wrong, correct):
                return "generic_plural_countable"
        if any(token in f"{wrong_key} {correct_key}" for token in ["long time", "whole day"]):
            return "article_time_expression"
        if any(token in f"{wrong_key} {correct_key}" for token in ["seat", "book", "laptop", "look", "nap", "drink", "crowd"]):
            if has_inserted_article(wrong, correct):
                return "article_object_single"
        if any(token in f"{wrong_key} {correct_key}" for token in ["picture", "boss", "owner", "public"]):
            return "article_specific_context"
        if any(token in f"{wrong_key} {correct_key}" for token in ["work", "works", "advice", "information", "food"]):
            return "article_uncountable_noun"
        if phrase_has_regular_plural_change(wrong, correct) and ("泛指" in text or "所有" in text or "類" in text):
            return "generic_plural_countable"
        if has_inserted_article(wrong, correct):
            return "article_missing"
        if has_article_choice_change(wrong, correct):
            return "article_choice"
        if "可數" in text and "冠詞" in text:
            return "article_missing"

    if category == "Grammar":
        wrong_key = normalize_group_key(wrong)
        correct_key = normalize_group_key(correct)
        if any(key in f"{wrong_key} {correct_key}" for key in ["sit", "seat", "sits", "seats"]):
            return "sit_seat_word_form"
        if any(marker in text_lower for marker in ["there is + n", "there are + n", "there is", "there are"]) and has_ving_pair(wrong, correct):
            return "there_is_n_ving"
        if any(marker in text for marker in ["一句兩動詞", "第二個動詞", "主要動詞"]) and has_ving_pair(wrong, correct):
            return "second_verb_ving"
        if ("介詞" in text or "介係詞" in text or any(token in wrong_key for token in ["instead of", "without", "of "])) and (
            has_ving_pair(wrong, correct) or correct_key.endswith("ing")
        ):
            return "after_preposition_ving"
        if any(marker in text for marker in ["對等", "平行", "列舉"]) and (has_ving_pair(wrong, correct) or correct_key.endswith("ing")):
            return "parallel_ving"
        if any(key in f"{wrong_key} {correct_key}" for key in ["fulled", "full", "depress", "depressed", "relax", "relaxing", "crowed", "crowded"]):
            return "adjective_participle_form"
        if wrong_key.startswith("there have") or correct_key.startswith("there is") or correct_key.startswith("there are"):
            if "there have" in wrong_key:
                return "there_have_structure"
        if any(key in {wrong_key, correct_key} for key in ["so", "but"]) and any(marker in text for marker in ["Because", "Although", "雖然", "因為"]):
            return "because_so_although_but"
        if wrong_key in {"that", "where"} or correct_key == "where" or "關係子句" in text or "關係副詞" in text:
            return "relative_clause_place"
        if "no any" in wrong_key:
            return "quantifier_no_any"
        subjunctive_pair_keys = {
            "if i am",
            "if i'm",
            "if i was",
            "i will",
            "i was",
            "will",
            "i'll",
            "if i were",
            "i would",
            "i were",
            "would",
        }
        if ("假設語氣" in text or "subjunctive" in text_lower) and (
            wrong_key in subjunctive_pair_keys or correct_key in subjunctive_pair_keys
        ):
            return "subjunctive_if"
        if "助動詞" in text or "modal" in text_lower or re.search(r"\b(would|could|should|must|can)\s*(?:\+|後|後面|接)", text_lower):
            return "modal_base_verb"
        if "假設語氣" in text or "subjunctive" in text_lower or normalize_group_key(wrong).startswith("if i ") or normalize_group_key(correct).startswith("if i "):
            return "subjunctive_if"
        if "v-ing" in text_lower or "現在分詞" in text or "動名詞" in text or has_ving_pair(wrong, correct):
            return "participle_ving"

    if category == "Tense":
        wrong_key = normalize_group_key(wrong)
        correct_key = normalize_group_key(correct)
        if "假設語氣" in text or "subjunctive" in text_lower:
            return "subjunctive_tense"
        if any(token in f"{wrong_key} {correct_key}" for token in ["haven't", "hasn't", "have not", "has not", "has just", "have just"]):
            return "present_perfect_current_result"
        if any(key in {wrong_key, correct_key} for key in ["are sleeping", "is sleeping", "are studying", "are chatting"]) or (
            correct_key.endswith("ing") and "圖片" in text
        ):
            return "picture_present_progressive"
        if any(
            pair in f"{wrong_key}->{correct_key}"
            for pair in ["didn't->don't", "were->are", "was->is", "could->can", "occupied->occupy", "chose->choose"]
        ):
            return "picture_present_tense"
        if "圖片" in text and ("現在式" in text or "目前" in text or "事實陳述" in text):
            return "picture_present_tense"

    return exact_teaching_family(entry)


def is_exact_teaching_family(family: str) -> bool:
    return family.startswith("exact:")


def cluster_sort_key(entry: dict[str, str]) -> tuple[int, int, int, int]:
    return (
        int(entry.get("source_count", "0") or 0),
        int(entry.get("frequency", "0") or 0),
        DEDUPE_CATEGORY_PRIORITY.get(entry.get("category", ""), 0),
        meaningful_token_count(entry.get("wrong", "")),
    )


def pattern_members_for_cluster(members: list[dict[str, str]], *, limit: int = CLUSTER_PATTERN_LIMIT) -> list[dict[str, str]]:
    selected: list[dict[str, str]] = []
    seen_pairs: set[tuple[str, str]] = set()
    for entry in sorted(members, key=cluster_sort_key, reverse=True):
        if int(entry.get("source_count", "0") or 0) < CLUSTER_MEMBER_PATTERN_MIN_SOURCE and selected:
            continue
        pair_key = (normalize_group_key(entry.get("wrong", "")), normalize_group_key(entry.get("correct", "")))
        if pair_key in seen_pairs:
            continue
        seen_pairs.add(pair_key)
        selected.append(entry)
        if len(selected) >= limit:
            break
    return selected


def representative_cluster_entry(members: list[dict[str, str]]) -> dict[str, str]:
    return max(
        members,
        key=lambda entry: (
            int(entry.get("source_count", "0") or 0),
            int(entry.get("frequency", "0") or 0),
            1 if entry.get("source_sentence", "") else 0,
            meaningful_token_count(entry.get("wrong", "")),
        ),
    )


def build_knowledge_cluster(cluster_key: str, category: str, teaching_family: str, members: list[dict[str, str]]) -> KnowledgeCluster | None:
    if not members:
        return None
    representative = representative_cluster_entry(members)
    file_ids: set[str] = set()
    for entry in members:
        file_ids.update(decode_file_ids(entry.get("file_ids", "")))
    frequency = sum(int(entry.get("frequency", "0") or 0) for entry in members)
    source_count = len(file_ids) if file_ids else sum(int(entry.get("source_count", "0") or 0) for entry in members)
    patterns = [pair_summary(entry, include_frequency=True) for entry in pattern_members_for_cluster(members)]
    title = TEACHING_FAMILY_META.get(teaching_family, {}).get("title") or representative.get("title", "")
    source_sentence = next((entry.get("source_sentence", "") for entry in sorted(members, key=cluster_sort_key, reverse=True) if entry.get("source_sentence")), "")
    return KnowledgeCluster(
        cluster_key=cluster_key,
        category=category,
        teaching_family=teaching_family,
        title=title,
        member_entries=members,
        frequency=frequency,
        source_count=source_count,
        patterns=patterns,
        representative_source_sentence=source_sentence or representative.get("source_sentence", ""),
    )


def knowledge_cluster_to_entry(cluster: KnowledgeCluster) -> dict[str, str]:
    representative = representative_cluster_entry(cluster.member_entries)
    meta = TEACHING_FAMILY_META.get(cluster.teaching_family, {})
    selected_patterns = pattern_members_for_cluster(cluster.member_entries)
    wrongs = "；".join(preview_text(entry.get("wrong", ""), 34) for entry in selected_patterns)
    corrections = "；".join(preview_text(entry.get("correct", ""), 34) for entry in selected_patterns)
    patterns = "；".join(cluster.patterns)
    is_list = cluster.category in LIST_ENTRY_CATEGORIES
    explanation = clean_teacher_text(str(meta.get("explanation") or representative.get("explanation", "")), 320)
    if not explanation:
        explanation = make_rule_explanation(
            CorrectionItem("", 0, cluster.category, representative.get("wrong", ""), representative.get("correct", ""), "", cluster.representative_source_sentence, ""),
            max(1, cluster.frequency),
        )
    example = clean_teacher_text(str(meta.get("example") or representative.get("example", "")), 220)
    if not example:
        example = make_rule_example(
            CorrectionItem("", 0, cluster.category, representative.get("wrong", ""), representative.get("correct", ""), "", cluster.representative_source_sentence, "")
        )
    return {
        "title": clean_teacher_text(cluster.title or representative.get("title", ""), 90),
        "wrong": wrongs if is_list else representative.get("wrong", ""),
        "correct": corrections if is_list else representative.get("correct", ""),
        "explanation": explanation,
        "example": patterns if is_list else example,
        "tip": clean_teacher_text(str(meta.get("tip") or representative.get("tip", "")), 180)
        or CATEGORY_TIPS.get(cluster.category, "把這類錯誤整理成固定檢查點，寫完作文後逐句檢查。"),
        "source_sentence": cluster.representative_source_sentence,
        "category": cluster.category,
        "frequency": str(cluster.frequency),
        "source_count": str(cluster.source_count),
        "file_ids": encode_file_ids(
            {
                file_id
                for entry in cluster.member_entries
                for file_id in decode_file_ids(entry.get("file_ids", ""))
            }
        ),
        "patterns": patterns,
        "is_list": "true" if is_list else "false",
        "is_cluster": "true",
        "cluster_key": cluster.cluster_key,
        "teaching_family": cluster.teaching_family,
        "member_count": str(len(cluster.member_entries)),
        "source_lookup": "disabled",
    }


def build_knowledge_cluster_entries(
    items: list[CorrectionItem],
    *,
    min_source_count: int | None = None,
    list_min_source_count: int | None = None,
) -> list[dict[str, str]]:
    effective_count = effective_essay_count(items=items)
    min_source_count = min_source_count or teacher_correction_source_threshold(effective_count)
    list_min_source_count = list_min_source_count or teacher_list_source_threshold(effective_count)
    exact_entries = dedupe_entry_pairs(build_grouped_entries(items))
    groups: dict[tuple[str, str], list[dict[str, str]]] = defaultdict(list)
    for entry in exact_entries:
        category = normalized_entry_category(entry)
        entry = {**entry, "category": category}
        if category == "Writing Strategy":
            continue
        if category not in LIST_ENTRY_CATEGORIES and is_noisy_wrong(entry.get("wrong", "")):
            continue
        if is_dirty_correction(entry.get("correct", "")):
            continue
        if is_unhelpful_pair(entry.get("wrong", ""), entry.get("correct", "")) and not (
            category in LIST_ENTRY_CATEGORIES and is_category_list_candidate(entry, category)
        ):
            continue
        teaching_family = teaching_family_for_entry(entry)
        if category in LIST_ENTRY_CATEGORIES and not teaching_family.endswith("_list"):
            continue
        groups[(category, teaching_family)].append(entry)

    clusters: list[dict[str, str]] = []
    for (category, teaching_family), members in groups.items():
        cluster_key = f"{category}:{teaching_family}"
        cluster = build_knowledge_cluster(cluster_key, category, teaching_family, members)
        if not cluster:
            continue
        threshold = list_min_source_count if category in LIST_ENTRY_CATEGORIES else min_source_count
        if cluster.source_count < threshold:
            continue
        entry = knowledge_cluster_to_entry(cluster)
        if is_exact_teaching_family(teaching_family) and not is_main_rule_entry(entry):
            continue
        clusters.append(entry)

    clusters.sort(key=teacher_language_rank, reverse=True)
    return clusters


def count_topic_term(corpus: str, raw_corpus: str, term: str) -> int:
    if re.search(r"[\u4e00-\u9fff]", term):
        return raw_corpus.count(term.lower())
    normalized = normalize_lookup(term)
    if not normalized:
        return 0
    return len(re.findall(rf"(?<![a-z0-9']){re.escape(normalized)}(?![a-z0-9'])", corpus))


def detect_batch_topic(records: list[EssayRecord], items: list[CorrectionItem]) -> str:
    corpus_parts: list[str] = []
    for record in records:
        corpus_parts.append(record.teacher_feedback)
    for item in items:
        corpus_parts.extend([item.wrong, item.correction, item.source_sentence])
    raw_corpus = compact_cjk(" ".join(corpus_parts)).lower()
    corpus = normalize_lookup(raw_corpus)
    topic_terms = {
        "landmark_promotion": [
            "landmark",
            "landmarks",
            "landscape",
            "scenery",
            "natural",
            "culture",
            "tourist",
            "tourism",
            "social media",
            "campaign",
            "promote",
            "promotion",
            "global wonders",
            "voting",
            "competition",
            "recognition",
            "景點",
            "自然",
            "文化",
            "社群媒體",
            "推廣",
            "宣傳",
            "投票",
            "比賽",
            "國際",
            "全球",
            "代表台灣",
        ],
        "advice_boundaries": [
            "advice",
            "advices",
            "unsolicited",
            "suggestion",
            "suggestions",
            "privacy",
            "private",
            "boundary",
            "boundaries",
            "intrusion",
            "intrusive",
            "well meaning",
            "well-meaning",
            "good intention",
            "positive intention",
            "relative",
            "aunt",
            "classmate",
            "comment",
            "comments",
            "建議",
            "隱私",
            "界線",
            "干涉",
            "善意",
            "親戚",
            "同學",
            "評論",
        ],
        "grade_appeal": [
            "grade",
            "grading",
            "final grade",
            "teacher",
            "mr lin",
            "ms yang",
            "assignment",
            "exam score",
            "成績",
            "評分",
            "老師",
            "申訴",
        ],
        "cafe_seating": [
            "coffee",
            "cafe",
            "café",
            "coffee shop",
            "seat",
            "seats",
            "customer",
            "customers",
            "occupy",
            "occupied",
            "table",
            "drink",
            "咖啡",
            "咖啡廳",
            "咖啡店",
            "座位",
            "顧客",
            "佔用",
        ],
    }
    scores = {
        topic: sum(count_topic_term(corpus, raw_corpus, term) for term in terms)
        for topic, terms in topic_terms.items()
    }
    if scores["landmark_promotion"] >= 35 and scores["landmark_promotion"] > max(
        scores["advice_boundaries"], scores["grade_appeal"], scores["cafe_seating"]
    ):
        return "landmark_promotion"
    if scores["advice_boundaries"] >= 35 and scores["advice_boundaries"] > max(scores["grade_appeal"], scores["cafe_seating"]):
        return "advice_boundaries"
    if scores["cafe_seating"] >= 30 and scores["cafe_seating"] > scores["grade_appeal"]:
        return "cafe_seating"
    if scores["grade_appeal"] >= 20 and scores["grade_appeal"] >= scores["cafe_seating"]:
        return "grade_appeal"
    return "generic"


def select_strategy_rules(records: list[EssayRecord], items: list[CorrectionItem]) -> list[dict[str, object]]:
    topic = detect_batch_topic(records, items)
    if topic == "landmark_promotion":
        return LANDMARK_STRATEGY_RULES
    if topic == "advice_boundaries":
        return ADVICE_STRATEGY_RULES
    if topic == "cafe_seating":
        return CAFE_STRATEGY_RULES
    if topic == "grade_appeal":
        return STRATEGY_RULES
    return GENERIC_STRATEGY_RULES


def find_strategy_source_sentence(rule: dict[str, object], items: list[CorrectionItem]) -> str:
    raw_keywords = rule.get("source_keywords") or rule.get("keywords") or []
    keywords = [normalize_lookup(str(keyword)) for keyword in raw_keywords if normalize_lookup(str(keyword))]
    if not keywords:
        return ""
    best_sentence = ""
    best_score = 0
    for item in items:
        source = item.source_sentence.strip()
        if not source:
            continue
        source_key = normalize_lookup(source)
        score = 0
        for keyword in keywords:
            if re.search(rf"(?<![a-z0-9']){re.escape(keyword)}(?![a-z0-9'])", source_key):
                score += 3 if " " in keyword else 1
        if score > best_score:
            best_sentence = source
            best_score = score
    return preview_text(best_sentence, 260) if best_score > 0 else ""


def build_strategy_entries(records: list[EssayRecord], items: list[CorrectionItem]) -> list[dict[str, str]]:
    feedback_texts = [record.teacher_feedback for record in records if record.teacher_feedback.strip()]
    if not feedback_texts:
        return []
    threshold = teacher_strategy_source_threshold(effective_essay_count(records))
    topic_rules = list(select_strategy_rules(records, items))
    rules = list(topic_rules)
    if topic_rules != GENERIC_STRATEGY_RULES:
        existing_titles = {str(rule["title"]) for rule in rules}
        rules.extend(rule for rule in GENERIC_STRATEGY_RULES if str(rule["title"]) not in existing_titles)
    entries: list[dict[str, str]] = []
    scored_entries: list[tuple[int, dict[str, str]]] = []
    for rule in rules:
        count = 0
        for feedback in feedback_texts:
            if any(keyword.lower() in feedback.lower() for keyword in rule["keywords"]):
                count += 1
        entry = {
            "title": str(rule["title"]),
            "wrong": str(rule["wrong"]),
            "correct": str(rule["correct"]),
            "explanation": str(rule["explanation"]),
            "example": str(rule["example"]),
            "tip": str(rule["tip"]),
            "source_sentence": find_strategy_source_sentence(rule, items),
            "category": "Writing Strategy",
            "frequency": str(count),
            "source_count": str(count),
            "source_lookup": "disabled",
        }
        scored_entries.append((count, entry))
        if count >= threshold:
            entries.append(entry)
    if len(entries) < TEACHER_TARGET_ENTRIES_PER_UNIT:
        selected_titles = {entry["title"] for entry in entries}
        scored_entries.sort(key=lambda pair: pair[0], reverse=True)
        for count, entry in scored_entries:
            if len(entries) >= TEACHER_TARGET_ENTRIES_PER_UNIT:
                break
            if count <= 0 or entry["title"] in selected_titles:
                continue
            selected_titles.add(entry["title"])
            entries.append(entry)
    entries.sort(key=lambda entry: int(entry["frequency"]), reverse=True)
    return entries[:5]


def build_manual_units(records: list[EssayRecord], items: list[CorrectionItem]) -> list[dict]:
    grouped_entries = build_grouped_entries(items)
    if not grouped_entries:
        return MANUAL_UNITS

    effective_count = effective_essay_count(records, items)
    cluster_candidates = build_knowledge_cluster_entries(
        items,
        min_source_count=teacher_correction_source_threshold(effective_count),
        list_min_source_count=teacher_list_source_threshold(effective_count),
    )
    main_candidates = [
        entry
        for entry in cluster_candidates
        if entry.get("category") not in LIST_ENTRY_CATEGORIES and is_main_rule_entry(entry)
    ]
    main_candidates.sort(key=teacher_language_rank, reverse=True)
    selected_ids: set[int] = set()
    main_entries: list[dict[str, str]] = []
    for category in MAIN_CATEGORY_ORDER:
        category_entries = [entry for entry in main_candidates if entry["category"] == category]
        for entry in category_entries[: CATEGORY_MAIN_LIMITS.get(category, 0)]:
            if len(main_entries) >= MAX_MAIN_ENTRIES:
                break
            selected_ids.add(id(entry))
            main_entries.append(entry)
        if len(main_entries) >= MAX_MAIN_ENTRIES:
            break
    for entry in main_candidates:
        if len(main_entries) >= MAX_MAIN_ENTRIES:
            break
        if id(entry) in selected_ids:
            continue
        selected_ids.add(id(entry))
        main_entries.append(entry)

    entries_by_category: dict[str, list[dict[str, str]]] = defaultdict(list)
    for entry in main_entries:
        entries_by_category[entry["category"]].append(entry)

    for entry in cluster_candidates:
        if entry.get("category") in LIST_ENTRY_CATEGORIES:
            entries_by_category[entry["category"]].append(entry)

    used_categories: set[str] = set()
    units: list[dict] = []
    for unit in UNIT_GROUPS:
        unit_entries: list[dict[str, str]] = []
        for category in unit["categories"]:
            unit_entries.extend(entries_by_category.get(category, []))
            used_categories.add(category)
        unit_entries.sort(key=teacher_language_rank, reverse=True)
        if unit_entries:
            units.append({**unit, "entries": unit_entries})

    remaining_entries: list[dict[str, str]] = []
    for category, category_entries in entries_by_category.items():
        if category not in used_categories:
            remaining_entries.extend(category_entries)
    if remaining_entries:
        remaining_entries.sort(key=teacher_language_rank, reverse=True)
        units.append(
            {
                "unit": len(units) + 1,
                "title": "其他高頻問題",
                "description": "本單元收錄本批次中未歸入前述分類，但仍具教學價值的高頻標註。",
                "categories": [],
                "entries": remaining_entries,
            }
        )

    for index, unit in enumerate(units, start=1):
        unit["unit"] = index
    return units


def build_high_frequency_list_entries(
    items: list[CorrectionItem],
    skip_categories: set[str] | None = None,
    *,
    min_source_count: int | None = None,
) -> list[dict[str, str]]:
    skip_categories = skip_categories or set()
    min_source_count = min_source_count or teacher_list_source_threshold(effective_essay_count(items=items))
    list_entries = [
        entry
        for entry in build_knowledge_cluster_entries(
            items,
            min_source_count=teacher_correction_source_threshold(effective_essay_count(items=items)),
            list_min_source_count=min_source_count,
        )
        if entry.get("category") in LIST_ENTRY_CATEGORIES and entry.get("category") not in skip_categories
    ]
    list_entries.sort(
        key=lambda entry: (
            int(entry.get("source_count", "0") or 0),
            int(entry.get("frequency", "0") or 0),
            DEDUPE_CATEGORY_PRIORITY.get(entry.get("category", ""), 0),
        ),
        reverse=True,
    )
    return list_entries


def reset_ai_status() -> None:
    AI_ANALYSIS_STATUS.update({"mode": "rules", "provider": "", "model": "", "error": ""})


def should_use_ai(use_ai: bool | None = None) -> bool:
    mode = os.environ.get("MANUAL_ANALYSIS_MODE", "auto").strip().lower()
    if use_ai is not None:
        return bool(use_ai)
    if mode in {"0", "false", "off", "rules", "rule"}:
        return False
    return bool(os.environ.get("GEMINI_API_KEY", "").strip())


def ai_candidate_limit() -> int:
    value = os.environ.get("AI_CANDIDATE_LIMIT", "180").strip()
    try:
        return max(60, min(360, int(value)))
    except ValueError:
        return 180


def build_ai_candidate_pack(
    items: list[CorrectionItem],
    limit: int | None = None,
    *,
    min_source_count: int | None = None,
) -> list[dict[str, object]]:
    limit = limit or ai_candidate_limit()
    effective_count = effective_essay_count(items=items)
    min_source_count = min_source_count or teacher_correction_source_threshold(effective_count)
    usable = build_knowledge_cluster_entries(
        items,
        min_source_count=min_source_count,
        list_min_source_count=teacher_list_source_threshold(effective_count),
    )
    usable.sort(
        key=lambda entry: (
            int(entry.get("source_count", "0")),
            int(entry["frequency"]),
            DEDUPE_CATEGORY_PRIORITY.get(entry.get("category", ""), 0),
            meaningful_token_count(entry["wrong"]),
        ),
        reverse=True,
    )

    per_category_limits = {
        "Word Choice": 36,
        "Grammar": 36,
        "Agreement": 28,
        "Tense": 22,
        "Article": 24,
        "Preposition": 24,
        "Spelling": 18,
        "Punctuation": 18,
        "Capitalization": 8,
        "Style": 18,
        "Redundancy": 8,
    }
    selected: list[dict[str, str]] = []
    selected_ids: set[int] = set()
    for category, category_limit in per_category_limits.items():
        category_entries = [entry for entry in usable if entry.get("category") == category]
        for entry in category_entries[:category_limit]:
            if len(selected) >= limit:
                break
            selected.append(entry)
            selected_ids.add(id(entry))
    for entry in usable:
        if len(selected) >= limit:
            break
        if id(entry) in selected_ids:
            continue
        selected.append(entry)
        selected_ids.add(id(entry))

    candidates: list[dict[str, object]] = []
    for index, entry in enumerate(selected, start=1):
        candidates.append(
            {
                "id": f"K{index:03d}",
                "clusterKey": entry.get("cluster_key", ""),
                "teachingFamily": entry.get("teaching_family", ""),
                "category": entry.get("category", ""),
                "frequency": int(entry.get("frequency", "0") or 0),
                "sourceCount": int(entry.get("source_count", "0") or 0),
                "fileIds": sorted(decode_file_ids(entry.get("file_ids", ""))),
                "wrong": preview_text(entry.get("wrong", ""), 140),
                "correct": preview_text(entry.get("correct", ""), 160),
                "patterns": limit_pattern_text(entry.get("patterns", "")),
                "explanation": preview_text(entry.get("explanation", ""), 260),
                "sourceSentence": preview_text(entry.get("source_sentence", ""), 260),
                "scoreBand": "",
                "isCluster": entry.get("is_cluster") == "true",
                "isList": entry.get("is_list") == "true",
                "memberCount": int(entry.get("member_count", "1") or 1),
                "title": entry.get("title", ""),
                "example": entry.get("example", ""),
                "tip": entry.get("tip", ""),
            }
        )
    return candidates


def build_ai_feedback_pack(records: list[EssayRecord], limit: int = 16) -> list[dict[str, object]]:
    scored = [record for record in records if record.teacher_feedback.strip()]
    scored.sort(key=lambda record: record.essay_score)
    low = scored[: max(4, limit // 3)]
    mid_start = max(0, len(scored) // 2 - 2)
    mid = scored[mid_start : mid_start + 4]
    high = scored[-max(4, limit // 3) :]
    picked: list[EssayRecord] = []
    seen: set[str] = set()
    for record in low + mid + high:
        if record.file_id in seen:
            continue
        seen.add(record.file_id)
        picked.append(record)
        if len(picked) >= limit:
            break
    return [
        {
            "score": record.essay_score,
            "scoreBand": score_band(record.essay_score),
            "feedback": preview_text(record.teacher_feedback, 360),
        }
        for record in picked
    ]


def ai_response_schema() -> dict[str, object]:
    entry = {
        "type": "object",
        "properties": {
            "clusterIds": {"type": "array", "items": {"type": "string"}},
            "candidateIds": {"type": "array", "items": {"type": "string"}},
            "title": {"type": "string"},
            "category": {"type": "string"},
            "rationale": {"type": "string"},
        },
        "required": ["title", "category"],
    }
    unit = {
        "type": "object",
        "properties": {
            "title": {"type": "string"},
            "description": {"type": "string"},
            "entries": {"type": "array", "items": entry},
        },
        "required": ["title", "description", "entries"],
    }
    return {
        "type": "object",
        "properties": {
            "units": {"type": "array", "items": unit},
            "notes": {"type": "string"},
        },
        "required": ["units"],
    }


def extract_gemini_text(payload: dict) -> str:
    candidates = payload.get("candidates") or []
    if not candidates:
        raise RuntimeError("Gemini 沒有回傳 candidates。")
    parts = (((candidates[0] or {}).get("content") or {}).get("parts") or [])
    texts = [part.get("text", "") for part in parts if isinstance(part, dict) and part.get("text")]
    if not texts:
        raise RuntimeError("Gemini 沒有回傳文字內容。")
    return "\n".join(texts).strip()


def parse_json_text(text: str) -> dict:
    value = text.strip()
    if value.startswith("```"):
        value = re.sub(r"^```(?:json)?\s*", "", value)
        value = re.sub(r"\s*```$", "", value)
    try:
        return json.loads(value)
    except json.JSONDecodeError:
        start = value.find("{")
        end = value.rfind("}")
        if start >= 0 and end > start:
            return json.loads(value[start : end + 1])
        raise


def call_gemini_json(prompt: str, schema: dict[str, object]) -> dict:
    api_key = os.environ.get("GEMINI_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("尚未設定 GEMINI_API_KEY。")
    model = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash-lite").strip() or "gemini-2.5-flash-lite"
    url = f"{GEMINI_ENDPOINT}/models/{urllib.parse.quote(model, safe='')}:generateContent?key={urllib.parse.quote(api_key)}"
    body = {
        "systemInstruction": {
            "parts": [
                {
                    "text": (
                        "你是一位台灣高中英文作文教師與教材編輯。"
                        "你的任務是根據批改標註資料整理專屬英語正則手冊，"
                        "必須重視可教學性、規則泛化能力、學生原句可驗證性，不能編造來源。"
                    )
                }
            ]
        },
        "contents": [{"role": "user", "parts": [{"text": prompt}]}],
        "generationConfig": {
            "responseMimeType": "application/json",
            "responseSchema": schema,
            "temperature": float(os.environ.get("GEMINI_TEMPERATURE", "0.2")),
            "maxOutputTokens": int(os.environ.get("GEMINI_MAX_OUTPUT_TOKENS", "12000")),
        },
    }
    request = urllib.request.Request(
        url,
        data=json.dumps(body, ensure_ascii=False).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=int(os.environ.get("GEMINI_TIMEOUT_SECONDS", "120"))) as response:
            payload = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"Gemini API HTTP {exc.code}: {detail[:800]}") from exc
    text = extract_gemini_text(payload)
    debug_path = OUTPUT_DIR / "gemini_raw_response.json.txt"
    try:
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        debug_path.write_text(text, encoding="utf-8")
    except OSError:
        pass
    return parse_json_text(text)


def build_ai_prompt(records: list[EssayRecord], items: list[CorrectionItem], candidates: list[dict[str, object]]) -> str:
    category_counts = Counter(item.category for item in items)
    scores = [record.essay_score for record in records]
    prompt_payload = {
        "batchName": BATCH_NAME,
        "summary": {
            "pdfs": len(records),
            "zeroScore": sum(record.unanswered for record in records),
            "markedItems": len(items),
            "withSourceSentences": sum(1 for item in items if item.source_sentence.strip()),
            "averageScore": round(mean(scores), 2) if scores else 0,
            "medianScore": median(scores) if scores else 0,
            "categories": category_counts.most_common(),
        },
        "candidateClusters": [
            {key: value for key, value in candidate.items() if key != "fileIds"}
            for candidate in candidates
        ],
        "teacherFeedbackSamples": build_ai_feedback_pack(records),
    }
    return (
        "請根據下方 JSON 資料，產出本批次專屬的英語正則手冊架構。\n"
        "要求：\n"
        "1. 產出 2 到 5 個 Unit，每個 Unit 依錯誤邏輯分類，不要只照頻率排序。\n"
        "2. 每個 Unit 優先產出 8 到 12 條；若某 Unit 候選不足，寧可少於 8 條，不得編造。\n"
        "3. 全手冊語言修正條目目標 40 到 50 條；只收高頻、影響理解、可泛化成規則的錯誤。\n"
        "4. 你只需要回傳 Unit、條目標題、category、clusterIds、rationale；不要回傳 wrong/correct/example/tip。\n"
        "5. 每個條目必須引用 1 到 3 個 clusterIds，clusterIds 必須來自 candidateClusters。\n"
        "6. 嚴格禁止把不同錯誤邏輯塞進同一條；不同 category、不同詞性問題、不同句構問題必須拆開。\n"
        "7. 除 Spelling、Capitalization、Punctuation 清單外，建議每條只引用 1 個最具代表性的 clusterId。\n"
        "8. candidateClusters 已經是系統端聚合後的知識點；除清單型項目外，通常每條只引用 1 個 clusterId。\n"
        "9. 不要產出 Writing Strategy、作文架構、段落策略或寫作策略條目；本手冊只收語言正則。\n"
        "10. rationale 使用繁體中文，20-60 字，說明為何這些候選應合併成一條規則。\n"
        "11. Unit title、description、entry title、rationale 都要使用繁體中文，可保留必要英文術語。\n"
        "12. Unit title 不要包含 Unit 編號，例如不要寫 Unit 1、第一單元，只寫「遣詞與慣用搭配」這類標題。\n"
        "13. 拼字、標點、大小寫請集中到清單型條目，不要混入文法句構單元。\n"
        "14. teacherFeedbackSamples 只用來理解本批作文主題，不要把單句文法錯誤硬改成寫作策略。\n"
        "15. candidateClusters 已先排除低頻個案，請優先選 sourceCount 較高的候選。\n\n"
        f"資料：\n{json.dumps(prompt_payload, ensure_ascii=False)}"
    )


def sanitize_ai_text(value: object, fallback: str = "", limit: int = 500) -> str:
    text = compact_cjk(str(value or fallback))
    if len(text) > limit:
        return text[: limit - 3].rstrip() + "..."
    return text


def sanitize_ai_unit_title(value: object, fallback: str = "英語正則") -> str:
    title = sanitize_ai_text(value, fallback, 60)
    title = re.sub(r"^(Unit|UNIT)\s*\d+\s*[:：.\-]?\s*", "", title).strip()
    title = re.sub(r"^第\s*[一二三四五六七八九十\d]+\s*單元\s*[:：.\-]?\s*", "", title).strip()
    return title or fallback


def is_overbroad_ai_rationale(text: str) -> bool:
    value = compact_cjk(text)
    if not value:
        return False
    markers = [
        "冠詞",
        "介系詞",
        "關係代名詞",
        "方位詞",
        "數量詞",
        "存現句",
        "拼寫",
        "大小寫",
        "標點",
        "時態",
        "主詞",
        "代名詞",
        "中式英文",
    ]
    marker_hits = sum(1 for marker in markers if marker in value)
    if "此類錯誤包含" in value and (value.count("、") >= 2 or marker_hits >= 3):
        return True
    if len(value) > 150 and (value.count("、") >= 4 or marker_hits >= 4):
        return True
    return False


def unique_join(values: list[str], limit: int = 4, preview_limit: int = 80) -> str:
    output: list[str] = []
    seen: set[str] = set()
    for value in values:
        text = compact_cjk(value)
        key = normalize_group_key(text)
        if not text or key in seen:
            continue
        seen.add(key)
        output.append(preview_text(text, preview_limit))
        if len(output) >= limit:
            break
    return "；".join(output)


def candidate_frequency(candidate: dict[str, object]) -> int:
    return int(candidate.get("frequency", 0) or 0)


def candidate_source_count(candidate: dict[str, object]) -> int:
    return int(candidate.get("sourceCount", 0) or 0)


def candidate_file_ids(candidate: dict[str, object]) -> set[str]:
    return decode_file_ids(candidate.get("fileIds", []))


def select_representative_candidate(candidates: list[dict[str, object]]) -> dict[str, object]:
    clean_candidates = [
        candidate
        for candidate in candidates
        if not is_dirty_correction(str(candidate.get("correct", "")))
        and not is_unhelpful_pair(str(candidate.get("wrong", "")), str(candidate.get("correct", "")))
    ] or candidates
    return max(
        clean_candidates,
        key=lambda candidate: (
            candidate_frequency(candidate),
            candidate_source_count(candidate),
            meaningful_token_count(str(candidate.get("wrong", ""))),
        ),
    )


def candidates_can_form_list(candidates: list[dict[str, object]]) -> bool:
    categories = {str(candidate.get("category", "")) for candidate in candidates}
    if len(categories) != 1:
        return False
    category = next(iter(categories))
    if category not in LIST_ENTRY_CATEGORIES:
        return False
    if any(candidate.get("isList") for candidate in candidates):
        return True
    if len(candidates) < 2:
        return False
    for candidate in candidates:
        entry = {
            "wrong": str(candidate.get("wrong", "")),
            "correct": str(candidate.get("correct", "")),
            "category": category,
            "frequency": str(candidate_frequency(candidate)),
            "source_count": str(candidate_source_count(candidate)),
        }
        if not is_category_list_candidate(entry, category):
            return False
    return True


def build_ai_list_entry(raw_entry: dict, matched: list[dict[str, object]]) -> dict[str, str]:
    category = str(matched[0].get("category", ""))
    selected = sorted(
        matched,
        key=lambda candidate: (
            candidate_frequency(candidate),
            candidate_source_count(candidate),
            meaningful_token_count(str(candidate.get("wrong", ""))),
        ),
        reverse=True,
    )[:8]
    pair_parts: list[str] = []
    seen_pair_parts: set[str] = set()
    for candidate in selected:
        raw_patterns = str(candidate.get("patterns", "") or "")
        if raw_patterns:
            raw_parts = re.split(r"；|;", raw_patterns)
        else:
            raw_parts = [
                f"{preview_text(str(candidate.get('wrong', '')), 34)} → {preview_text(str(candidate.get('correct', '')), 34)}"
                f"（出現 {candidate_frequency(candidate)} 次）"
            ]
        for part in raw_parts:
            text = compact_cjk(part)
            key = normalize_group_key(strip_frequency_suffix(text))
            if not text or key in seen_pair_parts:
                continue
            seen_pair_parts.add(key)
            pair_parts.append(text)
            if len(pair_parts) >= CLUSTER_PATTERN_LIMIT:
                break
        if len(pair_parts) >= CLUSTER_PATTERN_LIMIT:
            break
    pairs = "；".join(pair_parts)
    file_ids: set[str] = set()
    for candidate in selected:
        file_ids.update(candidate_file_ids(candidate))
    title = sanitize_ai_text(raw_entry.get("title"), MERGED_LIST_TITLES.get(category, "本批常見修正清單"), 90)
    if "清單" not in title:
        title = MERGED_LIST_TITLES.get(category, title)
    rationale = sanitize_ai_text(raw_entry.get("rationale"), "", 120)
    label = CATEGORY_LABELS.get(category, category)
    explanation = rationale or f"這些是可直接合併複習的{label}修正，適合整理成考前檢查清單。"
    if not explanation.endswith(("。", ".", "！", "!", "?", "？")):
        explanation += "。"
    return {
        "title": title,
        "wrong": "；".join(preview_text(str(candidate.get("wrong", "")), 34) for candidate in selected),
        "correct": "；".join(preview_text(str(candidate.get("correct", "")), 34) for candidate in selected),
        "explanation": explanation,
        "example": pairs,
        "tip": CATEGORY_TIPS.get(category, "把這類錯誤整理成固定檢查點，寫完作文後逐句檢查。"),
        "source_sentence": next((str(candidate.get("sourceSentence", "")) for candidate in selected if candidate.get("sourceSentence")), ""),
        "category": category,
        "frequency": str(sum(candidate_frequency(candidate) for candidate in selected)),
        "source_count": str(len(file_ids) if file_ids else sum(candidate_source_count(candidate) for candidate in selected)),
        "file_ids": encode_file_ids(file_ids),
        "patterns": pairs,
        "is_list": "true",
        "is_cluster": "true" if any(candidate.get("isCluster") for candidate in selected) else "false",
        "cluster_key": unique_join([str(candidate.get("clusterKey", "")) for candidate in selected], limit=4, preview_limit=80),
        "teaching_family": unique_join([str(candidate.get("teachingFamily", "")) for candidate in selected], limit=4, preview_limit=80),
        "source_lookup": "disabled",
    }


def build_ai_entry_from_candidates(raw_entry: dict, matched: list[dict[str, object]]) -> dict[str, str]:
    if candidates_can_form_list(matched):
        return build_ai_list_entry(raw_entry, matched)
    representative = select_representative_candidate(matched)
    category = str(representative.get("category", "")) or sanitize_ai_text(raw_entry.get("category"), "", 60)
    frequency = candidate_frequency(representative)
    source_sentence = str(representative.get("sourceSentence", ""))
    file_ids = candidate_file_ids(representative)
    wrong = clean_entry_value(str(representative.get("wrong", "")), 100)
    correct = clean_entry_value(str(representative.get("correct", "")), 120)
    rationale = sanitize_ai_text(raw_entry.get("rationale"), "", 180)
    overbroad_rationale = is_overbroad_ai_rationale(rationale)
    if overbroad_rationale:
        rationale = ""
    base_explanation = sanitize_ai_text(representative.get("explanation"), "", 260)
    explanation_parts = [part for part in [rationale, base_explanation] if part]
    explanation = " ".join(explanation_parts) or make_rule_explanation(
        CorrectionItem("", 0, category, wrong, correct, "", source_sentence, ""),
        max(1, frequency),
    )
    if explanation and not explanation.endswith(("。", ".", "！", "!", "?", "？")):
        explanation += "。"
    example = f"修正方向：{wrong} → {correct}"
    if source_sentence and wrong and correct and "；" not in wrong and "；" not in correct:
        pattern = re.escape(wrong)
        if re.match(r"^[A-Za-z0-9' -]+$", wrong):
            pattern = rf"(?<![A-Za-z0-9']){re.escape(wrong)}(?![A-Za-z0-9'])"
        replaced = re.sub(pattern, correct, source_sentence, count=1, flags=re.I)
        if replaced != source_sentence:
            example = preview_text(replaced, 180)
    title = sanitize_ai_text(raw_entry.get("title"), make_rule_title_from_candidate(representative), 90)
    if overbroad_rationale:
        title = make_rule_title_from_candidate(representative)
    patterns = limit_pattern_text(str(representative.get("patterns", "") or ""))
    candidate_example = sanitize_ai_text(representative.get("example"), "", 220)
    if candidate_example and not patterns:
        example = candidate_example
    return {
        "title": title,
        "wrong": wrong,
        "correct": correct,
        "explanation": explanation,
        "example": example,
        "tip": sanitize_ai_text(representative.get("tip"), "", 180)
        or CATEGORY_TIPS.get(category, "把這類錯誤整理成固定檢查點，寫完作文後逐句檢查。"),
        "source_sentence": source_sentence,
        "category": category,
        "frequency": str(max(1, frequency)),
        "source_count": str(len(file_ids) if file_ids else candidate_source_count(representative)),
        "file_ids": encode_file_ids(file_ids),
        "patterns": patterns,
        "is_cluster": "true" if representative.get("isCluster") else "false",
        "cluster_key": str(representative.get("clusterKey", "")),
        "teaching_family": str(representative.get("teachingFamily", "")),
        "source_lookup": "disabled",
    }


def normalize_ai_manual_units(ai_payload: dict, candidates: list[dict[str, object]]) -> list[dict]:
    candidate_by_id = {str(candidate["id"]): candidate for candidate in candidates}
    units: list[dict] = []
    for raw_unit in ai_payload.get("units", []):
        raw_entries = raw_unit.get("entries", []) if isinstance(raw_unit, dict) else []
        entries: list[dict[str, str]] = []
        for raw_entry in raw_entries:
            if not isinstance(raw_entry, dict):
                continue
            candidate_ids = [str(value) for value in (raw_entry.get("clusterIds") or raw_entry.get("candidateIds") or [])]
            if len(candidate_ids) > 3:
                candidate_ids = candidate_ids[:3]
            matched = [candidate_by_id[candidate_id] for candidate_id in candidate_ids if candidate_id in candidate_by_id]
            if not matched:
                continue
            raw_category = sanitize_ai_text(raw_entry.get("category"), "", 60)
            if raw_category == "Writing Strategy":
                continue
            matched_categories = {str(candidate.get("category", "")) for candidate in matched}
            if len(matched_categories) > 1:
                matched = [select_representative_candidate(matched)]
            elif not candidates_can_form_list(matched):
                matched = [select_representative_candidate(matched)]
            entry = build_ai_entry_from_candidates(raw_entry, matched)
            if entry["wrong"] and entry["correct"] and entry["explanation"]:
                entries.append(entry)
        if entries:
            units.append(
                {
                    "unit": len(units) + 1,
                    "title": sanitize_ai_unit_title(raw_unit.get("title")),
                    "description": sanitize_ai_text(raw_unit.get("description"), "本單元整理本批作文標註中可教學的高頻修正。", 220),
                    "categories": [],
                    "entries": entries[:14],
                }
            )
        if len(units) >= 6:
            break
    if sum(len(unit["entries"]) for unit in units) < 1:
        raise RuntimeError("Gemini 回傳的可用條目太少，已改用規則式歸納。")
    return units


def make_rule_title_from_candidate(candidate: dict[str, object]) -> str:
    if candidate.get("title"):
        return sanitize_ai_text(candidate.get("title"), "高頻修正", 90)
    category = str(candidate.get("category", ""))
    label = CATEGORY_LABELS.get(category, category or "錯誤修正")
    return f"{label}：{preview_text(str(candidate.get('wrong', '')), 30)} → {preview_text(str(candidate.get('correct', '')), 30)}"


def refine_teacher_entry_copy(entry: dict[str, str]) -> dict[str, str]:
    wrong_key = normalize_group_key(entry.get("wrong", ""))
    correct_key = normalize_group_key(entry.get("correct", ""))
    refined = dict(entry)
    if wrong_key == "nature" and correct_key == "natural":
        refined.update(
            {
                "title": "名詞修飾名詞時要確認詞性",
                "explanation": "nature 是名詞；修飾 landscapes 時應使用形容詞 natural。除非是固定複合名詞，否則不要直接把中文的「自然」逐字放到名詞前。",
                "tip": "看到中文「X 的 Y」，先判斷英文中 X 要用名詞、形容詞還是所有格。",
            }
        )
    elif wrong_key == "in north taiwan" and correct_key == "in northern taiwan":
        refined.update(
            {
                "title": "方位名詞修飾地區時用形容詞形式",
                "explanation": "north 是名詞或方位詞；修飾 Taiwan 這類地區名稱時，正式寫作通常用形容詞 northern。",
                "tip": "寫地區方位時，檢查 north/south/east/west 是否要改成 northern/southern/eastern/western。",
            }
        )
    elif wrong_key == "chose" and correct_key == "choose":
        refined.update(
            {
                "title": "描述現在的提名或選擇時用現在式",
                "explanation": "如果文章正在說明現在要提名或選擇哪個景點，應使用現在式 choose，而不是過去式 chose。",
                "tip": "先判斷動作發生在過去經驗，還是目前論述中的選擇。",
            }
        )
    elif wrong_key == "the nature" and correct_key == "nature":
        refined.update(
            {
                "title": "泛指大自然時通常不加 the",
                "explanation": "nature 表示「大自然」這個抽象整體時通常不加 the；只有指特定性質或特定自然環境時才需要冠詞。",
                "tip": "抽象或不可數名詞泛指整體概念時，先檢查是否真的需要 the。",
            }
        )
    return refined


def build_ai_manual_units(records: list[EssayRecord], items: list[CorrectionItem]) -> list[dict]:
    model = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash-lite").strip() or "gemini-2.5-flash-lite"
    AI_ANALYSIS_STATUS.update({"mode": "ai", "provider": "gemini", "model": model, "error": ""})
    effective_count = effective_essay_count(records)
    candidates = build_ai_candidate_pack(
        items,
        min_source_count=teacher_correction_source_threshold(effective_count),
    )
    payload = call_gemini_json(build_ai_prompt(records, items, candidates), ai_response_schema())
    units = normalize_ai_manual_units(payload, candidates)
    existing_list_categories = {
        entry.get("category")
        for unit in units
        for entry in unit.get("entries", [])
        if entry.get("is_list") == "true"
    }
    list_entries = build_high_frequency_list_entries(
        items,
        existing_list_categories,
        min_source_count=teacher_list_source_threshold(effective_count),
    )
    if list_entries:
        units.append(
            {
                "unit": len(units) + 1,
                "title": "拼字、大小寫與標點",
                "description": "本單元整理本批中可合併複習的表層修正清單。",
                "categories": list(LIST_ENTRY_CATEGORIES),
                "entries": list_entries,
            }
        )
    for index, unit in enumerate(units, start=1):
        unit["unit"] = index
    return units


def parse_marked_items(section: str, file_id: str, band: str, essay_text: str) -> list[CorrectionItem]:
    match = MARKED_ITEMS_RE.search(section)
    if not match:
        return []
    end_match = re.search(r"(?:作文\s*單\s*字\s*升\s*級|單\s*字\s*升\s*級)\s*[:：]?", section[match.end() :])
    end = match.end() + end_match.start() if end_match else len(section)
    chunk = section[match.end() : end]
    segments = re.split(r"(?m)(?=^\s*\d+\s*[.、)])", chunk)
    items: list[CorrectionItem] = []
    for segment in segments:
        item_match = re.match(r"\s*(\d+)\s*[.、)]\s*(.*)", segment, flags=re.S)
        if not item_match or "→" not in segment:
            continue
        item_no = int(item_match.group(1))
        body = item_match.group(2).strip()
        wrong, rest = body.split("→", 1)
        wrong = compact_cjk(wrong).strip()
        correction, explanation = split_correction_and_explanation(rest, wrong)
        if not correction or correction == DELETE_CORRECTION:
            recovered = recover_tail_arrow_correction(wrong)
            if recovered:
                wrong, correction, explanation = recovered
        if not wrong or not correction:
            continue
        category = infer_category(wrong, correction, explanation)
        items.append(
            CorrectionItem(
                file_id=file_id,
                item_no=item_no,
                category=category,
                wrong=wrong,
                correction=correction,
                explanation=explanation,
                source_sentence=find_source_sentence(wrong, essay_text),
                score_band=band,
            )
        )
    return items


def emit_progress(progress_callback: ProgressCallback | None, **payload) -> None:
    if progress_callback:
        progress_callback(payload)


def parse_essay_scores(section: str) -> tuple[float, float, float, float, float] | None:
    detailed_re = re.compile(
        r"作文\s*評\s*分[:：]\s*([0-9]+(?:\.[0-9]+)?)\s*/\s*20\s*"
        r"內容\s*[:：]\s*([0-9]+(?:\.[0-9]+)?)\s*分\s*；\s*"
        r"結\s*構\s*[:：]\s*([0-9]+(?:\.[0-9]+)?)\s*分\s*；\s*"
        r"文\s*法\s*[:：]\s*([0-9]+(?:\.[0-9]+)?)\s*分\s*；\s*"
        r"詞\s*彙\s*[:：]\s*([0-9]+(?:\.[0-9]+)?)\s*分"
    )
    match = detailed_re.search(section)
    if match:
        return tuple(float(x) for x in match.groups())  # type: ignore[return-value]

    simple_match = re.search(
        r"(?:作文\s*(?:評\s*分|得\s*分)|寫作\s*得\s*分)\s*[:：]?\s*([0-9]+(?:\.[0-9]+)?)\s*/\s*20",
        section,
    )
    if simple_match:
        score = float(simple_match.group(1))
        return score, 0.0, 0.0, 0.0, 0.0
    return None


def extract_records(progress_callback: ProgressCallback | None = None) -> tuple[list[EssayRecord], list[CorrectionItem]]:
    records: list[EssayRecord] = []
    items: list[CorrectionItem] = []

    pdf_paths = sorted(PDF_DIR.rglob("*.pdf"))
    if not pdf_paths:
        raise RuntimeError(f"No PDF files found under: {PDF_DIR}")

    total = len(pdf_paths)
    emit_progress(progress_callback, stage="analyzing", current=0, total=total, message="開始讀取 PDF 作文區塊")

    for index, path in enumerate(pdf_paths, start=1):
        emit_progress(
            progress_callback,
            stage="analyzing",
            current=index - 1,
            total=total,
            currentFile=path.name,
            message=f"正在分析第 {index} / {total} 份 PDF",
        )
        file_id = pdf_file_id(path)
        text = read_pdf_text(path)
        section = find_essay_section(text)
        if not section:
            raise RuntimeError(f"Essay section not found: {path.name}")
        scores = parse_essay_scores(section)
        if not scores:
            raise RuntimeError(f"Essay score not found: {path.name}")
        score, content, structure, grammar, vocab = scores
        error_count, error_stats = parse_error_stats(section)
        band = score_band(score)
        essay_text = extract_student_essay_text(section)
        teacher_feedback = extract_teacher_feedback(section)
        marked_items = parse_marked_items(section, file_id, band, essay_text)
        items.extend(marked_items)
        records.append(
            EssayRecord(
                file_id=file_id,
                file_name=path.name,
                student_id=infer_student_id(path),
                essay_score=score,
                content_score=content,
                structure_score=structure,
                grammar_score=grammar,
                vocabulary_score=vocab,
                error_count=error_count,
                error_stats=error_stats,
                marked_item_count=len(marked_items),
                unanswered=score == 0,
                teacher_feedback=teacher_feedback,
            )
        )
        emit_progress(
            progress_callback,
            stage="analyzing",
            current=index,
            total=total,
            currentFile=path.name,
            message=f"已完成第 {index} / {total} 份 PDF",
        )
    return records, items


def write_intermediate(records: list[EssayRecord], items: list[CorrectionItem]) -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    with SUMMARY_CSV.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=[
                "file_id",
                "file_name",
                "student_id",
                "essay_score",
                "content_score",
                "structure_score",
                "grammar_score",
                "vocabulary_score",
                "error_count",
                "marked_item_count",
                "unanswered",
                "teacher_feedback",
                "error_stats",
            ],
        )
        writer.writeheader()
        for record in records:
            row = asdict(record)
            row["error_stats"] = json.dumps(record.error_stats, ensure_ascii=False)
            writer.writerow(row)

    with ITEMS_CSV.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=[
                "file_id",
                "item_no",
                "category",
                "wrong",
                "correction",
                "explanation",
                "source_sentence",
                "score_band",
            ],
        )
        writer.writeheader()
        for item in items:
            writer.writerow(asdict(item))

    ITEMS_JSON.write_text(
        json.dumps([asdict(item) for item in items], ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    name: str | None = None,
) -> None:
    font_name = name or PRESET["font"]
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), PRESET["east_asia_font"])
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = PRESET["page"]["width"]
    section.page_height = PRESET["page"]["height"]
    section.top_margin = PRESET["page"]["margin"]
    section.right_margin = PRESET["page"]["margin"]
    section.bottom_margin = PRESET["page"]["margin"]
    section.left_margin = PRESET["page"]["margin"]
    section.header_distance = PRESET["page"]["header_footer"]
    section.footer_distance = PRESET["page"]["header_footer"]

    normal = doc.styles["Normal"]
    normal.font.name = PRESET["font"]
    normal._element.rPr.rFonts.set(qn("w:ascii"), PRESET["font"])
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), PRESET["font"])
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), PRESET["east_asia_font"])
    normal.font.size = Pt(PRESET["body_size"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(PRESET["body_after"])
    normal.paragraph_format.line_spacing = PRESET["body_line"]

    heading_tokens = [("Heading 1", PRESET["h1"]), ("Heading 2", PRESET["h2"]), ("Heading 3", PRESET["h3"])]
    for style_name, token in heading_tokens:
        style = doc.styles[style_name]
        style.font.name = PRESET["font"]
        style._element.rPr.rFonts.set(qn("w:ascii"), PRESET["font"])
        style._element.rPr.rFonts.set(qn("w:hAnsi"), PRESET["font"])
        style._element.rPr.rFonts.set(qn("w:eastAsia"), PRESET["east_asia_font"])
        style.font.size = Pt(token["size"])
        style.font.color.rgb = token["color"]
        style.font.bold = True
        style.paragraph_format.space_before = Pt(token["before"])
        style.paragraph_format.space_after = Pt(token["after"])
        style.paragraph_format.line_spacing = PRESET["body_line"]

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = PRESET["font"]
        style._element.rPr.rFonts.set(qn("w:eastAsia"), PRESET["east_asia_font"])
        style.font.size = Pt(PRESET["body_size"])
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = PRESET["body_line"]


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def setup_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(f"學測作文英語正則｜{BATCH_NAME}")
    set_run_font(run, size=9, color=PRESET["muted"])

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Page ")
    set_run_font(run, size=9, color=PRESET["muted"])
    add_field(footer, "PAGE")


def set_cell_text(cell, text: str, *, bold: bool = False, color: RGBColor | None = None, size: float = 10.5) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], *, indent_dxa: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    tbl.insert(0, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell, **PRESET["cell_margins_dxa"])


def add_title_page(
    doc: Document,
    records: list[EssayRecord],
    items: list[CorrectionItem],
    category_counts: Counter[str],
    manual_units: list[dict],
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("英語正則手冊")
    set_run_font(run, size=12, color=PRESET["muted"], bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"英語正則手冊｜{BATCH_NAME}")
    set_run_font(run, size=25, color=PRESET["ink"], bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    run = p.add_run("Error Correction & Practical English Usage")
    set_run_font(run, size=13, color=PRESET["muted"], italic=True)

    table = doc.add_table(rows=3, cols=2)
    set_table_geometry(table, [2300, 7060])
    rows = [
        ("批次名稱", BATCH_NAME),
        ("內容範圍", "僅整理作文批改紀錄中的英語正則內容"),
        ("使用方式", "先看正則說明與例句，再對照學生原句判斷語境。"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        shade_cell(row.cells[0], PRESET["table_header_fill"])
        set_cell_text(row.cells[0], label, bold=True, color=PRESET["ink"])
        set_cell_text(row.cells[1], value)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("本手冊不列學生姓名、學號或原始檔名；學生原句僅作匿名教學判斷使用。")
    set_run_font(run, size=10, color=PRESET["muted"])


def add_usage_guide(doc: Document) -> None:
    doc.add_heading("英語正則工具書使用攻略", level=1)
    points = [
        "本手冊由本批作文批改紀錄萃取高頻語言正則，供老師課堂講解與學生複習使用。",
        "本批資料中的批改點次數與涉及作文份數僅供排序參考，不等於學生人數。",
        "閱讀每條正則時，先看規則與例句，再看匿名學生原句，確認該錯誤在真實語境中的樣子。",
        "遇到清單型條目時，建議把常見錯法當成考前檢查清單逐項核對。",
    ]
    for point in points:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(point)
        set_run_font(run)


def add_manual_toc(doc: Document, manual_units: list[dict]) -> None:
    doc.add_heading("目錄", level=1)
    table = doc.add_table(rows=1 + len(manual_units), cols=2)
    set_table_geometry(table, [1800, 7560])
    shade_cell(table.rows[0].cells[0], PRESET["table_header_fill"])
    shade_cell(table.rows[0].cells[1], PRESET["table_header_fill"])
    set_cell_text(table.rows[0].cells[0], "Unit", bold=True, color=PRESET["ink"])
    set_cell_text(table.rows[0].cells[1], "內容", bold=True, color=PRESET["ink"])
    for row, unit in zip(table.rows[1:], manual_units):
        set_cell_text(row.cells[0], f"Unit {unit['unit']}", bold=True, color=PRESET["ink"])
        set_cell_text(row.cells[1], f"{unit['title']}：{unit_description_text(unit)}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Error Correction & Practical English Usage【英語正則】")
    set_run_font(run, size=13, bold=True, color=PRESET["ink"])


def add_stats_section(
    doc: Document,
    records: list[EssayRecord],
    items: list[CorrectionItem],
    category_counts: Counter[str],
    manual_units: list[dict],
) -> None:
    scores = [r.essay_score for r in records]
    doc.add_heading("本次作文資料摘要", level=1)
    table = doc.add_table(rows=6, cols=2)
    set_table_geometry(table, [2700, 6660])
    summary_rows = [
        ("PDF 份數", str(len(records))),
        ("未作答或 0 分", str(sum(r.unanswered for r in records))),
        ("作文分數", f"平均 {mean(scores):.2f} / 20；中位數 {median(scores):.1f}；最高 {max(scores):.1f}"),
        ("四項平均", f"內容 {mean(r.content_score for r in records):.2f}；結構 {mean(r.structure_score for r in records):.2f}；文法 {mean(r.grammar_score for r in records):.2f}；詞彙 {mean(r.vocabulary_score for r in records):.2f}"),
        ("可抽出標註", f"{len(items)} 條"),
        ("高頻錯誤類型", "、".join(f"{name} {count}" for name, count in category_counts.most_common(9))),
    ]
    for row, (label, value) in zip(table.rows, summary_rows):
        shade_cell(row.cells[0], PRESET["table_header_fill"])
        set_cell_text(row.cells[0], label, bold=True, color=PRESET["ink"])
        set_cell_text(row.cells[1], value)
    p = doc.add_paragraph()
    manual_entry_count = sum(len(unit["entries"]) for unit in manual_units)
    run = p.add_run(f"統計說明：錯誤類型依 PDF 中的作文錯誤統計彙總；手冊正文由本批 {len(items)} 筆標註自動聚合、去重與排序後，收錄 {manual_entry_count} 條專屬規則。")
    set_run_font(run, size=10, color=PRESET["muted"])


def add_labeled_paragraph(doc: Document, label: str, text: str, *, color: RGBColor | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    label_run = p.add_run(label)
    set_run_font(label_run, bold=True, color=color or PRESET["ink"])
    text_run = p.add_run(text)
    set_run_font(text_run)


def source_match_score(entry: dict[str, str], item: CorrectionItem) -> int:
    if not item.source_sentence:
        return 0
    entry_wrong = normalize_lookup(entry["wrong"])
    item_wrong = normalize_lookup(item.wrong)
    source = normalize_lookup(item.source_sentence)
    score = 0

    if entry_wrong and entry_wrong in source:
        score += 100
    if item_wrong and item_wrong in entry_wrong:
        score += 45
    if entry_wrong and entry_wrong in item_wrong:
        score += 45

    entry_tokens = {tok for tok in entry_wrong.split() if len(tok) >= 4}
    item_tokens = {tok for tok in item_wrong.split() if len(tok) >= 4}
    source_tokens = {tok for tok in source.split() if len(tok) >= 4}
    score += 4 * len(entry_tokens & source_tokens)
    score += 6 * len(entry_tokens & item_tokens)

    correction_norm = normalize_lookup(entry["correct"])
    correction_tokens = {tok for tok in correction_norm.split() if len(tok) >= 4}
    item_correction_tokens = {tok for tok in normalize_lookup(item.correction).split() if len(tok) >= 4}
    score += 2 * len(correction_tokens & item_correction_tokens)
    return score


def find_entry_source_sentence(entry: dict[str, str], items: list[CorrectionItem]) -> str:
    if entry.get("source_lookup") == "disabled":
        return entry.get("source_sentence", "")
    best_item: CorrectionItem | None = None
    best_score = 0
    for item in items:
        score = source_match_score(entry, item)
        if score > best_score:
            best_item = item
            best_score = score
    if best_item and best_score >= 12:
        return best_item.source_sentence
    return entry.get("source_sentence", "")


def clean_teacher_text(text: object, limit: int = 360) -> str:
    value = compact_cjk(str(text or ""))
    value = re.sub(r"本批次(?:中)?同型[^。.!?！？]*[。.!?！？]?", "", value)
    value = re.sub(r"本批(?:老師總評中)?有\s*\d+\s*份[^。.!?！？]*[。.!?！？]?", "", value)
    value = re.sub(r"本批(?:出現|資料|標註)[^。.!?！？]*[。.!?！？]?", "", value)
    value = re.sub(r"由\s*Gemini\s*[^。.!?！？]*[。.!?！？]?", "", value, flags=re.I)
    value = re.sub(r"\s+", " ", value).strip()
    if len(value) > limit:
        value = value[: limit - 3].rstrip() + "..."
    return value


def entry_stat_text(entry: dict[str, str]) -> str:
    frequency = int(entry.get("frequency", "0") or 0)
    source_count = int(entry.get("source_count", "0") or 0)
    if frequency <= 0 and source_count <= 0:
        return ""
    if source_count > 0:
        return f"批改點共 {frequency} 次，涉及 {source_count} 份作文（僅供排序參考，非人次）"
    return f"批改點共 {frequency} 次（僅供排序參考，非人次）"


HEADING_CATEGORY_LABELS = {
    "Word Choice": "遣詞",
    "Style": "表達",
    "Redundancy": "冗詞",
    "Grammar": "文法",
    "Agreement": "一致",
    "Tense": "時態",
    "Article": "冠詞",
    "Preposition": "介系詞",
    "Spelling": "拼字",
    "Capitalization": "大小寫",
    "Punctuation": "標點",
    "Writing Strategy": "策略",
}

MAIN_ERROR_TYPE_LABELS = {
    "Word Choice": "用詞不當",
    "Style": "表達不自然",
    "Redundancy": "冗詞或語意重複",
    "Grammar": "文法或句構錯誤",
    "Agreement": "一致性錯誤",
    "Tense": "時態錯誤",
    "Article": "冠詞或可數性錯誤",
    "Preposition": "介系詞搭配錯誤",
    "Spelling": "拼字錯誤",
    "Capitalization": "大小寫錯誤",
    "Punctuation": "標點錯誤",
    "Writing Strategy": "作文架構與表達策略",
}


def teacher_heading_text(number: int, entry: dict[str, str]) -> str:
    label = HEADING_CATEGORY_LABELS.get(entry.get("category", ""), entry.get("category", "正則"))
    title = clean_teacher_text(entry.get("title", ""), 90) or "高頻修正"
    return f"{number}. 【{label}】 {title}"


def common_error_patterns(entry: dict[str, str]) -> str:
    if entry.get("patterns"):
        return entry["patterns"]
    wrong = clean_teacher_text(entry.get("wrong", ""), 120)
    frequency = int(entry.get("frequency", "0") or 0)
    if wrong and frequency > 0:
        return f"{wrong}（出現 {frequency} 次）"
    return wrong


def correction_contrast_text(entry: dict[str, str]) -> str:
    if entry.get("patterns"):
        return entry["patterns"]
    wrong = clean_teacher_text(entry.get("wrong", ""), 160)
    correct = clean_teacher_text(entry.get("correct", ""), 180)
    if wrong and correct:
        return f"{wrong} → {correct}"
    return ""


def unit_description_text(unit: dict) -> str:
    description = clean_teacher_text(str(unit.get("description", "")), 220)
    note = "頻率僅供排序參考，仍須依文章語境判斷。"
    if note not in description:
        description = f"{description} {note}".strip()
    return description


def is_teacher_list_entry(entry: dict[str, str]) -> bool:
    return entry.get("is_list") == "true"


def is_strategy_entry(entry: dict[str, str]) -> bool:
    return entry.get("category") == "Writing Strategy"


def teacher_entry_frequency(entry: dict[str, str]) -> int:
    return int(entry.get("frequency", "0") or 0)


def teacher_entry_source_count(entry: dict[str, str]) -> int:
    return int(entry.get("source_count", "0") or 0)


def teacher_language_rank(entry: dict[str, str]) -> tuple[int, int, int, int, int]:
    return (
        teacher_entry_source_count(entry),
        teacher_entry_frequency(entry),
        DEDUPE_CATEGORY_PRIORITY.get(entry.get("category", ""), 0),
        1 if is_teacher_list_entry(entry) else 0,
        meaningful_token_count(entry.get("wrong", "")),
    )


def relaxed_teacher_correction_source_threshold(strict_threshold: int) -> int:
    return strict_threshold


def relaxed_teacher_list_source_threshold(strict_threshold: int) -> int:
    return strict_threshold


UNIT_CATEGORY_HINTS = [
    (("遣詞", "搭配", "詞彙", "用詞"), {"Word Choice", "Style", "Redundancy"}),
    (("文法", "句構", "動詞形式"), {"Grammar"}),
    (("一致", "時態", "單複數"), {"Agreement", "Tense"}),
    (("冠詞", "介系詞", "可數"), {"Article", "Preposition"}),
    (("拼字", "大小寫", "標點"), {"Spelling", "Capitalization", "Punctuation"}),
]


def unit_category_set(unit: dict) -> set[str]:
    categories = {str(category) for category in unit.get("categories", []) if str(category)}
    categories.update(str(entry.get("category", "")) for entry in unit.get("entries", []) if entry.get("category"))
    title = str(unit.get("title", ""))
    for hints, hinted_categories in UNIT_CATEGORY_HINTS:
        if any(hint in title for hint in hints):
            categories.update(hinted_categories)
    return {category for category in categories if category}


def strip_frequency_suffix(text: str) -> str:
    return re.sub(r"（\s*出現\s*\d+\s*次\s*）\s*$", "", compact_cjk(text)).strip()


def entry_pair_keys(entry: dict[str, str]) -> set[tuple[str, str, str]]:
    category = entry.get("category", "")
    keys: set[tuple[str, str, str]] = set()
    if entry.get("cluster_key"):
        keys.add((category, f"cluster:{normalize_group_key(entry.get('cluster_key', ''))}", normalize_group_key(entry.get("teaching_family", ""))))
    if (entry.get("patterns") or (entry.get("is_list") == "true" and entry.get("example"))):
        for part in re.split(r"；|;", entry.get("patterns") or entry.get("example", "")):
            if "→" not in part:
                continue
            wrong, correct = part.split("→", 1)
            keys.add((category, normalize_group_key(strip_frequency_suffix(wrong)), normalize_group_key(strip_frequency_suffix(correct))))
    if not keys:
        keys.add((category, normalize_group_key(entry.get("wrong", "")), normalize_group_key(entry.get("correct", ""))))
    return {key for key in keys if key[1] and key[2]}


def teacher_entry_meets_threshold(
    entry: dict[str, str],
    *,
    correction_threshold: int,
    list_threshold: int,
    strategy_threshold: int,
) -> bool:
    if is_strategy_entry(entry):
        return teacher_entry_source_count(entry) >= strategy_threshold
    if is_teacher_list_entry(entry):
        return teacher_entry_source_count(entry) >= list_threshold
    return teacher_entry_source_count(entry) >= correction_threshold


def candidate_can_supplement_unit(entry: dict[str, str], categories: set[str], threshold: int) -> bool:
    category = entry.get("category", "")
    if category not in categories or category == "Writing Strategy":
        return False
    if teacher_entry_source_count(entry) < threshold:
        return False
    if is_teacher_list_entry(entry):
        return category in LIST_ENTRY_CATEGORIES and bool(entry.get("patterns"))
    is_cluster_entry = entry.get("is_cluster") == "true" and bool(entry.get("patterns"))
    if (not is_cluster_entry and is_noisy_wrong(entry.get("wrong", ""))) or is_dirty_correction(entry.get("correct", "")):
        return False
    if is_unhelpful_pair(entry.get("wrong", ""), entry.get("correct", "")):
        return False
    if category in LIST_ENTRY_CATEGORIES:
        return is_category_list_candidate(entry, category)
    return is_main_rule_entry(entry) or teacher_entry_source_count(entry) >= threshold


def clean_supplement_entry(entry: dict[str, str], items: list[CorrectionItem], fallback_sources: list[str], offset: int) -> dict[str, str] | None:
    category = entry.get("category", "")
    family_meta = TEACHING_FAMILY_META.get(entry.get("teaching_family", ""), {})
    wrong = clean_entry_value(entry.get("wrong", ""), 120)
    correct = clean_entry_value(entry.get("correct", ""), 140)
    if not wrong or not correct:
        return None
    is_cluster_entry = entry.get("is_cluster") == "true" and bool(entry.get("patterns"))
    if (not is_cluster_entry and is_noisy_wrong(wrong)) or is_dirty_correction(correct) or is_unhelpful_pair(wrong, correct):
        return None
    source = ensure_entry_source_sentence(entry, items, fallback_sources, offset)
    explanation = clean_teacher_text(entry.get("explanation", ""), 320)
    if not explanation:
        explanation = make_rule_explanation(
            CorrectionItem("", 0, category, wrong, correct, "", source, ""),
            int(entry.get("frequency", "1") or 1),
        )
    example = clean_teacher_text(entry.get("example", ""), 220)
    if not example or "；" in example:
        example = make_rule_example(CorrectionItem("", 0, category, wrong, correct, "", source, ""))
    patterns = limit_pattern_text(entry.get("patterns", ""))
    cleaned_entry = {
        **entry,
        "title": clean_teacher_text(family_meta.get("title") or entry.get("title") or make_rule_title_from_candidate(entry), 90),
        "wrong": wrong,
        "correct": correct,
        "explanation": clean_teacher_text(family_meta.get("explanation") or explanation, 320),
        "example": clean_teacher_text(family_meta.get("example") or example, 220),
        "tip": clean_teacher_text(family_meta.get("tip", ""), 180)
        or clean_teacher_text(entry.get("tip", ""), 180)
        or CATEGORY_TIPS.get(category, "把這類錯誤整理成固定檢查點，寫完作文後逐句檢查。"),
        "source_sentence": source,
        "patterns": patterns,
        "source_count": str(int(entry.get("source_count", "0") or 0)),
        "frequency": str(int(entry.get("frequency", "0") or 0)),
        "source_lookup": "disabled",
    }
    return refine_teacher_entry_copy(cleaned_entry)


def supplement_teacher_units(
    cleaned_units: list[dict],
    items: list[CorrectionItem],
    effective_count: int,
    fallback_sources: list[str],
) -> list[dict]:
    strict_correction_threshold = teacher_correction_source_threshold(effective_count)
    strict_list_threshold = teacher_list_source_threshold(effective_count)
    strategy_threshold = teacher_strategy_source_threshold(effective_count)
    relaxed_correction_threshold = relaxed_teacher_correction_source_threshold(strict_correction_threshold)
    relaxed_list_threshold = relaxed_teacher_list_source_threshold(strict_list_threshold)

    category_to_unit: dict[str, dict] = {}
    standard_units: list[dict] = []
    for group in UNIT_GROUPS:
        unit = {
            **group,
            "unit": len(standard_units) + 1,
            "entries": [],
        }
        standard_units.append(unit)
        for category in group.get("categories", []):
            category_to_unit[category] = unit

    other_entries: list[dict[str, str]] = []
    for unit in cleaned_units:
        for entry in unit.get("entries", []):
            category = entry.get("category", "")
            if category == "Writing Strategy":
                continue
            target_unit = category_to_unit.get(category)
            if target_unit:
                target_unit["entries"].append(entry)
            else:
                other_entries.append(entry)

    if other_entries:
        other_unit = {
            "unit": len(standard_units) + 1,
            "title": "其他高頻問題",
            "description": "本單元整理本批中未歸入主要分類，但仍可作為複習提醒的標註。",
            "categories": [],
            "entries": other_entries,
        }
        standard_units.append(other_unit)

    candidate_pool = [
        entry
        for entry in build_knowledge_cluster_entries(
            items,
            min_source_count=strict_correction_threshold,
            list_min_source_count=strict_list_threshold,
        )
        if entry.get("category") != "Writing Strategy"
    ]
    candidate_pool.sort(key=teacher_language_rank, reverse=True)

    used_keys: set[tuple[str, str, str]] = set()
    for unit in standard_units:
        for entry in unit.get("entries", []):
            used_keys.update(entry_pair_keys(entry))

    supplemented_units: list[dict] = []
    source_offset = 0
    for unit in standard_units:
        categories = unit_category_set(unit)
        unit_entries: list[dict[str, str]] = []
        relaxed_entries: list[dict[str, str]] = []
        for entry in sorted(unit.get("entries", []), key=teacher_language_rank, reverse=True):
            if teacher_entry_meets_threshold(
                entry,
                correction_threshold=strict_correction_threshold,
                list_threshold=strict_list_threshold,
                strategy_threshold=strategy_threshold,
            ):
                unit_entries.append(entry)
            elif is_strategy_entry(entry) and teacher_entry_source_count(entry) > 0:
                relaxed_entries.append(entry)
            elif not is_strategy_entry(entry) and teacher_entry_source_count(entry) >= relaxed_correction_threshold:
                relaxed_entries.append(entry)

        for entry in relaxed_entries:
            if len(unit_entries) >= TEACHER_TARGET_ENTRIES_PER_UNIT:
                break
            unit_entries.append(entry)

        if "Writing Strategy" not in categories:
            for candidate in candidate_pool:
                if len(unit_entries) >= TEACHER_TARGET_ENTRIES_PER_UNIT:
                    break
                category = candidate.get("category", "")
                threshold = relaxed_list_threshold if category in LIST_ENTRY_CATEGORIES else relaxed_correction_threshold
                if not candidate_can_supplement_unit(candidate, categories, threshold):
                    continue
                keys = entry_pair_keys(candidate)
                if used_keys & keys:
                    continue
                cleaned_entry = clean_supplement_entry(candidate, items, fallback_sources, source_offset)
                source_offset += 1
                if not cleaned_entry:
                    continue
                used_keys.update(entry_pair_keys(cleaned_entry))
                unit_entries.append(cleaned_entry)

        if unit_entries:
            unit_entries.sort(key=teacher_language_rank, reverse=True)
            supplemented_units.append(
                {
                    **unit,
                    "unit": len(supplemented_units) + 1,
                    "entries": unit_entries[:TEACHER_MAX_ENTRIES_PER_UNIT],
                }
            )
    return supplemented_units


def filter_high_frequency_teacher_entries(manual_units: list[dict], effective_count: int) -> list[dict]:
    correction_threshold = teacher_correction_source_threshold(effective_count)
    list_threshold = teacher_list_source_threshold(effective_count)
    strategy_threshold = teacher_strategy_source_threshold(effective_count)
    filtered_units: list[dict] = []
    for unit in manual_units:
        unit_entries: list[dict[str, str]] = []
        for entry in unit.get("entries", []):
            if teacher_entry_meets_threshold(
                entry,
                correction_threshold=correction_threshold,
                list_threshold=list_threshold,
                strategy_threshold=strategy_threshold,
            ):
                unit_entries.append(entry)
        if unit_entries:
            filtered_units.append(
                {
                    **unit,
                    "unit": len(filtered_units) + 1,
                    "entries": unit_entries,
                }
            )
    return filtered_units


def limit_teacher_manual_entries(manual_units: list[dict], max_entries: int = TEACHER_MAX_MANUAL_ENTRIES) -> list[dict]:
    total_entries = sum(len(unit.get("entries", [])) for unit in manual_units)
    if total_entries <= max_entries:
        return manual_units

    sorted_units = []
    for unit in manual_units:
        entries = sorted(unit.get("entries", []), key=teacher_language_rank, reverse=True)
        if entries:
            sorted_units.append({**unit, "entries": entries})

    selected_by_unit: dict[int, list[dict[str, str]]] = {index: [] for index in range(len(sorted_units))}
    selected_count = 0
    max_depth = max((len(unit["entries"]) for unit in sorted_units), default=0)
    for depth in range(max_depth):
        for unit_index, unit in enumerate(sorted_units):
            if selected_count >= max_entries:
                break
            if depth >= len(unit["entries"]):
                continue
            selected_by_unit[unit_index].append(unit["entries"][depth])
            selected_count += 1
        if selected_count >= max_entries:
            break

    limited_units: list[dict] = []
    for unit_index, unit in enumerate(sorted_units):
        entries = selected_by_unit.get(unit_index, [])
        if not entries:
            continue
        entries.sort(key=teacher_language_rank, reverse=True)
        limited_units.append({**unit, "unit": len(limited_units) + 1, "entries": entries})
    return limited_units


def fallback_source_sentences(items: list[CorrectionItem]) -> list[str]:
    seen: set[str] = set()
    sources: list[str] = []
    for item in sorted(items, key=lambda candidate: len(candidate.source_sentence), reverse=True):
        source = preview_text(item.source_sentence, 260)
        key = normalize_group_key(source)
        if not source or key in seen:
            continue
        seen.add(key)
        sources.append(source)
        if len(sources) >= 40:
            break
    return sources


def source_supports_wrong(wrong: str, source: str) -> bool:
    wrong_key = normalize_lookup(wrong)
    source_key = normalize_lookup(source)
    if not wrong_key or not source_key:
        return False
    if wrong_key in source_key:
        return True
    wrong_tokens = {token for token in wrong_key.split() if len(token) >= 4}
    if not wrong_tokens:
        return False
    source_tokens = {token for token in source_key.split() if len(token) >= 4}
    return len(wrong_tokens & source_tokens) >= max(2, min(4, len(wrong_tokens)))


def find_supported_source_for_wrong(wrong: str, items: list[CorrectionItem]) -> str:
    for item in items:
        if item.source_sentence and source_supports_wrong(wrong, item.source_sentence):
            return preview_text(item.source_sentence, 260)
    return ""


def ensure_entry_source_sentence(entry: dict[str, str], items: list[CorrectionItem], fallback_sources: list[str], offset: int) -> str:
    source = preview_text(entry.get("source_sentence", ""), 260)
    if source and (
        entry.get("category") == "Writing Strategy"
        or entry.get("is_list") == "true"
        or entry.get("is_cluster") == "true"
        or source_supports_wrong(entry.get("wrong", ""), source)
    ):
        return source
    if entry.get("category") == "Writing Strategy":
        return ""
    if entry.get("category") != "Writing Strategy" and entry.get("is_list") != "true":
        supported = find_supported_source_for_wrong(entry.get("wrong", ""), items)
        if supported:
            return supported
    probe = dict(entry)
    probe.pop("source_lookup", None)
    source = preview_text(find_entry_source_sentence(probe, items), 260)
    if source and (entry.get("category") == "Writing Strategy" or entry.get("is_list") == "true" or source_supports_wrong(entry.get("wrong", ""), source)):
        return source
    if fallback_sources:
        return fallback_sources[offset % len(fallback_sources)]
    return ""


def clean_manual_units_for_teacher(
    manual_units: list[dict],
    items: list[CorrectionItem],
    records: list[EssayRecord] | None = None,
) -> list[dict]:
    effective_count = effective_essay_count(records, items)
    fallback_sources = fallback_source_sentences(items)
    cleaned_units: list[dict] = []
    seen_pairs: set[tuple[str, str, str]] = set()
    source_offset = 0

    for unit in manual_units:
        unit_title = sanitize_ai_unit_title(unit.get("title"), "英語正則")
        description = clean_teacher_text(str(unit.get("description", "")), 180)
        if not description or len(description) < 12 or description.endswith(("由", "中", "本單元由")) or "Gemini" in description or "AI" in description:
            description = "本單元整理本批作文標註中可直接教學與複習的正則。"

        entries: list[dict[str, str]] = []
        for raw_entry in unit.get("entries", []):
            if not isinstance(raw_entry, dict):
                continue
            entry = {str(key): str(value) for key, value in raw_entry.items() if value is not None}
            category = entry.get("category", "")
            if category == "Writing Strategy":
                continue
            is_list = entry.get("is_list") == "true" or (
                category in LIST_ENTRY_CATEGORIES
                and ("；" in entry.get("wrong", "") or "；" in entry.get("correct", "") or entry.get("patterns"))
            )

            if is_list:
                if category not in LIST_ENTRY_CATEGORIES:
                    continue
                family_meta = TEACHING_FAMILY_META.get(entry.get("teaching_family", ""), {})
                patterns = limit_pattern_text(entry.get("patterns") or entry.get("example", ""), limit=10)
                if not patterns:
                    continue
                source = ensure_entry_source_sentence(entry, items, fallback_sources, source_offset)
                source_offset += 1
                cleaned = {
                    **entry,
                    "title": clean_teacher_text(family_meta.get("title") or entry.get("title") or MERGED_LIST_TITLES.get(category, "常見修正清單"), 90),
                    "explanation": clean_teacher_text(family_meta.get("explanation") or entry.get("explanation", ""), 260)
                    or f"這些是可直接合併複習的{CATEGORY_LABELS.get(category, category)}修正，適合整理成考前檢查清單。",
                    "example": clean_teacher_text(family_meta.get("example") or patterns, 220),
                    "patterns": patterns,
                    "tip": clean_teacher_text(family_meta.get("tip"), 180)
                    or clean_teacher_text(entry.get("tip", ""), 180)
                    or CATEGORY_TIPS.get(category, "把這類錯誤整理成固定檢查點，寫完作文後逐句檢查。"),
                    "source_sentence": source,
                    "source_count": str(int(entry.get("source_count", "0") or 0)),
                    "frequency": str(int(entry.get("frequency", "0") or 0)),
                    "is_list": "true",
                    "source_lookup": "disabled",
                }
                entries.append(cleaned)
                continue

            wrong = clean_entry_value(entry.get("wrong", ""), 120)
            correct = clean_entry_value(entry.get("correct", ""), 140)
            if not wrong or not correct:
                continue
            is_cluster_entry = entry.get("is_cluster") == "true" and bool(entry.get("patterns"))
            if (not is_cluster_entry and is_noisy_wrong(wrong)) or is_dirty_correction(correct) or is_unhelpful_pair(wrong, correct):
                continue
            key_entry = {**entry, "wrong": wrong, "correct": correct}
            keys = entry_pair_keys(key_entry)
            if seen_pairs & keys:
                continue
            seen_pairs.update(keys)
            source = ensure_entry_source_sentence(entry, items, fallback_sources, source_offset)
            source_offset += 1
            explanation = clean_teacher_text(entry.get("explanation", ""), 320)
            if not explanation:
                explanation = make_rule_explanation(
                    CorrectionItem("", 0, category, wrong, correct, "", source, ""),
                    int(entry.get("frequency", "1") or 1),
                )
            example = clean_teacher_text(entry.get("example", ""), 220)
            if not example or "；" in example:
                example = make_rule_example(CorrectionItem("", 0, category, wrong, correct, "", source, ""))
            patterns = limit_pattern_text(entry.get("patterns", ""))
            cleaned_entry = {
                **entry,
                "title": clean_teacher_text(entry.get("title") or make_rule_title_from_candidate(entry), 90),
                "wrong": wrong,
                "correct": correct,
                "explanation": explanation,
                "example": example,
                "tip": clean_teacher_text(entry.get("tip", ""), 180)
                or CATEGORY_TIPS.get(category, "把這類錯誤整理成固定檢查點，寫完作文後逐句檢查。"),
                "source_sentence": source,
                "patterns": patterns,
                "source_count": str(int(entry.get("source_count", "0") or 0)),
                "frequency": str(int(entry.get("frequency", "0") or 0)),
                "source_lookup": "disabled",
            }
            entries.append(refine_teacher_entry_copy(cleaned_entry))

        if entries:
            cleaned_units.append(
                {
                    **unit,
                    "unit": len(cleaned_units) + 1,
                    "title": unit_title,
                    "description": description,
                    "entries": entries[:14],
                }
            )

    supplemented_units = supplement_teacher_units(cleaned_units, items, effective_count, fallback_sources)
    return limit_teacher_manual_entries(supplemented_units)


def add_entry(doc: Document, number: int, entry: dict[str, str], items: list[CorrectionItem]) -> None:
    doc.add_heading(teacher_heading_text(number, entry), level=3)
    add_labeled_paragraph(doc, "【正則說明】", entry["explanation"])
    add_labeled_paragraph(doc, "【例】", entry["example"])
    contrast = correction_contrast_text(entry)
    if contrast:
        add_labeled_paragraph(doc, "【不要寫／建議寫】", contrast)
    stats = entry_stat_text(entry)
    if stats:
        add_labeled_paragraph(doc, "【本批資料】", stats, color=PRESET["muted"])
    source_sentence = find_entry_source_sentence(entry, items)
    if source_sentence:
        add_labeled_paragraph(doc, "【學生原句】", source_sentence)
    add_labeled_paragraph(doc, "【小提醒】", entry["tip"])


def add_units(doc: Document, items: list[CorrectionItem], manual_units: list[dict]) -> None:
    doc.add_heading("Error Correction & Practical English Usage【英語正則】", level=1)
    for unit in manual_units:
        doc.add_heading(f"Unit {unit['unit']} {unit['title']}", level=2)
        for index, entry in enumerate(unit["entries"], start=1):
            add_entry(doc, index, entry, items)


def add_appendix(doc: Document, items: list[CorrectionItem]) -> None:
    doc.add_heading("附錄：語言修稿檢查清單", level=1)
    checks = [
        "每個句子是否都有主詞與主要動詞？",
        "敘述個人經驗時，過去式是否一致？",
        "可數名詞單數前是否有 a/an/the 或所有格？",
        "advice、communication、food 等不可數名詞是否誤加 s？",
        "feel 後面是否接形容詞，而不是情緒名詞？",
        "長句是否有 run-on sentence 或 comma splice？",
        "receive、occasion、necessary、business 等高頻拼字是否正確？",
    ]
    for check in checks:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(check)
        set_run_font(run)


def build_docx(records: list[EssayRecord], items: list[CorrectionItem], manual_units: list[dict] | None = None) -> None:
    doc = Document()
    style_document(doc)
    setup_header_footer(doc)
    manual_units = manual_units or build_manual_units(records, items)
    category_counts = Counter()
    for record in records:
        category_counts.update(record.error_stats)

    add_units(doc, items, manual_units)
    doc.save(MANUAL_PATH)


def summarize_run(records: list[EssayRecord], items: list[CorrectionItem], manual_units: list[dict] | None = None) -> dict:
    scores = [record.essay_score for record in records]
    category_counts: Counter[str] = Counter()
    for record in records:
        category_counts.update(record.error_stats)
    if not category_counts:
        category_counts = Counter(item.category for item in items)

    manual_units = manual_units or build_manual_units(records, items)
    bands = Counter(score_band(score) for score in scores)
    band_order = ["0", "1-5", "5.5-10", "10.5-15", "15.5-20"]
    return {
        "batchName": BATCH_NAME,
        "pdfs": len(records),
        "zeroScore": sum(record.unanswered for record in records),
        "markedItems": len(items),
        "withSourceSentences": sum(1 for item in items if item.source_sentence.strip()),
        "manualEntries": sum(len(unit["entries"]) for unit in manual_units),
        "analysisMode": AI_ANALYSIS_STATUS.get("mode", "rules"),
        "aiProvider": AI_ANALYSIS_STATUS.get("provider", ""),
        "aiModel": AI_ANALYSIS_STATUS.get("model", ""),
        "aiError": AI_ANALYSIS_STATUS.get("error", ""),
        "averageScore": round(mean(scores), 2) if scores else 0,
        "medianScore": median(scores) if scores else 0,
        "categories": [{"name": name, "count": count} for name, count in category_counts.most_common()],
        "scoreBands": [{"label": label, "count": bands[label]} for label in band_order],
        "files": {
            "docx": str(MANUAL_PATH),
            "summary_csv": str(SUMMARY_CSV),
            "items_csv": str(ITEMS_CSV),
            "items_json": str(ITEMS_JSON),
            "manual_json": str(OUTPUT_DIR / "manual_units.json"),
        },
    }


def run_pipeline(
    pdf_dir: Path | str,
    output_dir: Path | str,
    batch_name: str,
    progress_callback: ProgressCallback | None = None,
    use_ai: bool | None = None,
) -> dict:
    reset_ai_status()
    configure_paths(pdf_dir, output_dir, batch_name)
    if not PDF_DIR.exists():
        raise RuntimeError(f"PDF folder not found: {PDF_DIR}")
    records, items = extract_records(progress_callback)
    emit_progress(
        progress_callback,
        stage="writing_outputs",
        current=len(records),
        total=len(records),
        message="正在寫入 CSV 與 JSON 中繼資料",
    )
    write_intermediate(records, items)

    manual_units: list[dict]
    if should_use_ai(use_ai):
        emit_progress(
            progress_callback,
            stage="ai_summarizing",
            current=len(records),
            total=len(records),
            message="正在使用 Gemini 歸納本批專屬手冊架構",
        )
        try:
            manual_units = build_ai_manual_units(records, items)
        except Exception as exc:
            AI_ANALYSIS_STATUS.update(
                {
                    "mode": "rules_fallback",
                    "provider": "gemini",
                    "model": os.environ.get("GEMINI_MODEL", "gemini-2.5-flash-lite"),
                    "error": str(exc),
                }
            )
            manual_units = build_manual_units(records, items)
    else:
        manual_units = build_manual_units(records, items)

    manual_units = clean_manual_units_for_teacher(manual_units, items, records)

    (OUTPUT_DIR / "manual_units.json").write_text(
        json.dumps(manual_units, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    emit_progress(
        progress_callback,
        stage="building_docx",
        current=len(records),
        total=len(records),
        message="正在產出 Word 手冊",
    )
    build_docx(records, items, manual_units)
    summary = summarize_run(records, items, manual_units)
    emit_progress(
        progress_callback,
        stage="completed",
        current=len(records),
        total=len(records),
        message="分析完成",
    )
    return summary


def main() -> None:
    summary = run_pipeline(PDF_DIR, OUTPUT_DIR, BATCH_NAME)
    print(json.dumps({"date": date.today().isoformat(), **summary}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
